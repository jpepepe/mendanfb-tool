# -*- coding: utf-8 -*-
"""
求人提案 FB自動生成ツール
使い方: streamlit run 求人提案FB_ツール.py
"""

import streamlit as st
import json, re, zipfile, os, tempfile, unicodedata, io
from pathlib import Path
import anthropic
import requests
from gdrive import upload_json as gdrive_upload_json, upload_proposal_summary as gdrive_upload_proposal_summary

SLACK_WEBHOOK = st.secrets.get("SLACK_WEBHOOK_URL", "") if hasattr(st, 'secrets') else ""

st.set_page_config(
    page_title="求人提案 FB ツール",
    page_icon="💼",
    layout="wide"
)

st.markdown("""
<style>
.score-box { background:#f0f4ff; border-left:4px solid #2E75B6; padding:10px 16px; margin:6px 0; border-radius:4px; }
.score-high { border-left-color:#375623; background:#e2efda; }
.score-mid  { border-left-color:#c55a11; background:#fce4d6; }
.score-low  { border-left-color:#c00000; background:#fcecea; }
.check-ok   { color:#375623; font-weight:bold; }
.check-ng   { color:#c00000; font-weight:bold; }
.check-warn { color:#c55a11; font-weight:bold; }
.section-title { background:#1F3864; color:white; padding:8px 16px; border-radius:6px; font-weight:bold; margin:16px 0 8px 0; }
.phase-title { background:#2E75B6; color:white; padding:7px 16px; border-radius:5px; font-weight:bold; margin:12px 0 6px 0; }
.sub-title { background:#2E75B6; color:white; padding:6px 14px; border-radius:4px; font-weight:bold; margin:10px 0 6px 0; font-size:0.9rem; }
.concern-hit  { background:#e2efda; border-left:4px solid #375623; padding:8px 12px; border-radius:4px; margin:6px 0; }
.concern-miss { background:#fcecea; border-left:4px solid #c00000; padding:8px 12px; border-radius:4px; margin:6px 0; }
.bt-hit  { background:#EBF5FB; border-left:4px solid #2E75B6; padding:8px 12px; border-radius:4px; margin:6px 0; }
.bt-miss { background:#FEF9E7; border-left:4px solid #F39C12; padding:8px 12px; border-radius:4px; margin:6px 0; }
.phrase-box { background:#1F3864; color:white; padding:10px 14px; border-radius:6px; margin:6px 0; font-family:monospace; }
.warning-box { background:#fff3cd; border-left:4px solid #ffc107; padding:10px 16px; margin:8px 0; border-radius:4px; }
.good-flag { background:#e2efda; border-left:4px solid #375623; padding:10px 16px; margin:8px 0; border-radius:4px; }
.rate-good { background:#e2efda; border-left:4px solid #375623; padding:10px 16px; margin:8px 0; border-radius:4px; font-size:1.1rem; }
.rate-bad  { background:#fcecea; border-left:4px solid #c00000; padding:10px 16px; margin:8px 0; border-radius:4px; font-size:1.1rem; }
.phase-skip { background:#f5f5f5; border-left:4px solid #aaa; padding:8px 14px; border-radius:4px; margin:6px 0; color:#888; }
</style>
""", unsafe_allow_html=True)

# ── ヘルパー ──────────────────────────────────────────────
def sq(text: str) -> str:
    t = (text or '').strip()
    if t.startswith('「') and t.endswith('」'):
        return t
    return f'「{t}」'

def score_bar(s, max_s=3):
    return '█' * s + '░' * (max_s - s)

def score_cls(s):
    return 'score-high' if s >= 2 else ('score-mid' if s == 1 else 'score-low')


# ── Word文書生成（求人提案版） ────────────────────────────
def _generate_word_doc(d: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    def h(text, level=1, color=(31,56,100)):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)

    def body(text, bold=False, color=None, size=10.5):
        p = doc.add_paragraph()
        r = p.add_run(str(text))
        r.bold = bold; r.font.size = Pt(size)
        if color: r.font.color.rgb = RGBColor(*color)

    def kv(key, val, kc=(46,117,182)):
        p = doc.add_paragraph()
        r1 = p.add_run(f'{key}：'); r1.bold = True
        r1.font.color.rgb = RGBColor(*kc); r1.font.size = Pt(10.5)
        r2 = p.add_run(str(val or '')); r2.font.size = Pt(10.5)

    def tbl(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
        for i, h_ in enumerate(headers):
            c = t.rows[0].cells[i]; c.text = h_
            if c.paragraphs[0].runs: c.paragraphs[0].runs[0].bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row): cells[i].text = str(v or '')
        doc.add_paragraph()

    gd      = d.get('grip_drivers', {})
    bh      = d.get('behaviors', {})
    ov      = d.get('overall', {})
    p1      = d.get('phase1_opening', {})
    p2      = d.get('phase2_basics', {})
    p3      = d.get('phase3_pitch', {})
    p4      = d.get('phase4_proposal', {})
    p5      = d.get('phase5_agreement', {})
    p6      = d.get('phase6_closing', {})
    ca_an   = d.get('concern_analysis', {})
    phrases = d.get('next_phrases', [])
    axes    = ['求人マッチング説明','動機付け','懸念解消','認識統一','次回接続']
    total   = sum(gd.get(ax,{}).get('score',0) for ax in axes)
    grade   = ov.get('grade','') or ('S' if total>=13 else 'A' if total>=10 else 'B' if total>=7 else 'C' if total>=4 else 'D')
    proposed = d.get('提案件数') or 0
    accepted = d.get('応諾件数') or 0

    t0 = doc.add_heading('求人提案 分析レポート', 0)
    t0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kv('CA名', d.get('ca','')); kv('候補者名', d.get('candidate',''))
    kv('グリップ', d.get('grip','')); kv('提案件数', f'{proposed}件 / 応諾{accepted}件')
    doc.add_paragraph()

    h('■ 1. 総合評価')
    p = doc.add_paragraph()
    r = p.add_run(f'グレード：{grade}　　総合スコア：{total}/15')
    r.bold = True; r.font.size = Pt(14)
    if proposed > 0:
        rate = accepted / proposed * 100
        body(f'応諾率：{rate:.0f}%（{accepted}/{proposed}件）', bold=True,
             color=(55,86,35) if rate >= 50 else (192,0,0))
    if ov.get('grade_reason'): body(ov['grade_reason'])
    if ov.get('top_strength'): body(f'【最大の強み】{ov["top_strength"]}', color=(55,86,35))
    doc.add_paragraph()

    h('■ 2. ルーブリック採点（5軸）')
    tbl(
        ['評価軸','スコア','強み','改善点','次のアクション'],
        [[ax, f'{gd.get(ax,{}).get("score",0)}/3',
          gd.get(ax,{}).get('strength',''),
          gd.get(ax,{}).get('weakness',''),
          gd.get(ax,{}).get('next_action','')] for ax in axes]
    )

    if ov.get('top_issues'):
        h('■ 3. 優先改善ポイント TOP3')
        for i, issue in enumerate(ov['top_issues'][:3], 1):
            if isinstance(issue, dict):
                body(f'{i}. {issue.get("issue","")}', bold=True)
                if issue.get('detail'): body(f'　📌 {issue["detail"]}')
                if issue.get('fix'):    body(f'　💡 改善例：{issue["fix"]}', color=(55,86,35))
        doc.add_paragraph()

    if ov.get('one_thing'):
        h('■ 4. 次の面談で必ず試してほしいこと')
        body(ov['one_thing'], color=(46,117,182), bold=True)
        doc.add_paragraph()

    h('■ 5. フェーズ別評価')
    h('フェーズ1：冒頭・前回振り返り', level=2)
    body(f'振り返り：{"あり" if p1.get("振り返り実施") else "なし"}　温度感確認：{"あり" if p1.get("温度感確認") else "なし"}　メリデメ：{p1.get("メリデメ説明スコア",0)}/3')
    if p1.get('振り返り評価'): body(p1['振り返り評価'])
    if p1.get('改善アドバイス'): body(f'💡 {p1["改善アドバイス"]}', color=(55,86,35))

    h('フェーズ2：転職基礎説明', level=2)
    if not p2.get('対象'):
        body('対象外（転職活動経験あり）')
    else:
        body(f'実施：{"あり" if p2.get("実施") else "なし"}　フロー：{"✅" if p2.get("転職フロー説明") else "❌"}　'
             f'選考数字：{"✅" if p2.get("選考基準数字説明") else "❌"}　必要応募数：{"✅" if p2.get("必要応募数説明") else "❌"}')
        if p2.get('コメント'): body(p2['コメント'])

    h('フェーズ3：提案職種説明（★最重要）', level=2, color=(192,0,0))
    body(f'提案スタイル：{p3.get("提案スタイル","─")}型　ミライ型：{p3.get("ミライ型スコア",0)}/3　マッチング：{p3.get("マッチングスコア",0)}/3'
         + ('　★プロフェッショナル' if p3.get('プロフェッショナル') else ''))
    if p3.get('ミライ型根拠'): body(p3['ミライ型根拠'])
    if p3.get('改善アドバイス'): body(f'💡 {p3["改善アドバイス"]}', color=(55,86,35))

    h('フェーズ4：求人提案', level=2)
    if p4:
        body(f'価値提供：{p4.get("価値提供スコア",0)}/3　難易度伝達：{p4.get("難易度伝達スコア",0)}/3　会話バランス：{p4.get("CA一方向比率スコア",0)}/3')
        if p4.get('改善アドバイス'): body(f'💡 {p4["改善アドバイス"]}', color=(55,86,35))

    h('フェーズ5：応募意思確認（★重要）', level=2, color=(192,0,0))
    if p5:
        body(f'能動性：{p5.get("能動性スコア",0)}/3　応諾の質：{p5.get("応諾の質スコア",0)}/3　応諾後不安：{"あり⚠️" if p5.get("応諾後不安") else "なし"}')
        if p5.get('NGワード検出'):
            body(f'⚠️ NGワード：{p5.get("NGワード内容","")}', color=(192,0,0))
        if p5.get('渋り検出'):
            body(f'渋り検出：{p5.get("渋り場面","")}')
        if p5.get('改善アドバイス'): body(f'💡 {p5["改善アドバイス"]}', color=(55,86,35))

    h('フェーズ6：次回アポ', level=2)
    if p6:
        body(f'面接対策アポ：{"あり" if p6.get("面接対策アポ") else "なし"}　日程具体性：{p6.get("日程具体性","─")}　書類確認：{"あり" if p6.get("書類次ステップ確認") else "なし"}')
        if p6.get('改善アドバイス'): body(f'💡 {p6["改善アドバイス"]}', color=(55,86,35))
    doc.add_paragraph()

    h('■ 6. 懸念深掘り分析')
    if ca_an.get('summary'): body(ca_an['summary'])
    for ms in (ca_an.get('missed_scenes') or [])[:5]:
        if not isinstance(ms, dict): continue
        body(f'求職者：{ms.get("cd_text","")}')
        body(f'実際：{ms.get("ca_actual","")}', color=(192,0,0))
        body(f'→ {ms.get("ca_suggested","")}', color=(55,86,35))
        doc.add_paragraph()
    for gs in (ca_an.get('good_scenes') or [])[:5]:
        if not isinstance(gs, dict): continue
        body(f'✅ 求職者：{gs.get("cd_text","")}')
        body(f'CA：{gs.get("ca_text","")}', color=(55,86,35))
        doc.add_paragraph()

    h('■ 7. 行動指標')
    ok_w = lambda v, thr, rev=False: '✅' if (v <= thr if rev else v >= thr) else '❌'
    tbl(
        ['指標','値','目標','判定'],
        [
            ('求職者発話比率',  f'{round(bh.get("求職者発話比率",0)*100)}%', '35%以上', ok_w(bh.get('求職者発話比率',0)*100, 35)),
            ('懸念スルー率',    f'{bh.get("懸念スルー率",0)}%',              '50%以下', ok_w(bh.get('懸念スルー率',0), 50, rev=True)),
            ('バックトラッキング',f'{bh.get("バックトラッキング",0)}回',      '5回以上', ok_w(bh.get('バックトラッキング',0), 5)),
            ('マッチング説明',  f'{bh.get("マッチング説明数",0)}回',          '2回以上', ok_w(bh.get('マッチング説明数',0), 2)),
            ('名前呼称回数',    f'{bh.get("名前呼称回数",0)}回',              '3回以上', ok_w(bh.get('名前呼称回数',0), 3)),
            ('フィラー回数',    f'{bh.get("フィラー回数",0)}回',              '30回以下',ok_w(bh.get('フィラー回数',0), 30, rev=True)),
            ('ポジティブ反応',  f'{bh.get("ポジティブ反応",0)}回',            '5回以上', ok_w(bh.get('ポジティブ反応',0), 5)),
        ]
    )

    if phrases:
        h('■ 8. 次の面談で使えるフレーズ集')
        for ph in phrases:
            if not isinstance(ph, dict) or not ph.get('phrase'): continue
            body(f'【{ph.get("situation","")}】', bold=True)
            p_ph = doc.add_paragraph()
            r_ph = p_ph.add_run(ph.get('phrase',''))
            r_ph.font.color.rgb = RGBColor(31,56,100)
            r_ph.bold = True; r_ph.font.size = Pt(11)
            if ph.get('why'): body(f'　💡 {ph["why"]}')
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Slack送信（求人提案版） ────────────────────────────────
def _send_slack(d: dict, webhook_url: str) -> tuple:
    gd      = d.get('grip_drivers', {})
    ov      = d.get('overall', {})
    p3      = d.get('phase3_pitch', {})
    p5      = d.get('phase5_agreement', {})
    ca_an   = d.get('concern_analysis', {})
    phrases = d.get('next_phrases', [])

    axes  = ['求人マッチング説明','動機付け','懸念解消','認識統一','次回接続']
    total = sum(gd.get(ax,{}).get('score',0) for ax in axes)
    grade = ov.get('grade','') or ('S' if total>=13 else 'A' if total>=10 else 'B' if total>=7 else 'C' if total>=4 else 'D')
    g_emoji = {'S':'🏆','A':'🟢','B':'🔵','C':'🟡','D':'🔴'}.get(grade,'⚪')

    proposed = d.get('提案件数') or 0
    accepted = d.get('応諾件数') or 0
    rate_str = f'{accepted/proposed*100:.0f}%（{accepted}/{proposed}件）' if proposed else '─'

    score_lines = '\n'.join(
        f"　{ax}：{'█'*gd.get(ax,{}).get('score',0)}{'░'*(3-gd.get(ax,{}).get('score',0))} "
        f"{gd.get(ax,{}).get('score',0)}/3　{gd.get(ax,{}).get('weakness','')}"
        for ax in axes)

    issues_text = ''
    for i, issue in enumerate(ov.get('top_issues',[])[:3], 1):
        if isinstance(issue, dict):
            issues_text += f'{i}. *{issue.get("issue","")}*\n　{issue.get("detail","")}\n　💡 {issue.get("fix","")}\n'

    p3_text = f'スタイル：{p3.get("提案スタイル","─")}型　ミライ型：{p3.get("ミライ型スコア","─")}/3　マッチング：{p3.get("マッチングスコア","─")}/3'
    if p3.get('プロフェッショナル'): p3_text += '　★プロフェッショナル'

    p5_text = f'能動性：{p5.get("能動性スコア","─")}/3　応諾の質：{p5.get("応諾の質スコア","─")}/3'
    if p5.get('渋り検出'):
        p5_text += f'\n　渋り検出あり → アウト返し効果：{p5.get("アウト返し効果","─")}'

    concern_text = ca_an.get('summary','')
    for ms in (ca_an.get('missed_scenes') or [])[:2]:
        if not isinstance(ms, dict): continue
        concern_text += f'\n　スルー：「{ms.get("cd_text","")}」\n　→ `{ms.get("ca_suggested","")}`'

    phrases_text = ''
    for ph in phrases:
        if not isinstance(ph, dict) or not ph.get('phrase'): continue
        phrases_text += f'*{ph.get("situation","")}*\n　`{ph.get("phrase","")}` — {ph.get("why","")}\n'

    blocks = [
        {"type": "header", "text": {"type": "plain_text",
            "text": f'{g_emoji} 求人提案FB｜{d.get("ca","")}CA × {d.get("candidate","")}　グリップ{d.get("grip","")}'}},
        {"type": "section", "fields": [
            {"type": "mrkdwn", "text": f"*グレード:* {grade}　総合スコア：{total}/15"},
            {"type": "mrkdwn", "text": f"*応諾率:* {rate_str}"},
        ]},
        {"type": "section", "text": {"type": "mrkdwn", "text": f"*📋 グレード判定理由*\n{ov.get('grade_reason','─')}"}},
        {"type": "section", "text": {"type": "mrkdwn", "text": f"*🎯 ルーブリック（5軸）*\n{score_lines}"}},
        {"type": "section", "text": {"type": "mrkdwn", "text": f"*⭐ フェーズ3*\n{p3_text}"}},
        {"type": "section", "text": {"type": "mrkdwn", "text": f"*⭐ フェーズ5*\n{p5_text}"}},
        {"type": "section", "text": {"type": "mrkdwn", "text": f"*💪 強み*\n{ov.get('top_strength','─')}"}},
        {"type": "section", "text": {"type": "mrkdwn", "text": f"*🎯 優先改善ポイント*\n{issues_text or '─'}"}},
    ]
    if concern_text:
        blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": f"*🔍 懸念深掘り*\n{concern_text}"}})
    if ov.get('one_thing'):
        blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": f"*🚀 次に試すこと*\n{ov['one_thing']}"}})
    if phrases_text:
        blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": f"*🗣️ フレーズ集*\n{phrases_text.strip()}"}})

    try:
        resp = requests.post(webhook_url, json={"blocks": blocks}, timeout=10)
        if resp.status_code == 200:
            return True, '送信成功'
        return False, f'エラー: {resp.status_code} {resp.text}'
    except Exception as e:
        return False, str(e)

# ── 求人提案専用パターン ─────────────────────────────────
FILLER_PAT   = re.compile(r'えー+|あのー+|えっと+|まあ+|うーん+|んー+|あー+|なんか(?!ら)|ちょっと(?=\s|、|。|$)')
CONCERN_PAT  = re.compile(r'難し|不安|心配|どうかな|考えたい|迷って|迷っている|微妙|きつい|きつそう|遠い|遠くて|給料.*低|年収.*低|条件.*合わ|残業.*多|勤務地.*遠|転勤.*ある|転勤.*嫌|ちょっと.*難|ちょっと.*不安|ちょっと.*心配|ちょっと.*迷|ちょっと.*考え')
INTEREST_PAT = re.compile(r'いいです|面白そう|気になる|ぜひ|やってみたい|好き|魅力|惹か|チャレンジ|いいかも|ありです|前向き')
AGREE_PAT    = re.compile(r'応募します|お願いします|やります|ぜひ.*お願い|進めて|受けてみ|出してみ|書類.*出|送って')
BT_PAT       = re.compile(r'おっしゃって|つまり|要は|ということ|まとめると|先ほど|さっき|先ほどの|おっしゃった.*が|言ってた|言われてた')
POS_PAT      = re.compile(r'すごい|いいです|素晴らしい|さすが|面白い|いいですね|なるほど.*そう|それは.*いい|それって.*すごい')
NEXT_PAT     = re.compile(r'面接|書類|日程|スケジュール|次回|別日|改めて.*お時間|推薦|応募|出し|送り|LINE|ライン|また.*お話')
MATCH_PAT    = re.compile(r'おっしゃって.*いた|以前.*おっしゃ|条件.*合って|ご希望.*通り|おすすめ|向いて|ぴったり|マッチ|合って|合います|合うかな')

OUTPUT_JSON = Path(__file__).parent / "output" / "json_proposal"
OUTPUT_UTT  = Path(__file__).parent / "output" / "utterances_proposal"
OUTPUT_JSON.mkdir(parents=True, exist_ok=True)
OUTPUT_UTT.mkdir(parents=True, exist_ok=True)

# ── ファイル名パース ─────────────────────────────────────
def parse_filename(stem: str) -> dict:
    stem = unicodedata.normalize('NFC', stem)
    parts = re.split(r'[_＿]', re.sub(r'-グリップ', '_グリップ', stem))
    result = {'ca': '', 'grip': '未入力', 'candidate': '', '提案件数': None, '応諾件数': None}
    for i, p in enumerate(parts):
        m = re.search(r'グリップ([ABCD])', p)
        if m:
            result['ca']   = parts[0] if i > 0 else ''
            result['grip'] = m.group(1)
            if i + 1 < len(parts): result['candidate'] = parts[i + 1]
    m_prop = re.search(r'提案(\d+)', stem)
    m_acc  = re.search(r'応諾(\d+)', stem)
    if m_prop: result['提案件数'] = int(m_prop.group(1))
    if m_acc:  result['応諾件数'] = int(m_acc.group(1))
    return result

# ── 話者分離 ──────────────────────────────────────────────
def detect_speakers(full_text):
    from collections import Counter
    cands = re.findall(r'(?<![^\s。！？\n])([぀-鿿a-zA-Zー]{2,10}):\s', full_text)
    freq = Counter(cands)
    return [n for n, c in freq.most_common(10) if c >= 3][:2]

def parse_docx(file_bytes):
    from xml.etree import ElementTree as ET
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        f.write(file_bytes); tmp = f.name
    try:
        with zipfile.ZipFile(tmp) as z:
            with z.open('word/document.xml') as f:
                tree = ET.parse(f)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        paragraphs = tree.getroot().findall('.//w:p', ns)
        full_text = ''
        for p in paragraphs:
            line = ''.join(t.text or '' for t in p.findall('.//w:t', ns)).strip()
            if line: full_text += line + '\n'
        full_text = full_text.replace('：', ':')
        full_text = re.sub(r'^\d{2}:\d{2}:\d{2}\n', '', full_text, flags=re.MULTILINE)
        full_text = re.sub(r'^\d{4}年.+\n', '', full_text, flags=re.MULTILINE)
        full_text = re.sub(r'^会議\s.+\n', '', full_text, flags=re.MULTILINE)
        speakers = detect_speakers(full_text)
        if not speakers: return [], full_text, 'docx'
        pat = re.compile(r'(?:' + '|'.join(re.escape(s) for s in speakers) + r'):\s*')
        utterances = []
        matches = list(pat.finditer(full_text))
        for i, m in enumerate(matches):
            speaker = m.group().rstrip(': \t')
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
            text = re.sub(r'\s+', ' ', full_text[start:end]).strip()
            if text: utterances.append({'speaker': speaker, 'text': text})
        return utterances, full_text, 'docx'
    finally:
        os.unlink(tmp)

def parse_txt_with_haiku(file_bytes, ca_name, cand_name, client):
    raw = file_bytes.decode('utf-8', errors='ignore')
    raw = re.sub(r'^\d+\t', '', raw, flags=re.MULTILINE)
    prompt = f"""以下は人材紹介会社の求人提案面談の音声文字起こしです。話者は2名です：
- CA（キャリアアドバイザー）: {ca_name}
- 求職者: {cand_name}
各発言をCAまたは求職者に割り当て、JSON配列のみ返してください。
形式: [{{"speaker": "CA", "text": "発話内容"}}, ...]
文字起こし:
{raw[:12000]}"""
    resp = client.messages.create(
        model='claude-haiku-4-5-20251001', max_tokens=16000,
        messages=[{'role': 'user', 'content': prompt}])
    content = re.sub(r'```(?:json)?\s*', '', resp.content[0].text.strip()).strip()
    try:
        parsed = json.loads(content)
        if isinstance(parsed, list) and parsed: return parsed, raw, 'txt'
    except: pass
    return [{'speaker': '不明', 'text': raw}], raw, 'txt'

# ── 行動指標計算（軽量版） ────────────────────────────────
def compute_metrics(utterances, ca_name, cand_name):
    ca_base   = re.sub(r'さん$', '', ca_name)
    cand_base = re.sub(r'さん$', '', cand_name)
    all_spk   = set(u['speaker'] for u in utterances)
    ca_spk    = {s for s in all_spk if ca_base in s or s == 'CA'}
    cd_spk_n  = {s for s in all_spk if cand_base in s or s == '求職者'}
    cd_spk    = cd_spk_n if cd_spk_n else (all_spk - ca_spk - {'不明'})

    def is_ca(s): return s in ca_spk
    def is_cd(s): return s in cd_spk

    ca_texts = [u['text'] for u in utterances if is_ca(u['speaker'])]
    cd_texts = [u['text'] for u in utterances if is_cd(u['speaker'])]
    all_ca   = ' '.join(ca_texts)
    all_cd   = ' '.join(cd_texts)
    total    = len(all_ca) + len(all_cd)

    # 会話回数
    turns, prev = 0, None
    for u in utterances:
        if len(u['text'].replace(' ', '')) < 10: continue
        sp = 'CA' if is_ca(u['speaker']) else ('CD' if is_cd(u['speaker']) else None)
        if sp and sp != prev: turns += 1; prev = sp

    # 名前呼称
    nm_base = re.sub(r'さん$', '', cand_name)
    nm_pat  = re.compile(re.escape(nm_base) + r'さん')
    name_cnt = sum(len(nm_pat.findall(t)) for t in ca_texts)

    # 後半比率
    n = len(utterances)
    latter = utterances[int(n * 0.5):]
    la_ca  = sum(len(u['text']) for u in latter if is_ca(u['speaker']))
    la_cd  = sum(len(u['text']) for u in latter if is_cd(u['speaker']))
    la_tot = la_ca + la_cd

    # 懸念ワード場面
    concern_moments_all    = []
    concern_slip_examples  = []
    concern_drill_examples = []
    for i, u in enumerate(utterances[:-1]):
        if is_ca(u['speaker']): continue
        m = CONCERN_PAT.search(u['text'])
        if not m or len(u['text']) < 8: continue
        nxt = utterances[i + 1]
        if not is_ca(nxt['speaker']): continue
        # 深掘りの簡易判定（疑問文・共感・掘り下げワード）
        drilled = bool(re.search(r'[？?]|なぜ|どう|もう少し|具体的|詳しく|どんな|どのあたり|どういう', nxt['text']))
        record = {
            'idx': i,
            'cd_text': u['text'][:80],
            'concern_word': m.group(),
            'ca_response': nxt['text'][:80],
            'drilled': drilled,
        }
        concern_moments_all.append(record)
        if drilled and len(concern_drill_examples) < 3:
            concern_drill_examples.append((u['text'][:70], nxt['text'][:80]))
        elif not drilled and len(concern_slip_examples) < 4:
            concern_slip_examples.append((u['text'][:70], nxt['text'][:80]))

    concern_total = len(concern_moments_all)
    concern_skip  = sum(1 for r in concern_moments_all if not r['drilled'])
    slip_rate     = round(concern_skip / concern_total * 100) if concern_total > 0 else 0

    # マッチング場面
    match_moments = []
    for i, u in enumerate(utterances):
        if not is_ca(u['speaker']): continue
        if MATCH_PAT.search(u['text']) and len(u['text']) > 15:
            match_moments.append({'ca_text': u['text'][:100], 'idx': i})

    # バックトラッキング
    backtrack_moments = []
    for i, u in enumerate(utterances):
        if not is_ca(u['speaker']): continue
        if BT_PAT.search(u['text']) and len(u['text']) > 10:
            referenced = ''
            for j in range(i - 1, max(0, i - 6), -1):
                if is_cd(utterances[j]['speaker']):
                    referenced = utterances[j]['text'][:60]; break
            backtrack_moments.append({
                'ca_text': u['text'][:100],
                'referenced': referenced,
                'idx': i,
            })

    # クロージング
    closing_start   = int(len(utterances) * 0.8)
    closing_ca_utts = [u['text'] for u in utterances[closing_start:] if is_ca(u['speaker']) and len(u['text']) > 10]
    closing_cd_utts = [u['text'] for u in utterances[closing_start:] if is_cd(u['speaker']) and len(u['text']) > 10]

    return {
        '求職者発話比率':    round(len(all_cd) / total, 3) if total > 0 else 0,
        '後半求職者比率':    round(la_cd / la_tot, 3) if la_tot > 0 else 0,
        '会話回数':          turns,
        '名前呼称回数':      name_cnt,
        'フィラー回数':      len(FILLER_PAT.findall(all_ca)),
        'バックトラッキング': sum(1 for t in ca_texts if BT_PAT.search(t)),
        'マッチング説明数':  len(match_moments),
        'ポジティブ反応':    sum(1 for t in ca_texts if POS_PAT.search(t)),
        '次回接続発話数':    sum(1 for t in ca_texts[-max(1, len(ca_texts) // 4):] if NEXT_PAT.search(t)),
        '懸念場面数':        concern_total,
        '懸念スルー数':      concern_skip,
        '懸念スルー率':      slip_rate,
        # 詳細（表示用）
        '懸念モーメント全件':  concern_moments_all,
        '懸念スルー例':        concern_slip_examples,
        '懸念深掘り例':        concern_drill_examples,
        'バックトラッキングモーメント': backtrack_moments,
        'マッチングモーメント': match_moments,
        'クロージングCA':      closing_ca_utts[-3:],
        'クロージングCD':      closing_cd_utts[-2:],
        # 内部用
        'ca_texts':  ca_texts,
        'cd_texts':  cd_texts,
        'utterances': utterances,
        'is_ca':     is_ca,
        'is_cd':     is_cd,
    }

# ── Claude Call 1: フェーズ1〜3 + 5軸スコア + 総合評価 ──
def phase_early_analysis(utterances, ca_name, cand_name, fmt,
                         proposed: int, accepted: int, client):
    transcript = '\n'.join(f"[{u['speaker']}] {u['text']}" for u in utterances)[:15000]
    prop_str = f"{proposed}件提案 / {accepted}件応諾" if proposed else "提案件数不明"

    prompt = f"""あなたは人材紹介会社のトレーナーです。以下は「求人提案面談」の文字起こしです。
フェーズ1〜3の評価と、5軸スコア・総合評価をJSONで返してください。

## スコア基準（0〜3点）
- 3: 根拠引用あり + 求職者の明示的な同意・前向きな反応が確認できる
- 2: 把握できているが確認が弱い or 求職者が「はい」止まり
- 1: 触れているが浅い・一方的な説明のみ
- 0: 未実施

## メタ情報
CA名: {ca_name} / 求職者名: {cand_name} / 形式: {fmt}
提案結果: {prop_str}

## 面談文字起こし
{transcript}

## 出力（JSONのみ・余計な文章不要）
{{
  "phase1_opening": {{
    "振り返り実施": true,
    "振り返り評価": "前回面談の内容をどう引き継いだか（1〜2文）",
    "温度感確認": true,
    "温度感評価": "求職者の気持ち変化を確認できたか（1文）",
    "メリデメ説明スコア": 0,
    "メリデメ説明コメント": "提案方向性のメリデメをどう伝えたか（1文）",
    "改善アドバイス": "フェーズ1で改善すべき点（なければ空文字）"
  }},
  "phase2_basics": {{
    "対象": true,
    "実施": false,
    "転職フロー説明": false,
    "選考基準数字説明": false,
    "必要応募数説明": false,
    "タイミング評価": "適切/遅い/なし",
    "コメント": "転職基礎説明の評価（初期フェーズ求職者でない場合は「対象外」）"
  }},
  "phase3_pitch": {{
    "提案スタイル": "A",
    "提案スタイル根拠": "A（職種軸）またはB（求人一つずつ）と判定した根拠（引用）",
    "提案方向数": 1,
    "ミライ型スコア": 0,
    "ミライ型根拠": "求人票型/ミライ型の判定根拠（引用）。ミライ型の例：将来像・業界成長・活躍イメージ。求人票型の例：残業時間・給与・福利厚生の読み上げ",
    "ミライ型具体例": "ミライ型説明の引用（あれば）",
    "市場トレンド説明": false,
    "市場トレンドコメント": "業界成長性・AI時代等のトレンドに触れたか（1文）",
    "マッチングスコア": 0,
    "マッチングスコア根拠": "0=読み上げのみ/1=条件面のみ/2=意向or適正と紐づけ/3=意向×適正×条件すべて。判定根拠を引用",
    "プロフェッショナル": false,
    "プロフェッショナル根拠": "★：意向と適正が一致する適職として言語化し求職者が気づいた場面（あれば引用）",
    "転職動機紐づけスコア": 0,
    "転職動機紐づけコメント": "求職者の転職動機・不安・ワクワクと職種を結びつけたか（1文）",
    "ワクワク反応スコア": 0,
    "ワクワク反応コメント": "求職者の自発的な質問・前向き反応があったか（1文）",
    "改善アドバイス": "フェーズ3で最も改善すべき点（例文付き）"
  }},
  "grip_drivers": {{
    "求人マッチング説明": {{"score":0,"evidence":[],"strength":"","weakness":"","next_action":""}},
    "動機付け":          {{"score":0,"evidence":[],"strength":"","weakness":"","next_action":""}},
    "懸念解消":          {{"score":0,"evidence":[],"strength":"","weakness":"","next_action":""}},
    "認識統一":          {{"score":0,"evidence":[],"strength":"","weakness":"","next_action":""}},
    "次回接続":          {{"score":0,"evidence":[],"strength":"","weakness":"","next_action":""}}
  }},
  "overall": {{
    "grade": "A",
    "grade_reason": "グレード判定の理由を1文で",
    "top_strength": "この面談の最大の強みを具体的に1文で",
    "top_issues": [
      {{"issue":"課題1のタイトル","detail":"何が問題だったか","fix":"代わりにこう言うべきだった（例文付き）"}},
      {{"issue":"課題2のタイトル","detail":"","fix":""}},
      {{"issue":"課題3のタイトル","detail":"","fix":""}}
    ],
    "one_thing": "次の面談で必ず1つだけ試してほしいこと（例文付き）"
  }},
  "notes": "文字起こし品質・話者分離の懸念など（なければ空文字）"
}}

## 判定基準
grade: S=全軸2.5以上+ミライ型+懸念深掘り◎, A=総合スコア10以上, B=7〜9, C=4〜6, D=3以下
phase2対象: 求職者が「転職活動を始めたばかり」「初めての転職」「転職の流れがわからない」等の発言がある場合にtrue
提案スタイルA: 「機械電気系・設備保全・施工管理の3軸で提案します」等、職種軸で整理して提案
提案スタイルB: 「1社目はスーレスコさんで、次に〜」等、求人を一社ずつ順番に説明"""

    try:
        resp = client.messages.create(
            model='claude-sonnet-4-6', max_tokens=6000,
            messages=[{'role': 'user', 'content': prompt}])
        content = re.sub(r'```(?:json)?\s*', '', resp.content[0].text.strip()).strip()
        if resp.stop_reason == 'max_tokens':
            content = content.rstrip(',\n ')
            for _ in range(10): content += '}'
            content += ']}}'
        try:
            return json.loads(content)
        except:
            m = re.search(r'\{.*\}', content, re.DOTALL)
            if m:
                try: return json.loads(m.group())
                except: pass
    except Exception:
        pass
    return {}

# ── Claude Call 2: フェーズ4〜6 + フレーズ集 ─────────────
def phase_late_analysis(utterances, ca_name, cand_name, client):
    transcript = '\n'.join(f"[{u['speaker']}] {u['text']}" for u in utterances)[:14000]

    prompt = f"""あなたは人材紹介会社の求人提案面談のコーチです。以下の文字起こしを分析し、
フェーズ4〜6の評価とフレーズ集をJSONで返してください。

## メタ情報
CA名: {ca_name} / 求職者名: {cand_name}

## 面談文字起こし
{transcript}

## 出力（JSONのみ・余計な文章不要）
{{
  "phase4_proposal": {{
    "価値提供スコア": 0,
    "価値提供コメント": "求人票の読み上げか、未来ビジョン型か。スコア根拠を引用（1〜2文）",
    "価値提供良い例": "ミライ型・価値提供型の引用（あれば）",
    "難易度伝達スコア": 0,
    "難易度伝達コメント": "選考難易度を数字で正直に伝えたか（1文）",
    "CA一方向比率スコア": 0,
    "CA一方向比率コメント": "CAが一方的に話し続けていないか（1文）",
    "提案不足理由": "",
    "改善アドバイス": "フェーズ4で最も改善すべき点（例文付き）"
  }},
  "phase5_agreement": {{
    "応募確認スタイル": "一社ずつ",
    "能動性スコア": 0,
    "能動性コメント": "求職者が受動的な「はい」だけか、自発的な質問・意見・懸念を出していたか（1文）",
    "応諾の質スコア": 0,
    "応諾の質コメント": "流された応諾か、納得した上での応諾か（1文）",
    "応諾後不安": false,
    "応諾後不安コメント": "辞退リスクにつながる不安が残っているか（1文）",
    "覚悟の醸成": false,
    "覚悟の醸成コメント": "能動的な行動を引き出せたか（1文）",
    "NGワード検出": false,
    "NGワード内容": "検出したNGワード（例：一緒に頑張りましょう）",
    "自発的質問数": 0,
    "渋り検出": false,
    "渋り場面": "求職者の渋り発話の引用（あれば）",
    "実際に会ってみないとアウト返し": "できた/なし",
    "面接練習論アウト返し": "できた/なし",
    "アウト返し効果": "効果あり/なし/スルー",
    "アウト返しコメント": "渋り対応の評価（1〜2文）",
    "改善アドバイス": "フェーズ5で最も改善すべき点（例文付き）"
  }},
  "phase6_closing": {{
    "面接対策アポ": false,
    "日程具体性": "具体日時/曖昧/なし",
    "書類次ステップ確認": false,
    "コメント": "クロージング全体の評価（1〜2文）",
    "改善アドバイス": "フェーズ6で改善すべき点（なければ空文字）"
  }},
  "concern_analysis": {{
    "summary": "懸念深掘り全体の評価（2〜3文）",
    "missed_scenes": [
      {{
        "cd_text": "スルーされた求職者の懸念発話（引用）",
        "concern_word": "懸念ワード",
        "ca_actual": "実際のCAの返し（引用）",
        "ca_suggested": "こう返すべきだった（具体的な例文）",
        "why": "なぜこの返しが良いか（1文）"
      }}
    ],
    "good_scenes": [
      {{
        "cd_text": "うまく対応できた求職者の懸念発話（引用）",
        "ca_text": "CAの深掘り・解消発話（引用）",
        "why_good": "なぜ良かったか（1文）"
      }}
    ]
  }},
※missed_scenesは最も改善インパクトが大きい代表的な場面を上限5件、good_scenesは最も参考になる場面を上限5件に厳選すること。
  "next_phrases": [
    {{"situation": "懸念ワードが出た直後", "phrase": "「〜さんがおっしゃる〜という点、もう少し聞かせてもらえますか？」", "why": "なぜこのフレーズが効くか"}},
    {{"situation": "求人と求職者を紐づける場面", "phrase": "「〜さんが以前おっしゃっていた〜という軸で、この求人は〜」", "why": ""}},
    {{"situation": "渋りが出たとき", "phrase": "「実際に会ってみないとわからない、は本当にそうで。でも書類が通れば必ず会えますよね。まず1歩だけ踏み出してみませんか？」", "why": ""}},
    {{"situation": "応募意思を確認するタイミング", "phrase": "「今日お話しして、この求人いかがでしたか？率直に教えてください」", "why": ""}},
    {{"situation": "クロージングで次のステップを確認", "phrase": "「では書類はいつまでにご準備できそうですか？一緒にスケジュール決めましょう」", "why": ""}}
  ]
}}

## 判定基準
価値提供スコア: 0=項目の読み上げのみ, 1=条件説明のみ, 2=部分的に未来ビジョン, 3=全体的にミライ型・活躍イメージで説明
能動性スコア: 0=求職者が「はい」のみ, 1=たまに質問あり, 2=複数の自発的質問・意見, 3=積極的に懸念・条件確認・質問を出している
応諾の質スコア: 0=流れで「はい」と言わされている, 1=理解はしているが受動的, 2=納得して応諾, 3=求職者自身が「受けてみたい」と能動的に発言
渋り検出: 「ちょっと難しいかな」「どうかな」「考えたい」「実際会ってみないと」等の消極的反応"""

    try:
        resp = client.messages.create(
            model='claude-sonnet-4-6', max_tokens=8000,
            messages=[{'role': 'user', 'content': prompt}])
        content = re.sub(r'```(?:json)?\s*', '', resp.content[0].text.strip()).strip()
        if resp.stop_reason == 'max_tokens':
            content = content.rstrip(',\n ')
            for _ in range(10): content += '}'
            content += ']}}'
        try:
            return json.loads(content)
        except:
            m = re.search(r'\{.*\}', content, re.DOTALL)
            if m:
                try: return json.loads(m.group())
                except: pass
    except Exception:
        pass
    return {}

# ── 保存 ─────────────────────────────────────────────────
def save_results(ca, grip, candidate, fmt, proposed, accepted,
                 utterances, metrics, early_result, late_result=None):
    safe_grip = grip if grip != '未入力' else 'X'
    key = f"{ca}_{safe_grip}_{candidate}_求人提案_提案{proposed or 'X'}_応諾{accepted or 'X'}"

    utt_path = OUTPUT_UTT / f"{key}.json"
    speakers = list(set(u['speaker'] for u in utterances))
    utt_data = {
        "ca": ca, "grip": safe_grip, "candidate": candidate,
        "meeting_type": "求人提案", "format": fmt,
        "提案件数": proposed, "応諾件数": accepted,
        "speakers_detected": speakers,
        "utterance_count": len(utterances),
        "utterances": utterances,
    }
    utt_path.write_text(json.dumps(utt_data, ensure_ascii=False, indent=2), encoding='utf-8')
    try:
        gdrive_upload_json(f"{key}.json", utt_data, subfolder="utterances_proposal")
    except Exception as e:
        st.warning(f"話者分離データのDrive保存失敗（ローカルには保存済み）: {e}")

    bh_calc = {k: v for k, v in metrics.items()
               if k not in ('ca_texts', 'cd_texts', 'utterances', 'is_ca', 'is_cd',
                            '懸念モーメント全件', '懸念スルー例', '懸念深掘り例',
                            'バックトラッキングモーメント', 'マッチングモーメント',
                            'クロージングCA', 'クロージングCD')}

    phase3 = early_result.get('phase3_pitch', {})
    phase5 = (late_result or {}).get('phase5_agreement', {})
    phase6 = (late_result or {}).get('phase6_closing', {})

    # ダッシュボード用指標
    dashboard = {
        "提案スタイル":         phase3.get('提案スタイル', ''),
        "提案方向数":           phase3.get('提案方向数', None),
        "ミライ型スコア":       phase3.get('ミライ型スコア', None),
        "マッチングスコア":     phase3.get('マッチングスコア', None),
        "求職者発話比率":       metrics.get('求職者発話比率', None),
        "能動性スコア":         phase5.get('能動性スコア', None),
        "懸念スルー率":         metrics.get('懸念スルー率', None),
        "渋り検出":             phase5.get('渋り検出', None),
        "アウト返し実施":       phase5.get('実際に会ってみないとアウト返し', '') != 'なし',
        "フェーズ2実施":        early_result.get('phase2_basics', {}).get('実施', False),
        "次回アポ具体性":       phase6.get('日程具体性', ''),
        "5軸合計": sum(
            early_result.get('grip_drivers', {}).get(k, {}).get('score', 0)
            for k in ['求人マッチング説明', '動機付け', '懸念解消', '認識統一', '次回接続']
        ),
        "グレード":             early_result.get('overall', {}).get('grade', ''),
    }

    json_data = {
        "ca": ca, "grip": safe_grip, "candidate": candidate,
        "meeting_type": "求人提案", "format": fmt,
        "提案件数": proposed, "応諾件数": accepted,
        "phase1_opening":  early_result.get('phase1_opening', {}),
        "phase2_basics":   early_result.get('phase2_basics', {}),
        "phase3_pitch":    early_result.get('phase3_pitch', {}),
        "phase4_proposal": (late_result or {}).get('phase4_proposal', {}),
        "phase5_agreement":(late_result or {}).get('phase5_agreement', {}),
        "phase6_closing":  (late_result or {}).get('phase6_closing', {}),
        "grip_drivers":    early_result.get('grip_drivers', {}),
        "overall":         early_result.get('overall', {}),
        "concern_analysis":(late_result or {}).get('concern_analysis', {}),
        "next_phrases":    (late_result or {}).get('next_phrases', []),
        "behaviors":       {**bh_calc},
        "dashboard":       dashboard,
        "notes":           early_result.get('notes', ''),
    }
    json_path = OUTPUT_JSON / f"{key}.json"
    json_path.write_text(json.dumps(json_data, ensure_ascii=False, indent=2), encoding='utf-8')
    # Google Drive（共有ドライブ）にも保存
    try:
        gdrive_upload_json(f"{key}.json", json_data, subfolder="json_proposal")
        # 新規保存時はサマリーキャッシュを無効化（次回ダッシュボード表示時に再生成）
        try:
            gdrive_upload_proposal_summary([])  # 空で上書き → 再構築トリガー
        except Exception:
            pass
    except Exception as e:
        st.warning(f"Google Drive保存失敗（ローカルには保存済み）: {e}")
    return utt_path, json_path


# ═══════════════════════════════════════════════════════
# UI
# ═══════════════════════════════════════════════════════
st.title('💼 求人提案 自動FB ツール')
st.caption('求人提案面談の文字起こしを読み込んで、フェーズ別に採点・フィードバックを生成します')

api_key = os.environ.get('ANTHROPIC_API_KEY', '')

with st.sidebar:
    st.header('⚙️ 設定')
    if not api_key:
        st.error('⚠️ APIキー未設定\n`source ~/.zshrc` 後に再起動')
    else:
        st.success('✅ APIキー設定済み')
    st.divider()
    st.subheader('面談情報（自動入力）')
    ca_input    = st.text_input('CA名', placeholder='例：下川')
    cand_input  = st.text_input('求職者名', placeholder='例：岡寺さん')
    grip_input  = st.selectbox('グリップランク（任意）', ['未入力', 'A', 'B', 'C', 'D'])
    prop_input  = st.number_input('提案件数', min_value=0, max_value=50, value=0)
    acc_input   = st.number_input('応諾件数', min_value=0, max_value=50, value=0)
    st.divider()
    st.caption('💡 ファイル名規則: CA名_グリップA_求職者名_求人提案_提案N_応諾N')

# ── ファイル入力 ──────────────────────────────────────────
file_bytes = None
filename   = None

uploaded = st.file_uploader(
    'ファイルをアップロード（.docx または .txt）', type=['docx', 'txt'])
if uploaded:
    fn_nfc    = unicodedata.normalize('NFC', uploaded.name)
    stem      = fn_nfc.rsplit('.', 1)[0]
    meta      = parse_filename(stem)
    if not ca_input:   ca_input   = meta['ca']
    if not cand_input: cand_input = meta['candidate']
    if grip_input == '未入力' and meta['grip'] != '未入力': grip_input = meta['grip']
    if prop_input == 0 and meta['提案件数']: prop_input = meta['提案件数']
    if acc_input == 0 and meta['応諾件数']:  acc_input  = meta['応諾件数']
    file_bytes = uploaded.read()
    filename   = fn_nfc

# ── メタ情報表示 & 分析実行 ───────────────────────────────
if file_bytes and filename:
    fn_nfc = unicodedata.normalize('NFC', filename)

    col1, col2, col3, col4, col5 = st.columns(5)
    col1.metric('CA名',    ca_input or '未入力')
    col2.metric('求職者名', cand_input or '未入力')
    col3.metric('グリップ', grip_input)
    col4.metric('提案件数', f'{prop_input}件')
    col5.metric('応諾件数', f'{acc_input}件')

    # 警告・フラグ
    if prop_input == 1 and (acc_input == 0 or acc_input == 1):
        st.markdown(
            '<div class="warning-box">⚠️ <b>提案1件・応諾1件（または0件）の場合：</b>'
            '実際は1件のみ提案したパターンか、管理表未記載の可能性があります。実際の件数を確認してください。</div>',
            unsafe_allow_html=True)

    if prop_input > 0:
        rate = acc_input / prop_input * 100
        if rate == 100 and prop_input >= 6:
            st.markdown(
                '<div class="good-flag">🌟 <b>優良ケースフラグ：</b>'
                f'応諾率100%（{acc_input}/{prop_input}件）。ベテラン型の理想形です。</div>',
                unsafe_allow_html=True)
        elif rate == 0 and prop_input > 0:
            st.markdown(
                '<div class="warning-box">🚨 <b>応諾率0%フラグ：</b>全件断られたケースです。</div>',
                unsafe_allow_html=True)
        else:
            rate_cls = 'rate-good' if rate >= 50 else 'rate-bad'
            st.markdown(
                f'<div class="{rate_cls}">📊 応諾率: <b>{rate:.0f}%</b>　'
                f'（{acc_input}件 / {prop_input}件）</div>',
                unsafe_allow_html=True)

    # ── 途中失敗時の再開ボタン ────────────────────────────
    if '_proposal_retry' in st.session_state:
        _retry = st.session_state['_proposal_retry']
        client_retry = anthropic.Anthropic(api_key=api_key) if api_key else None

        if 'early_result' in _retry:
            st.warning('⚠️ フェーズ4〜6の分析が未完了です。下のボタンで再実行できます。')
            if st.button('🔄 フェーズ4〜6だけ再実行', type='primary', use_container_width=True):
                with st.spinner('🔍 フェーズ4〜6を再実行中...'):
                    try:
                        late_result = phase_late_analysis(
                            _retry['utterances'], _retry['ca'], _retry['candidate'], client_retry)
                        late_result = late_result or {}
                    except Exception as e:
                        st.error(f'再実行も失敗しました：{e}'); st.stop()
                utt_path, json_path = save_results(
                    _retry['ca'], _retry['grip'], _retry['candidate'], _retry['fmt'],
                    _retry['proposed'], _retry['accepted'],
                    _retry['utterances'], _retry['metrics'], _retry['early_result'], late_result)
                st.session_state['_proposal_res'] = {
                    'early_result': _retry['early_result'],
                    'late_result':  late_result,
                    'metrics':      _retry['metrics'],
                    'ca':           _retry['ca'],
                    'grip':         _retry['grip'],
                    'candidate':    _retry['candidate'],
                    'proposed':     _retry['proposed'],
                    'accepted':     _retry['accepted'],
                    'behaviors':    _retry['metrics'],
                    'ref_file':     json_path.name,
                }
                st.session_state.pop('_proposal_retry', None)
                st.rerun()
        else:
            st.warning('⚠️ フェーズ1〜3の分析が未完了です。下のボタンで最初から再実行できます。')
            if st.button('🔄 最初から再実行', type='primary', use_container_width=True):
                with st.spinner('🤖 フェーズ1〜3を再実行中...'):
                    early_result = phase_early_analysis(
                        _retry['utterances'], _retry['ca'], _retry['candidate'], _retry['fmt'],
                        _retry['proposed'], _retry['accepted'], client_retry)
                if not early_result.get('overall'):
                    st.error('再実行も失敗しました。しばらく待ってから再試行してください。'); st.stop()
                with st.spinner('🔍 フェーズ4〜6を分析中...'):
                    try:
                        late_result = phase_late_analysis(
                            _retry['utterances'], _retry['ca'], _retry['candidate'], client_retry)
                        late_result = late_result or {}
                    except Exception:
                        late_result = {}
                utt_path, json_path = save_results(
                    _retry['ca'], _retry['grip'], _retry['candidate'], _retry['fmt'],
                    _retry['proposed'], _retry['accepted'],
                    _retry['utterances'], _retry['metrics'], early_result, late_result)
                st.session_state['_proposal_res'] = {
                    'early_result': early_result,
                    'late_result':  late_result,
                    'metrics':      _retry['metrics'],
                    'ca':           _retry['ca'],
                    'grip':         _retry['grip'],
                    'candidate':    _retry['candidate'],
                    'proposed':     _retry['proposed'],
                    'accepted':     _retry['accepted'],
                    'behaviors':    _retry['metrics'],
                    'ref_file':     json_path.name,
                }
                st.session_state.pop('_proposal_retry', None)
                st.rerun()
        st.divider()

    if st.button('🔍 分析開始', type='primary', use_container_width=True):
        if not api_key:
            st.error('APIキーが設定されていません。'); st.stop()
        if not ca_input or not cand_input:
            st.error('CA名と求職者名を入力してください'); st.stop()

        client = anthropic.Anthropic(api_key=api_key)

        with st.spinner('📄 ファイルを解析中...'):
            if 'docx' in fn_nfc.lower():
                utterances, raw_text, fmt = parse_docx(file_bytes)
            else:
                utterances, raw_text, fmt = parse_txt_with_haiku(
                    file_bytes, ca_input, cand_input, client)

        if not utterances or (len(utterances) == 1 and utterances[0]['speaker'] == '不明'):
            st.error('話者分離に失敗しました。CA名・求職者名を確認してください。'); st.stop()

        speakers_detected = list(set(u['speaker'] for u in utterances))
        st.success(f'✅ 話者分離完了：{len(utterances)}発話 / 検出話者：{speakers_detected}')

        with st.spinner('📊 行動指標を計算中...'):
            metrics = compute_metrics(utterances, ca_input, cand_input)

        with st.spinner('🤖 フェーズ1〜3 + 5軸スコアを分析中...'):
            early_result = phase_early_analysis(
                utterances, ca_input, cand_input, fmt,
                prop_input or None, acc_input or None, client)

        if not early_result.get('overall'):
            # フェーズ1〜3が失敗 → retry情報を保存して再開ボタンを出す
            st.session_state['_proposal_retry'] = {
                'utterances': utterances, 'metrics': metrics,
                'ca': ca_input, 'grip': grip_input, 'candidate': cand_input,
                'fmt': fmt, 'proposed': prop_input or None, 'accepted': acc_input or None,
            }
            st.error('⚠️ フェーズ1〜3の分析に失敗しました。上の「再実行」ボタンで再試行できます。')
            st.rerun()

        with st.spinner('🔍 フェーズ4〜6（求人提案・応募確認・クロージング）を分析中...'):
            try:
                late_result = phase_late_analysis(utterances, ca_input, cand_input, client)
                late_result = late_result or {}
            except Exception:
                late_result = {}

        if not late_result.get('phase4_proposal'):
            # フェーズ4〜6が失敗 → early_resultは保存してretry情報に載せる
            st.session_state['_proposal_retry'] = {
                'utterances': utterances, 'metrics': metrics,
                'ca': ca_input, 'grip': grip_input, 'candidate': cand_input,
                'fmt': fmt, 'proposed': prop_input or None, 'accepted': acc_input or None,
                'early_result': early_result,
            }
            # フェーズ1〜3の結果だけ先に保存
            save_results(
                ca_input, grip_input, cand_input, fmt,
                prop_input or None, acc_input or None,
                utterances, metrics, early_result, {})
            st.warning('⚠️ フェーズ4〜6の分析が未完了です。上の「再実行」ボタンでフェーズ4〜6だけ再試行できます。')
            st.rerun()

        utt_path, json_path = save_results(
            ca_input, grip_input, cand_input, fmt,
            prop_input or None, acc_input or None,
            utterances, metrics, early_result, late_result)
        st.success(f'💾 保存完了　📊 分析結果：`{json_path.name}`')

        st.session_state['_proposal_res'] = {
            'early_result': early_result,
            'late_result':  late_result,
            'metrics':      metrics,
            'ca':           ca_input,
            'grip':         grip_input,
            'candidate':    cand_input,
            'proposed':     prop_input or None,
            'accepted':     acc_input or None,
            'behaviors':    metrics,
            'ref_file':     json_path.name,
        }
        st.session_state.pop('_proposal_retry', None)
        st.rerun()

# ── 結果表示（session_stateから） ────────────────────────
if '_proposal_res' in st.session_state:
    _res         = st.session_state['_proposal_res']
    early_result = _res['early_result']
    late_result  = _res['late_result']
    metrics      = _res['metrics']
    ca_input     = _res['ca']
    grip_input   = _res['grip']
    cand_input   = _res['candidate']


    # ═══════════════════════════════════════════════
    # 結果表示
    # ═══════════════════════════════════════════════
    st.divider()
    gd      = early_result.get('grip_drivers', {})
    overall = early_result.get('overall', {})
    total_score = sum(gd.get(k, {}).get('score', 0)
                      for k in ['求人マッチング説明', '動機付け', '懸念解消', '認識統一', '次回接続'])

    # ── グレード ──────────────────────────────────
    grade = overall.get('grade', '─')
    grade_colors = {'S': '#1a5276', 'A': '#1e8449', 'B': '#2471a3', 'C': '#d35400', 'D': '#c0392b'}
    grade_color  = grade_colors.get(grade, '#555')
    st.markdown(
        f'<div style="background:{grade_color};color:white;padding:16px 20px;'
        f'border-radius:10px;margin-bottom:12px">'
        f'<span style="font-size:2.4rem;font-weight:bold">{grade}</span>'
        f'&nbsp;&nbsp;<span style="font-size:1rem">{overall.get("grade_reason","")}</span>'
        f'</div>', unsafe_allow_html=True)

    # ── KPI ───────────────────────────────────────
    phase3 = early_result.get('phase3_pitch', {})
    phase5 = late_result.get('phase5_agreement', {}) if late_result else {}

    c1, c2, c3, c4, c5, c6 = st.columns(6)
    c1.metric('総合スコア', f'{total_score}/15',
              delta='良好' if total_score >= 10 else ('要改善' if total_score < 7 else '普通'))
    if prop_input > 0:
        rate = acc_input / prop_input * 100
        c2.metric('応諾率', f'{rate:.0f}%',
                  delta=('✅良好' if rate >= 50 else '⚠️目標50%以上'), delta_color='off')
    else:
        c2.metric('応諾率', '─')
    mts = phase3.get('ミライ型スコア', '─')
    c3.metric('ミライ型スコア', f'{mts}/3' if isinstance(mts, int) else '─',
              delta=('✅ミライ型' if isinstance(mts, int) and mts >= 2 else '⚠️求人票型'), delta_color='off')
    sl = metrics['懸念スルー率']
    c4.metric('懸念スルー率', f'{sl}%',
              delta=('✅良好' if sl <= 50 else '❌要改善'), delta_color='off')
    cr = metrics['求職者発話比率']
    c5.metric('求職者発話比率', f'{cr:.0%}',
              delta=('✅目標達成' if cr >= 0.35 else '⚠️目標35%'), delta_color='off')
    style = phase3.get('提案スタイル', '─')
    c6.metric('提案スタイル', f'{style}型' if style in ('A', 'B') else '─',
              delta=('✅職種軸型' if style == 'A' else ('⚠️求人一つずつ型' if style == 'B' else '')), delta_color='off')

    st.divider()

    # ── 面談サマリー ──────────────────────────────
    st.markdown('<div class="section-title">📋 面談サマリー</div>', unsafe_allow_html=True)
    col_s, col_i = st.columns([1, 1])
    with col_s:
        st.markdown('**💪 この面談の最大の強み**')
        st.success(overall.get('top_strength', '─'))
    with col_i:
        st.markdown('**🎯 優先改善ポイント（TOP3）**')
        for idx, issue in enumerate(overall.get('top_issues', [])[:3]):
            if isinstance(issue, dict):
                with st.expander(f'**{idx + 1}. {issue.get("issue", "")}**', expanded=(idx == 0)):
                    if issue.get('detail'): st.write('📌 ' + issue['detail'])
                    if issue.get('fix'):    st.code(issue['fix'], language=None)
            else:
                st.error(f'{idx + 1}. {issue}')

    one_thing = overall.get('one_thing', '')
    if one_thing:
        st.markdown('**🚀 次の面談で必ず1つ試してほしいこと**')
        st.warning(one_thing)

    st.divider()

    # ════════════════════════════════════════════════
    # フェーズ1：冒頭・前回振り返り
    # ════════════════════════════════════════════════

    # ルーブリック採点（5軸）
    # ════════════════════════════════════════════════
    st.markdown('<div class="section-title">🎯 ルーブリック採点（5軸）</div>', unsafe_allow_html=True)
    AXES = ['求人マッチング説明', '動機付け', '懸念解消', '認識統一', '次回接続']
    AXIS_LABELS = {
        '求人マッチング説明': '求人マッチング説明（WILL/CAN/MUSTに紐づけて説明できたか）',
        '動機付け':          '動機付け（求職者が前向きになったか）',
        '懸念解消':          '懸念解消（不安・迷いを引き出して対処できたか）',
        '認識統一':          '認識統一（どの求人に応募するか合意できたか）',
        '次回接続':          '次回接続（面接日程・書類等の次ステップを確認したか）',
    }
    for axis in AXES:
        d   = gd.get(axis, {})
        s   = d.get('score', 0)
        cls = score_cls(s)
        ev  = d.get('evidence', [])
        bar = score_bar(s)
        na  = d.get('next_action', '')
        html = (
            f'<div class="score-box {cls}">'
            f'<b>{AXIS_LABELS[axis]}</b> &nbsp; <code>{bar}</code> <b>{s}/3</b>'
            + (f'<br>💪 {d["strength"]}' if d.get('strength') else '')
            + (f'<br>📌 {d["weakness"]}' if d.get('weakness') else '')
            + (f'<br><small style="color:#595959">根拠：{sq(ev[0])}</small>' if ev else '')
            + (f'<br>🚀 <b>次のアクション：</b>{na}' if na else '')
            + '</div>'
        )
        st.markdown(html, unsafe_allow_html=True)

    st.divider()

    # ════════════════════════════════════════════════

    # 行動指標（数値詳細）
    # ════════════════════════════════════════════════
    st.markdown('<div class="section-title">📊 行動指標（数値詳細）</div>', unsafe_allow_html=True)
    bh_data = [
        ('求職者発話比率',     f'{metrics["求職者発話比率"]:.0%}',   '目標35%以上',  metrics['求職者発話比率'] >= 0.35),
        ('後半求職者比率',     f'{metrics["後半求職者比率"]:.0%}',   '目標40%以上',  metrics['後半求職者比率'] >= 0.4),
        ('会話回数（ラリー）', f'{metrics["会話回数"]}回',           '参考値',        True),
        ('名前呼称回数',       f'{metrics["名前呼称回数"]}回',       '目標3回以上',  metrics['名前呼称回数'] >= 3),
        ('フィラー回数',       f'{metrics["フィラー回数"]}回',       '目標30回以下', metrics['フィラー回数'] <= 30),
        ('バックトラッキング', f'{metrics["バックトラッキング"]}回', '目標5回以上',  metrics['バックトラッキング'] >= 5),
        ('マッチング説明',     f'{metrics["マッチング説明数"]}回',   '目標2回以上',  metrics['マッチング説明数'] >= 2),
        ('ポジティブ反応',     f'{metrics["ポジティブ反応"]}回',     '目標5回以上',  metrics['ポジティブ反応'] >= 5),
        ('懸念スルー率',       f'{metrics["懸念スルー率"]}%',        '目標50%以下',  metrics['懸念スルー率'] <= 50),
    ]
    cols = st.columns(3)
    for i, (label, val, target, ok) in enumerate(bh_data):
        with cols[i % 3]:
            icon = '✅' if ok else '⚠️'
            st.metric(f'{icon} {label}', val, help=target)

    st.divider()
    st.markdown('<div class="section-title">フェーズ別評価</div>', unsafe_allow_html=True)
    p1 = early_result.get('phase1_opening', {})
    with st.expander('📌 フェーズ1：冒頭・前回振り返り', expanded=True):
        c1, c2, c3 = st.columns(3)
        c1.metric('前回振り返り', '✅あり' if p1.get('振り返り実施') else '❌なし')
        c2.metric('温度感確認', '✅あり' if p1.get('温度感確認') else '❌なし')
        mds = p1.get('メリデメ説明スコア', 0)
        c3.metric('メリデメ説明', f'{mds}/3',
                  delta='✅良好' if mds >= 2 else ('⚠️' if mds == 1 else '❌なし'), delta_color='off')
        if p1.get('振り返り評価'):
            st.write('📋 ' + p1['振り返り評価'])
        if p1.get('温度感評価'):
            st.write('🌡️ ' + p1['温度感評価'])
        if p1.get('メリデメ説明コメント'):
            st.write('⚖️ ' + p1['メリデメ説明コメント'])
        if p1.get('改善アドバイス'):
            st.info('💡 改善: ' + p1['改善アドバイス'])

    # ════════════════════════════════════════════════
    # フェーズ2：転職基礎説明
    # ════════════════════════════════════════════════
    p2 = early_result.get('phase2_basics', {})
    with st.expander('📌 フェーズ2：転職基礎説明（条件付き）', expanded=False):
        if not p2.get('対象', False):
            st.markdown('<div class="phase-skip">この求職者は転職初期フェーズではないため、フェーズ2は対象外です。</div>',
                        unsafe_allow_html=True)
        else:
            c1, c2, c3, c4 = st.columns(4)
            c1.metric('実施', '✅あり' if p2.get('実施') else '❌なし')
            c2.metric('転職フロー説明', '✅' if p2.get('転職フロー説明') else '❌')
            c3.metric('選考基準数字', '✅' if p2.get('選考基準数字説明') else '❌')
            c4.metric('必要応募数', '✅' if p2.get('必要応募数説明') else '❌')
            timing = p2.get('タイミング評価', '')
            if timing:
                st.write(f'⏱️ タイミング: **{timing}**')
            if p2.get('コメント'):
                st.write('📋 ' + p2['コメント'])

    # ════════════════════════════════════════════════
    # フェーズ3：提案職種説明 ★最重要
    # ════════════════════════════════════════════════
    p3 = early_result.get('phase3_pitch', {})
    with st.expander('⭐ フェーズ3：提案職種説明（最重要）', expanded=True):
        # 提案スタイル・ミライ型スコア
        col_a, col_b, col_c, col_d = st.columns(4)
        style = p3.get('提案スタイル', '─')
        col_a.metric('提案スタイル', f'{style}型' if style in ('A', 'B') else '─',
                     delta='✅職種軸型（ベテラン）' if style == 'A' else '⚠️求人一つずつ型（新人）',
                     delta_color='off')
        col_b.metric('提案方向数', f'{p3.get("提案方向数", "─")}方向')
        mts = p3.get('ミライ型スコア', 0)
        col_c.metric('ミライ型スコア ★', f'{mts}/3',
                     delta='✅ミライ型' if mts >= 2 else ('⚠️部分的' if mts == 1 else '❌求人票型'),
                     delta_color='off')
        mcs = p3.get('マッチングスコア', 0)
        col_d.metric('マッチングスコア', f'{mcs}/3' + (' ★' if p3.get('プロフェッショナル') else ''),
                     delta='✅良好' if mcs >= 2 else ('⚠️' if mcs == 1 else '❌'), delta_color='off')

        # ミライ型詳細
        st.markdown('<div class="sub-title">🔮 求人票型 vs ミライ型</div>', unsafe_allow_html=True)
        st.markdown(f"""
    | スコア | 意味 |
    |---|---|
    | 0 | 求人票の読み上げのみ（残業時間・給与・福利厚生） |
    | 1 | 条件面の説明のみ |
    | 2 | 部分的にミライ・活躍イメージを使って説明 |
    | 3 | 全体的にミライ型（将来像・業界成長・活躍イメージで説明） |
    """)
        if p3.get('ミライ型根拠'):
            st.write('📋 判定根拠: ' + p3['ミライ型根拠'])
        if p3.get('ミライ型具体例'):
            st.markdown(
                f'<div class="concern-hit">✅ ミライ型の例：<b>{sq(p3["ミライ型具体例"])}</b></div>',
                unsafe_allow_html=True)

        # マッチングスコア詳細
        st.markdown('<div class="sub-title">🎯 価値観×適正マッチング説明</div>', unsafe_allow_html=True)
        if p3.get('マッチングスコア根拠'):
            st.write('📋 ' + p3['マッチングスコア根拠'])
        if p3.get('プロフェッショナル'):
            st.markdown(
                f'<div class="good-flag">🌟 <b>プロフェッショナル判定：</b>意向と適正が一致する適職として言語化し、求職者が気づいた場面が確認されました。<br>{sq(p3.get("プロフェッショナル根拠", ""))}</div>',
                unsafe_allow_html=True)

        # その他指標
        col_e, col_f = st.columns(2)
        with col_e:
            mkt = '✅あり' if p3.get('市場トレンド説明') else '❌なし'
            st.metric('市場・業界トレンド説明', mkt)
            if p3.get('市場トレンドコメント'):
                st.caption(p3['市場トレンドコメント'])
        with col_f:
            wks = p3.get('ワクワク反応スコア', 0)
            st.metric('求職者のワクワク反応', f'{wks}/3')
            if p3.get('ワクワク反応コメント'):
                st.caption(p3['ワクワク反応コメント'])

        if p3.get('転職動機紐づけコメント'):
            st.write('🔗 転職動機への紐づけ: ' + p3['転職動機紐づけコメント'])
        if p3.get('改善アドバイス'):
            st.warning('💡 改善: ' + p3['改善アドバイス'])

    # ════════════════════════════════════════════════
    # フェーズ4：求人提案
    # ════════════════════════════════════════════════
    p4 = late_result.get('phase4_proposal', {}) if late_result else {}
    with st.expander('📌 フェーズ4：求人提案', expanded=True):
        if not p4:
            st.warning('フェーズ4の分析データが取得できませんでした。')
        else:
            c1, c2, c3 = st.columns(3)
            vps = p4.get('価値提供スコア', 0)
            c1.metric('価値提供スコア', f'{vps}/3',
                      delta='✅良好' if vps >= 2 else ('⚠️' if vps == 1 else '❌'), delta_color='off')
            dts = p4.get('難易度伝達スコア', 0)
            c2.metric('難易度伝達', f'{dts}/3',
                      delta='✅良好' if dts >= 2 else ('⚠️' if dts == 1 else '❌'), delta_color='off')
            cas = p4.get('CA一方向比率スコア', 0)
            c3.metric('会話バランス', f'{cas}/3',
                      delta='✅良好' if cas >= 2 else ('⚠️' if cas == 1 else '❌一方向'), delta_color='off')
            if p4.get('価値提供コメント'):
                st.write('📋 ' + p4['価値提供コメント'])
            if p4.get('価値提供良い例'):
                st.markdown(
                    f'<div class="concern-hit">✅ 良い例：{sq(p4["価値提供良い例"])}</div>',
                    unsafe_allow_html=True)
            if p4.get('難易度伝達コメント'):
                st.write('📊 ' + p4['難易度伝達コメント'])
            if p4.get('提案不足理由'):
                st.markdown(
                    f'<div class="warning-box">⚠️ 提案件数が少ない理由: {p4["提案不足理由"]}</div>',
                    unsafe_allow_html=True)
            if p4.get('改善アドバイス'):
                st.warning('💡 改善: ' + p4['改善アドバイス'])

    # ════════════════════════════════════════════════
    # フェーズ5：応募意思確認
    # ════════════════════════════════════════════════
    with st.expander('⭐ フェーズ5：応募意思確認・渋り対応', expanded=True):
        if not phase5:
            st.warning('フェーズ5の分析データが取得できませんでした。')
        else:
            c1, c2, c3, c4 = st.columns(4)
            nos = phase5.get('能動性スコア', 0)
            c1.metric('求職者の能動性', f'{nos}/3',
                      delta='✅能動的' if nos >= 2 else ('⚠️' if nos == 1 else '❌受動的'), delta_color='off')
            aqs = phase5.get('応諾の質スコア', 0)
            c2.metric('応諾の質', f'{aqs}/3',
                      delta='✅納得応諾' if aqs >= 2 else ('⚠️' if aqs == 1 else '❌流れで応諾'), delta_color='off')
            c3.metric('応諾後の不安', '⚠️あり（辞退リスク）' if phase5.get('応諾後不安') else '✅なし',
                      delta=None)
            c4.metric('自発的質問数', f'{phase5.get("自発的質問数", 0)}回')

            if phase5.get('能動性コメント'):
                st.write('💬 ' + phase5['能動性コメント'])
            if phase5.get('応諾の質コメント'):
                st.write('✍️ ' + phase5['応諾の質コメント'])
            if phase5.get('覚悟の醸成'):
                st.success('✅ 覚悟の醸成: ' + phase5.get('覚悟の醸成コメント', ''))
            if phase5.get('NGワード検出'):
                st.markdown(
                    f'<div class="warning-box">⚠️ <b>NGワード検出：</b>{phase5.get("NGワード内容", "")}</div>',
                    unsafe_allow_html=True)

            # 渋り対応
            st.markdown('<div class="sub-title">🤝 渋り検出 → アウト返し評価</div>',
                        unsafe_allow_html=True)
            if phase5.get('渋り検出'):
                st.markdown(
                    f'<div class="warning-box">⚠️ 渋りを検出：{sq(phase5.get("渋り場面", ""))}</div>',
                    unsafe_allow_html=True)
                col_r1, col_r2 = st.columns(2)
                r1 = phase5.get('実際に会ってみないとアウト返し', 'なし')
                r2 = phase5.get('面接練習論アウト返し', 'なし')
                eff = phase5.get('アウト返し効果', '─')
                col_r1.metric('「実際に会ってみないと」返し', r1,
                              delta='✅' if r1 == 'できた' else '❌', delta_color='off')
                col_r2.metric('「面接練習論」返し', r2,
                              delta='✅' if r2 == 'できた' else '❌', delta_color='off')
                st.metric('アウト返しの効果', eff)
                if phase5.get('アウト返しコメント'):
                    st.write('📋 ' + phase5['アウト返しコメント'])
            else:
                st.markdown('<div class="phase-skip">渋りは検出されませんでした。</div>',
                            unsafe_allow_html=True)

            if phase5.get('改善アドバイス'):
                st.warning('💡 改善: ' + phase5['改善アドバイス'])

    # ════════════════════════════════════════════════
    # フェーズ6：次回アポ
    # ════════════════════════════════════════════════
    p6 = late_result.get('phase6_closing', {}) if late_result else {}
    with st.expander('📌 フェーズ6：次回アポ・クロージング', expanded=False):
        if not p6:
            st.warning('フェーズ6の分析データが取得できませんでした。')
        else:
            c1, c2, c3 = st.columns(3)
            c1.metric('面接対策アポ', '✅あり' if p6.get('面接対策アポ') else '❌なし')
            c2.metric('日程の具体性', p6.get('日程具体性', '─'))
            c3.metric('書類・次ステップ確認', '✅あり' if p6.get('書類次ステップ確認') else '❌なし')
            if p6.get('コメント'):
                st.write('📋 ' + p6['コメント'])
            if p6.get('改善アドバイス'):
                st.info('💡 改善: ' + p6['改善アドバイス'])

            ca_close = metrics.get('クロージングCA', [])
            cd_close = metrics.get('クロージングCD', [])
            if ca_close or cd_close:
                with st.expander('クロージング発話を確認'):
                    for t in ca_close: st.markdown('🔵 **CA：** ' + t[:120])
                    for t in cd_close: st.markdown('⚪ **求職者：** ' + t[:120])

    st.divider()

    # ════════════════════════════════════════════════
    # 懸念深掘り詳細（フェーズ4〜5の補足）
    # ════════════════════════════════════════════════
    st.markdown('<div class="section-title">🔍 懸念深掘りマップ</div>', unsafe_allow_html=True)

    ca_obj      = late_result.get('concern_analysis', {}) if late_result else {}
    all_moments = metrics.get('懸念モーメント全件', [])
    cn_total    = metrics['懸念場面数']
    cn_skip     = metrics['懸念スルー数']
    cn_drill    = cn_total - cn_skip

    col_a, col_b, col_c = st.columns(3)
    col_a.metric('懸念ワード出現', f'{cn_total}場面')
    col_b.metric('深掘りできた', f'{cn_drill}場面',
                 delta='✅' if cn_drill >= cn_total * 0.5 else '⚠️', delta_color='off')
    col_c.metric('スルーした', f'{cn_skip}場面',
                 delta='❌ 要改善' if cn_skip > 0 else '✅', delta_color='off')

    if ca_obj.get('summary'):
        st.info('📊 ' + ca_obj['summary'])

    missed_scenes = ca_obj.get('missed_scenes', [])
    if missed_scenes:
        st.markdown('<div class="sub-title">❌ スルーした場面 → こう返すべきだった</div>',
                    unsafe_allow_html=True)
        for ms in missed_scenes:
            col_l, col_r = st.columns([1, 1])
            with col_l:
                st.markdown(
                    f'<div class="concern-miss">'
                    f'<small>求職者</small><br><b>{sq(ms.get("cd_text", ""))}</b><br>'
                    f'<small style="color:#888">↓ 実際のCA</small><br>'
                    f'<span style="color:#c00000">{sq(ms.get("ca_actual", ""))}</span>'
                    f'</div>', unsafe_allow_html=True)
            with col_r:
                st.markdown(
                    f'<div class="concern-hit">'
                    f'<small>💡 こう返すと深掘りできた</small><br>'
                    f'<b style="color:#375623">{sq(ms.get("ca_suggested", ""))}</b><br>'
                    f'<small style="color:#555">→ {ms.get("why", "")}</small>'
                    f'</div>', unsafe_allow_html=True)

    good_scenes = ca_obj.get('good_scenes', [])
    if good_scenes:
        st.markdown('<div class="sub-title">✅ うまく対応できた場面</div>', unsafe_allow_html=True)
        for gs in good_scenes:
            st.markdown(
                f'<div class="concern-hit">'
                f'<small>求職者</small>　<b>{sq(gs.get("cd_text", ""))}</b><br>'
                f'<small style="color:#888">↓ CA（深掘り）</small><br>'
                f'<span style="color:#375623"><b>{sq(gs.get("ca_text", ""))}</b></span><br>'
                f'<small>💡 {gs.get("why_good", "")}</small>'
                f'</div>', unsafe_allow_html=True)

    st.divider()

    # ════════════════════════════════════════════════
    # 次の面談で使えるフレーズ集
    # ════════════════════════════════════════════════
    st.markdown('<div class="section-title">🗣️ 次の面談で使えるフレーズ集</div>', unsafe_allow_html=True)
    next_phrases = late_result.get('next_phrases', []) if late_result else []
    if next_phrases:
        for ph in next_phrases:
            if not ph.get('phrase'): continue
            st.markdown(
                f'<div style="margin:8px 0">'
                f'<small style="color:#888; background:#f0f0f0; padding:2px 8px; border-radius:10px">'
                f'{ph.get("situation", "")}</small><br>'
                f'<div class="phrase-box">{sq(ph.get("phrase", ""))}</div>'
                f'<small style="color:#555">💡 {ph.get("why", "")}</small>'
                f'</div>', unsafe_allow_html=True)

    st.divider()

    # ── Word出力 / Slack送信 / メモ・保存先 ──────────
    st.markdown('<div class="section-title">⚡ アクション</div>', unsafe_allow_html=True)

    _d_action = {
        'ca':       _res.get('ca', ''),
        'candidate':_res.get('candidate', ''),
        'grip':     _res.get('grip', ''),
        '提案件数': _res.get('proposed'),
        '応諾件数': _res.get('accepted'),
        **early_result,
        **(late_result or {}),
        'behaviors': _res.get('behaviors', {}),
    }

    act1, act2 = st.columns(2)

    with act1:
        st.markdown('**📄 Word文書として出力**')
        if st.button('Word文書を生成', key='fb_gen_word', use_container_width=True):
            with st.spinner('生成中...'):
                try:
                    docx_bytes = _generate_word_doc(_d_action)
                    fname = f'求人提案分析_{_d_action["ca"]}_{_d_action["candidate"]}.docx'
                    st.download_button(
                        '📥 ダウンロード', data=docx_bytes, file_name=fname,
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        use_container_width=True, key='fb_dl_word')
                except Exception as e:
                    st.error(f'生成失敗：{e}')

    with act2:
        st.markdown('**💬 Slackに送信**')
        if st.button('Slackに送信', key='fb_send_slack', use_container_width=True):
            if not SLACK_WEBHOOK:
                st.error('Slack Webhook URLが設定されていません')
            else:
                with st.spinner('送信中...'):
                    ok_flag, msg = _send_slack(_d_action, SLACK_WEBHOOK)
                    if ok_flag: st.success('✅ 送信しました')
                    else:       st.error(f'送信失敗：{msg}')

    st.divider()
    if early_result.get('notes'):
        with st.expander('📝 分析メモ'):
            st.write(early_result['notes'])
    with st.expander('💾 保存先の確認'):
        st.code(_res.get('ref_file', ''), language=None)

    st.divider()

    # ════════════════════════════════════════════════
if '_proposal_res' not in st.session_state and '_proposal_retry' not in st.session_state and not file_bytes:
    st.info('👆 ファイルをアップロードして分析を開始してください')
st.markdown("""
**対応ファイル形式**
- `.docx`：話者ラベルあり（CA名 / 求職者名 形式）→ 即解析
- `.txt`：音声文字起こし → Claude Haikuで話者推定

**ファイル命名規則**（CA名・求職者名・提案件数が自動入力されます）
```
CA名_グリップランク_求職者名_求人提案_提案N_応諾N.docx
例: 下川_グリップA_岡寺さん_求人提案_提案11_応諾11.docx
```

**このツールが分析すること**
- 📌 フェーズ1：冒頭・前回振り返り（温度感確認・メリデメ説明）
- 📌 フェーズ2：転職基礎説明（初期フェーズ求職者のみ）
- ⭐ フェーズ3：提案職種説明（ミライ型スコア・提案スタイルA/B・マッチング説明★）
- 📌 フェーズ4：求人提案（価値提供型 vs 読み上げ・難易度伝達）
- ⭐ フェーズ5：応募意思確認（渋り検出・アウト返し評価・能動性スコア）
- 📌 フェーズ6：次回アポ（日程具体性・書類確認）
- 🔍 懸念深掘りマップ（スルー場面ごとに改善例）
- 🗣️ 次の面談で使えるフレーズ集（この求職者のキーワードを使用）
""")
