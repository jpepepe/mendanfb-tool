# -*- coding: utf-8 -*-
"""
求人提案 分析ダッシュボード
"""

import streamlit as st
import json, sys, os, io
from pathlib import Path
from typing import Optional
import pandas as pd

sys.path.insert(0, str(Path(__file__).parent.parent))
from gdrive import (list_json_files, download_json as gdrive_download_json,
                    upload_json as gdrive_upload_json,
                    upload_proposal_summary, download_proposal_summary)
import anthropic
import requests



st.markdown("""
<style>
.section-title { background:#1F3864; color:white; padding:8px 16px; border-radius:6px;
                 font-weight:bold; margin:16px 0 8px 0; }
.sub-title { background:#2E75B6; color:white; padding:6px 14px; border-radius:4px;
             font-weight:bold; margin:10px 0 6px 0; font-size:0.9rem; }
.score-high { border-left:4px solid #375623; background:#e2efda; padding:8px 14px;
              border-radius:4px; margin:4px 0; }
.score-mid  { border-left:4px solid #c55a11; background:#fce4d6; padding:8px 14px;
              border-radius:4px; margin:4px 0; }
.score-low  { border-left:4px solid #c00000; background:#fcecea; padding:8px 14px;
              border-radius:4px; margin:4px 0; }
.concern-hit  { background:#e2efda; border-left:4px solid #375623; padding:8px 12px;
                border-radius:4px; margin:6px 0; }
.concern-miss { background:#fcecea; border-left:4px solid #c00000; padding:8px 12px;
                border-radius:4px; margin:6px 0; }
.phrase-box { background:#1F3864; color:white; padding:10px 14px; border-radius:6px;
              margin:6px 0; font-family:monospace; }
.grade-S { background:#1a5276; }
.grade-A { background:#1e8449; }
.grade-B { background:#2471a3; }
.grade-C { background:#d35400; }
.grade-D { background:#c0392b; }
</style>
""", unsafe_allow_html=True)

OUTPUT_JSON_PROPOSAL = Path(__file__).parent.parent / "output" / "json_proposal"

def sq(text: str) -> str:
    t = (text or '').strip()
    if t.startswith('「') and t.endswith('」'):
        return t
    return f'「{t}」'

SLACK_WEBHOOK = st.secrets.get("SLACK_WEBHOOK_URL", "") if hasattr(st, 'secrets') else ""
OUTPUT_UTT_PROPOSAL = Path(__file__).parent.parent / "output" / "utterances_proposal"

# ── Word文書生成（求人提案版） ────────────────────────────
def generate_word_doc_proposal(d: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    def h(text, level=1, color=(31,56,100)):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)

    def body(text, bold=False, color=None, size=10.5):
        p = doc.add_paragraph()
        r = p.add_run(str(text))
        r.bold = bold; r.font.size = Pt(size)
        if color: r.font.color.rgb = RGBColor(*color)

    def kv(key, val, kc=(46,117,182)):
        p = doc.add_paragraph()
        r1 = p.add_run(f'{key}：'); r1.bold = True
        r1.font.color.rgb = RGBColor(*kc); r1.font.size = Pt(10.5)
        r2 = p.add_run(str(val or '')); r2.font.size = Pt(10.5)

    def tbl(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
        for i, h_ in enumerate(headers):
            c = t.rows[0].cells[i]; c.text = h_
            if c.paragraphs[0].runs: c.paragraphs[0].runs[0].bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row): cells[i].text = str(v or '')
        doc.add_paragraph()

    gd      = d.get('grip_drivers', {})
    bh      = d.get('behaviors', {})
    ov      = d.get('overall', {})
    p1      = d.get('phase1_opening', {})
    p2      = d.get('phase2_basics', {})
    p3      = d.get('phase3_pitch', {})
    p4      = d.get('phase4_proposal', {})
    p5      = d.get('phase5_agreement', {})
    p6      = d.get('phase6_closing', {})
    ca_an   = d.get('concern_analysis', {})
    phrases = d.get('next_phrases', [])
    axes    = ['求人マッチング説明','動機付け','懸念解消','認識統一','次回接続']
    total   = sum(gd.get(ax,{}).get('score',0) for ax in axes)
    grade   = ov.get('grade','') or ('S' if total>=13 else 'A' if total>=10 else 'B' if total>=7 else 'C' if total>=4 else 'D')
    proposed = d.get('提案件数') or 0
    accepted = d.get('応諾件数') or 0

    # ══ タイトル ══
    t0 = doc.add_heading('求人提案 分析レポート', 0)
    t0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kv('CA名', d.get('ca','')); kv('候補者名', d.get('candidate',''))
    kv('グリップ', d.get('grip','')); kv('提案件数', f'{proposed}件 / 応諾{accepted}件')
    doc.add_paragraph()

    # ══ 1. 総合評価 ══
    h('■ 1. 総合評価')
    p = doc.add_paragraph()
    r = p.add_run(f'グレード：{grade}　　総合スコア：{total}/15')
    r.bold = True; r.font.size = Pt(14)
    if proposed > 0:
        rate = accepted / proposed * 100
        body(f'応諾率：{rate:.0f}%（{accepted}/{proposed}件）', bold=True,
             color=(55,86,35) if rate >= 50 else (192,0,0))
    if ov.get('grade_reason'): body(ov['grade_reason'])
    if ov.get('top_strength'): body(f'【最大の強み】{ov["top_strength"]}', color=(55,86,35))
    doc.add_paragraph()

    # ══ 2. ルーブリック採点（5軸） ══
    h('■ 2. ルーブリック採点（5軸）')
    tbl(
        ['評価軸','スコア','強み','改善点','次のアクション'],
        [[ax, f'{gd.get(ax,{}).get("score",0)}/3',
          gd.get(ax,{}).get('strength',''),
          gd.get(ax,{}).get('weakness',''),
          gd.get(ax,{}).get('next_action','')] for ax in axes]
    )

    # ══ 3. 優先改善ポイント ══
    if ov.get('top_issues'):
        h('■ 3. 優先改善ポイント TOP3')
        for i, issue in enumerate(ov['top_issues'][:3], 1):
            if isinstance(issue, dict):
                body(f'{i}. {issue.get("issue","")}', bold=True)
                if issue.get('detail'): body(f'　📌 {issue["detail"]}')
                if issue.get('fix'):    body(f'　💡 改善例：{issue["fix"]}', color=(55,86,35))
        doc.add_paragraph()

    if ov.get('one_thing'):
        h('■ 4. 次の面談で必ず試してほしいこと')
        body(ov['one_thing'], color=(46,117,182), bold=True)
        doc.add_paragraph()

    # ══ 5. フェーズ別評価 ══
    h('■ 5. フェーズ別評価')

    h('フェーズ1：冒頭・前回振り返り', level=2)
    body(f'振り返り実施：{"あり" if p1.get("振り返り実施") else "なし"}　'
         f'温度感確認：{"あり" if p1.get("温度感確認") else "なし"}　'
         f'メリデメ説明：{p1.get("メリデメ説明スコア",0)}/3')
    if p1.get('振り返り評価'): body(p1['振り返り評価'])
    if p1.get('改善アドバイス'): body(f'💡 {p1["改善アドバイス"]}', color=(55,86,35))

    h('フェーズ2：転職基礎説明', level=2)
    if not p2.get('対象'):
        body('対象外（転職活動経験あり）')
    else:
        body(f'実施：{"あり" if p2.get("実施") else "なし"}　'
             f'フロー説明：{"✅" if p2.get("転職フロー説明") else "❌"}　'
             f'選考基準数字：{"✅" if p2.get("選考基準数字説明") else "❌"}　'
             f'必要応募数：{"✅" if p2.get("必要応募数説明") else "❌"}')
        if p2.get('コメント'): body(p2['コメント'])

    h('フェーズ3：提案職種説明（★最重要）', level=2, color=(192,0,0))
    body(f'提案スタイル：{p3.get("提案スタイル","─")}型（A=職種軸/B=一社ずつ）　'
         f'提案方向数：{p3.get("提案方向数","─")}')
    body(f'ミライ型スコア：{p3.get("ミライ型スコア",0)}/3　'
         f'マッチングスコア：{p3.get("マッチングスコア",0)}/3'
         + ('　★プロフェッショナル判定' if p3.get('プロフェッショナル') else ''))
    if p3.get('ミライ型根拠'): body(f'ミライ型根拠：{p3["ミライ型根拠"]}')
    if p3.get('ミライ型具体例'): body(f'ミライ型の例：{p3["ミライ型具体例"]}', color=(55,86,35))
    if p3.get('マッチングスコア根拠'): body(p3['マッチングスコア根拠'])
    if p3.get('改善アドバイス'): body(f'💡 {p3["改善アドバイス"]}', color=(55,86,35))

    h('フェーズ4：求人提案', level=2)
    if p4:
        body(f'価値提供スコア：{p4.get("価値提供スコア",0)}/3　'
             f'難易度伝達：{p4.get("難易度伝達スコア",0)}/3　'
             f'会話バランス：{p4.get("CA一方向比率スコア",0)}/3')
        if p4.get('価値提供コメント'): body(p4['価値提供コメント'])
        if p4.get('改善アドバイス'): body(f'💡 {p4["改善アドバイス"]}', color=(55,86,35))

    h('フェーズ5：応募意思確認（★重要）', level=2, color=(192,0,0))
    if p5:
        body(f'能動性スコア：{p5.get("能動性スコア",0)}/3　'
             f'応諾の質：{p5.get("応諾の質スコア",0)}/3　'
             f'応諾後不安：{"あり⚠️" if p5.get("応諾後不安") else "なし"}')
        if p5.get('NGワード検出'):
            body(f'⚠️ NGワード検出：{p5.get("NGワード内容","")}', color=(192,0,0))
        if p5.get('渋り検出'):
            body(f'渋り検出：{p5.get("渋り場面","")}')
            body(f'「実際に会ってみないと」返し：{p5.get("実際に会ってみないとアウト返し","なし")}　'
                 f'「面接練習論」返し：{p5.get("面接練習論アウト返し","なし")}　'
                 f'効果：{p5.get("アウト返し効果","─")}')
        if p5.get('改善アドバイス'): body(f'💡 {p5["改善アドバイス"]}', color=(55,86,35))

    h('フェーズ6：次回アポ', level=2)
    if p6:
        body(f'面接対策アポ：{"あり" if p6.get("面接対策アポ") else "なし"}　'
             f'日程具体性：{p6.get("日程具体性","─")}　'
             f'書類確認：{"あり" if p6.get("書類次ステップ確認") else "なし"}')
        if p6.get('コメント'): body(p6['コメント'])
        if p6.get('改善アドバイス'): body(f'💡 {p6["改善アドバイス"]}', color=(55,86,35))
    doc.add_paragraph()

    # ══ 6. 懸念深掘り分析 ══
    h('■ 6. 懸念深掘り分析')
    if ca_an.get('summary'): body(ca_an['summary'])
    for ms in (ca_an.get('missed_scenes') or []):
        if not isinstance(ms, dict): continue
        body(f'求職者：{ms.get("cd_text","")}')
        body(f'実際のCA：{ms.get("ca_actual","")}', color=(192,0,0))
        body(f'→ こう返すべきだった：{ms.get("ca_suggested","")}', color=(55,86,35))
        if ms.get('why'): body(f'　理由：{ms["why"]}')
        doc.add_paragraph()
    for gs in (ca_an.get('good_scenes') or []):
        if not isinstance(gs, dict): continue
        body(f'✅ 求職者：{gs.get("cd_text","")}')
        body(f'CA（深掘り）：{gs.get("ca_text","")}', color=(55,86,35))
        if gs.get('why_good'): body(f'　良かった点：{gs["why_good"]}')
        doc.add_paragraph()

    # ══ 7. 行動指標 ══
    h('■ 7. 行動指標')
    ok = lambda v, thr, rev=False: '✅' if (v <= thr if rev else v >= thr) else '❌'
    tbl(
        ['指標','値','目標','判定'],
        [
            ('求職者発話比率',  f'{round(bh.get("求職者発話比率",0)*100)}%', '35%以上', ok(bh.get('求職者発話比率',0)*100, 35)),
            ('懸念スルー率',    f'{bh.get("懸念スルー率",0)}%',              '50%以下', ok(bh.get('懸念スルー率',0), 50, rev=True)),
            ('バックトラッキング',f'{bh.get("バックトラッキング",0)}回',      '5回以上', ok(bh.get('バックトラッキング',0), 5)),
            ('マッチング説明',  f'{bh.get("マッチング説明数",0)}回',          '2回以上', ok(bh.get('マッチング説明数',0), 2)),
            ('名前呼称回数',    f'{bh.get("名前呼称回数",0)}回',              '3回以上', ok(bh.get('名前呼称回数',0), 3)),
            ('フィラー回数',    f'{bh.get("フィラー回数",0)}回',              '30回以下',ok(bh.get('フィラー回数',0), 30, rev=True)),
            ('ポジティブ反応',  f'{bh.get("ポジティブ反応",0)}回',            '5回以上', ok(bh.get('ポジティブ反応',0), 5)),
        ]
    )

    # ══ 8. フレーズ集 ══
    if phrases:
        h('■ 8. 次の面談で使えるフレーズ集')
        for ph in phrases:
            if not isinstance(ph, dict) or not ph.get('phrase'): continue
            body(f'【{ph.get("situation","")}】', bold=True)
            p_ph = doc.add_paragraph()
            r_ph = p_ph.add_run(ph.get('phrase',''))
            r_ph.font.color.rgb = RGBColor(31,56,100)
            r_ph.bold = True; r_ph.font.size = Pt(11)
            if ph.get('why'): body(f'　💡 {ph["why"]}')
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Slack送信（求人提案版） ────────────────────────────────
def send_to_slack_proposal(webhook_url: str, d: dict) -> tuple:
    gd      = d.get('grip_drivers', {})
    ov      = d.get('overall', {})
    p3      = d.get('phase3_pitch', {})
    p5      = d.get('phase5_agreement', {})
    bh      = d.get('behaviors', {})
    ca_an   = d.get('concern_analysis', {})
    phrases = d.get('next_phrases', [])

    axes  = ['求人マッチング説明','動機付け','懸念解消','認識統一','次回接続']
    total = sum(gd.get(ax,{}).get('score',0) for ax in axes)
    grade = ov.get('grade','') or ('S' if total>=13 else 'A' if total>=10 else 'B' if total>=7 else 'C' if total>=4 else 'D')
    g_emoji = {'S':'🏆','A':'🟢','B':'🔵','C':'🟡','D':'🔴'}.get(grade,'⚪')

    proposed = d.get('提案件数') or 0
    accepted = d.get('応諾件数') or 0
    rate_str = f'{accepted/proposed*100:.0f}%（{accepted}/{proposed}件）' if proposed else '─'

    # ルーブリック
    score_lines = '\n'.join(
        f"　{ax}：{'█'*gd.get(ax,{}).get('score',0)}{'░'*(3-gd.get(ax,{}).get('score',0))} "
        f"{gd.get(ax,{}).get('score',0)}/3　{gd.get(ax,{}).get('weakness','')}"
        for ax in axes)

    # 改善ポイント
    issues_text = ''
    for i, issue in enumerate(ov.get('top_issues',[])[:3], 1):
        if isinstance(issue, dict):
            issues_text += f'{i}. *{issue.get("issue","")}*\n　{issue.get("detail","")}\n　💡 {issue.get("fix","")}\n'

    # フェーズ3サマリー
    style  = p3.get('提案スタイル','─')
    mirai  = p3.get('ミライ型スコア','─')
    match  = p3.get('マッチングスコア','─')
    p3_text = f'スタイル：{style}型　ミライ型：{mirai}/3　マッチング：{match}/3'
    if p3.get('プロフェッショナル'): p3_text += '　★プロフェッショナル'

    # フェーズ5サマリー
    p5_text = f'能動性：{p5.get("能動性スコア","─")}/3　応諾の質：{p5.get("応諾の質スコア","─")}/3'
    if p5.get('渋り検出'):
        p5_text += f'\n　渋り検出あり → アウト返し効果：{p5.get("アウト返し効果","─")}'

    # 懸念深掘り
    concern_text = ca_an.get('summary','')
    for ms in (ca_an.get('missed_scenes') or [])[:2]:
        if not isinstance(ms, dict): continue
        concern_text += f'\n　スルー：「{ms.get("cd_text","")}」\n　→ `{ms.get("ca_suggested","")}`'

    # フレーズ集
    phrases_text = ''
    for ph in phrases:
        if not isinstance(ph, dict) or not ph.get('phrase'): continue
        phrases_text += f'*{ph.get("situation","")}*\n　`{ph.get("phrase","")}` — {ph.get("why","")}\n'

    blocks = [
        {"type": "header", "text": {"type": "plain_text",
            "text": f'{g_emoji} 求人提案FB｜{d.get("ca","")}CA × {d.get("candidate","")}　グリップ{d.get("grip","")}'}},
        {"type": "section", "fields": [
            {"type": "mrkdwn", "text": f"*グレード:* {grade}　総合スコア：{total}/15"},
            {"type": "mrkdwn", "text": f"*応諾率:* {rate_str}"},
        ]},
        {"type": "section", "text": {"type": "mrkdwn",
            "text": f"*📋 グレード判定理由*\n{ov.get('grade_reason','─')}"}},
        {"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🎯 ルーブリック（5軸）*\n{score_lines}"}},
        {"type": "section", "text": {"type": "mrkdwn",
            "text": f"*⭐ フェーズ3（提案職種説明）*\n{p3_text}"}},
        {"type": "section", "text": {"type": "mrkdwn",
            "text": f"*⭐ フェーズ5（応募意思確認）*\n{p5_text}"}},
        {"type": "section", "text": {"type": "mrkdwn",
            "text": f"*💪 強み*\n{ov.get('top_strength','─')}"}},
        {"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🎯 優先改善ポイント*\n{issues_text or '─'}"}},
    ]
    if concern_text:
        blocks.append({"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🔍 懸念深掘り*\n{concern_text}"}})
    if ov.get('one_thing'):
        blocks.append({"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🚀 次に試すこと*\n{ov['one_thing']}"}})
    if phrases_text:
        blocks.append({"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🗣️ 次の面談で使えるフレーズ集*\n{phrases_text.strip()}"}})

    try:
        resp = requests.post(webhook_url, json={"blocks": blocks}, timeout=10)
        if resp.status_code == 200:
            return True, '送信成功'
        return False, f'エラー: {resp.status_code} {resp.text}'
    except Exception as e:
        return False, str(e)


# ── 再分析用API呼び出し（ツールモジュールに依存しない独立実装） ──
import re as _re

def _phase_early(utterances, ca_name, cand_name, fmt, proposed, accepted, client):
    transcript = '\n'.join(f"[{u['speaker']}] {u['text']}" for u in utterances)[:15000]
    prop_str = f"{proposed}件提案 / {accepted}件応諾" if proposed else "提案件数不明"
    prompt = f"""あなたは人材紹介会社のトレーナーです。以下は「求人提案面談」の文字起こしです。
フェーズ1〜3の評価と、5軸スコア・総合評価をJSONで返してください。

## スコア基準（0〜3点）
- 3: 根拠引用あり + 求職者の明示的な同意・前向きな反応が確認できる
- 2: 把握できているが確認が弱い or 求職者が「はい」止まり
- 1: 触れているが浅い・一方的な説明のみ
- 0: 未実施

## メタ情報
CA名: {ca_name} / 求職者名: {cand_name} / 形式: {fmt}
提案結果: {prop_str}

## 面談文字起こし
{transcript}

## 出力（JSONのみ・余計な文章不要）
{{
  "phase1_opening": {{"振り返り実施":true,"振り返り評価":"","温度感確認":true,"温度感評価":"","メリデメ説明スコア":0,"メリデメ説明コメント":"","改善アドバイス":""}},
  "phase2_basics": {{"対象":true,"実施":false,"転職フロー説明":false,"選考基準数字説明":false,"必要応募数説明":false,"タイミング評価":"なし","コメント":""}},
  "phase3_pitch": {{"提案スタイル":"B","提案スタイル根拠":"","提案方向数":1,"ミライ型スコア":0,"ミライ型根拠":"","ミライ型具体例":"","市場トレンド説明":false,"市場トレンドコメント":"","マッチングスコア":0,"マッチングスコア根拠":"","プロフェッショナル":false,"プロフェッショナル根拠":"","転職動機紐づけスコア":0,"転職動機紐づけコメント":"","ワクワク反応スコア":0,"ワクワク反応コメント":"","改善アドバイス":""}},
  "grip_drivers": {{
    "求人マッチング説明": {{"score":0,"evidence":[],"strength":"","weakness":"","next_action":""}},
    "動機付け":          {{"score":0,"evidence":[],"strength":"","weakness":"","next_action":""}},
    "懸念解消":          {{"score":0,"evidence":[],"strength":"","weakness":"","next_action":""}},
    "認識統一":          {{"score":0,"evidence":[],"strength":"","weakness":"","next_action":""}},
    "次回接続":          {{"score":0,"evidence":[],"strength":"","weakness":"","next_action":""}}
  }},
  "overall": {{"grade":"B","grade_reason":"","top_strength":"","top_issues":[{{"issue":"","detail":"","fix":""}}],"one_thing":""}},
  "notes": ""
}}

grade: S=全軸2.5以上+ミライ型+懸念深掘り◎, A=総合スコア10以上, B=7〜9, C=4〜6, D=3以下
提案スタイルA: 職種軸で整理して提案 / B: 求人を一社ずつ順番に説明"""
    try:
        resp = client.messages.create(
            model='claude-sonnet-4-6', max_tokens=6000,
            messages=[{'role': 'user', 'content': prompt}])
        content = _re.sub(r'```(?:json)?\s*', '', resp.content[0].text.strip()).strip()
        if resp.stop_reason == 'max_tokens':
            content = content.rstrip(',\n ')
            for _ in range(10): content += '}'
            content += ']}}'
        try:
            return json.loads(content)
        except:
            m = _re.search(r'\{.*\}', content, _re.DOTALL)
            if m:
                try: return json.loads(m.group())
                except: pass
    except Exception:
        pass
    return {}

def _phase_late(utterances, ca_name, cand_name, client):
    transcript = '\n'.join(f"[{u['speaker']}] {u['text']}" for u in utterances)[:14000]
    prompt = f"""あなたは人材紹介会社の求人提案面談のコーチです。以下の文字起こしを分析し、
フェーズ4〜6の評価とフレーズ集をJSONで返してください。

## メタ情報
CA名: {ca_name} / 求職者名: {cand_name}

## 面談文字起こし
{transcript}

## 出力（JSONのみ・余計な文章不要）
{{
  "phase4_proposal": {{"価値提供スコア":0,"価値提供コメント":"","価値提供良い例":"","難易度伝達スコア":0,"難易度伝達コメント":"","CA一方向比率スコア":0,"CA一方向比率コメント":"","提案不足理由":"","改善アドバイス":""}},
  "phase5_agreement": {{"応募確認スタイル":"一社ずつ","能動性スコア":0,"能動性コメント":"","応諾の質スコア":0,"応諾の質コメント":"","応諾後不安":false,"応諾後不安コメント":"","覚悟の醸成":false,"覚悟の醸成コメント":"","NGワード検出":false,"NGワード内容":"","自発的質問数":0,"渋り検出":false,"渋り場面":"","実際に会ってみないとアウト返し":"なし","面接練習論アウト返し":"なし","アウト返し効果":"スルー","アウト返しコメント":"","改善アドバイス":""}},
  "phase6_closing": {{"面接対策アポ":false,"日程具体性":"なし","書類次ステップ確認":false,"コメント":"","改善アドバイス":""}},
  "concern_analysis": {{"summary":"","missed_scenes":[],"good_scenes":[]}},
  "next_phrases": [
    {{"situation":"懸念ワードが出た直後","phrase":"","why":""}},
    {{"situation":"求人と求職者を紐づける場面","phrase":"","why":""}},
    {{"situation":"渋りが出たとき","phrase":"","why":""}},
    {{"situation":"応募意思を確認するタイミング","phrase":"","why":""}},
    {{"situation":"クロージングで次のステップを確認","phrase":"","why":""}}
  ]
}}

価値提供スコア: 0=読み上げのみ, 1=条件説明のみ, 2=部分的にミライ, 3=全体的にミライ型
能動性スコア: 0=はいのみ, 1=たまに質問, 2=複数の自発的発言, 3=積極的に意見・懸念
渋り検出: 「ちょっと難しいかな」「どうかな」「考えたい」「実際会ってみないと」等
missed_scenesは改善インパクトが最も大きい代表的な場面を上限5件、good_scenesは最も参考になる場面を上限5件に厳選すること。"""
    try:
        resp = client.messages.create(
            model='claude-sonnet-4-6', max_tokens=8000,
            messages=[{'role': 'user', 'content': prompt}])
        content = _re.sub(r'```(?:json)?\s*', '', resp.content[0].text.strip()).strip()
        if resp.stop_reason == 'max_tokens':
            content = content.rstrip(',\n ')
            for _ in range(10): content += '}'
            content += ']}}'
        try:
            return json.loads(content)
        except:
            m = _re.search(r'\{.*\}', content, _re.DOTALL)
            if m:
                try: return json.loads(m.group())
                except: pass
    except Exception:
        pass
    return {}

def _save_proposal(ca, grip, candidate, fmt, proposed, accepted,
                   utterances, behaviors, early_result, late_result):
    safe_grip = grip if grip != '未入力' else 'X'
    key = f"{ca}_{safe_grip}_{candidate}_求人提案_提案{proposed or 'X'}_応諾{accepted or 'X'}"
    p3 = early_result.get('phase3_pitch', {})
    p5 = (late_result or {}).get('phase5_agreement', {})
    p6 = (late_result or {}).get('phase6_closing', {})
    dashboard = {
        "提案スタイル": p3.get('提案スタイル', ''),
        "ミライ型スコア": p3.get('ミライ型スコア', None),
        "マッチングスコア": p3.get('マッチングスコア', None),
        "能動性スコア": p5.get('能動性スコア', None),
        "渋り検出": p5.get('渋り検出', None),
        "次回アポ具体性": p6.get('日程具体性', ''),
        "5軸合計": sum(early_result.get('grip_drivers', {}).get(k, {}).get('score', 0)
                      for k in ['求人マッチング説明','動機付け','懸念解消','認識統一','次回接続']),
        "グレード": early_result.get('overall', {}).get('grade', ''),
    }
    json_data = {
        "ca": ca, "grip": safe_grip, "candidate": candidate,
        "meeting_type": "求人提案", "format": fmt,
        "提案件数": proposed, "応諾件数": accepted,
        "phase1_opening":  early_result.get('phase1_opening', {}),
        "phase2_basics":   early_result.get('phase2_basics', {}),
        "phase3_pitch":    early_result.get('phase3_pitch', {}),
        "phase4_proposal": (late_result or {}).get('phase4_proposal', {}),
        "phase5_agreement":(late_result or {}).get('phase5_agreement', {}),
        "phase6_closing":  (late_result or {}).get('phase6_closing', {}),
        "grip_drivers":    early_result.get('grip_drivers', {}),
        "overall":         early_result.get('overall', {}),
        "concern_analysis":(late_result or {}).get('concern_analysis', {}),
        "next_phrases":    (late_result or {}).get('next_phrases', []),
        "behaviors":       behaviors,
        "dashboard":       dashboard,
        "notes":           early_result.get('notes', ''),
    }
    json_path = OUTPUT_JSON_PROPOSAL / f"{key}.json"
    json_path.write_text(json.dumps(json_data, ensure_ascii=False, indent=2), encoding='utf-8')
    try:
        gdrive_upload_json(f"{key}.json", json_data, subfolder="json_proposal")
    except Exception:
        pass
    return json_path

# ── 再分析（求人提案版） ─────────────────────────────────
def reanalyze_proposal(d: dict, client, rerun_phase: str = 'all'):
    """
    rerun_phase: 'all'=最初から / 'late'=フェーズ4〜6のみ
    """
    ca        = d.get('ca', '')
    grip      = d.get('grip', 'X')
    candidate = d.get('candidate', '')
    safe_grip = grip if grip != '未入力' else 'X'
    proposed  = d.get('提案件数')
    accepted  = d.get('応諾件数')

    # 発話分離データを探す（ローカル → Drive）
    utt_data = None
    for f in OUTPUT_UTT_PROPOSAL.glob(f"{ca}_{safe_grip}_{candidate}_求人提案*.json"):
        try:
            utt_data = json.loads(f.read_text(encoding='utf-8'))
            break
        except Exception:
            continue

    if not utt_data:
        try:
            drive_utts = list_json_files(subfolder="utterances_proposal")
            matched = [f for f in drive_utts
                       if f["name"].startswith(f"{ca}_{safe_grip}_{candidate}_求人提案")]
            if matched:
                utt_data = gdrive_download_json(matched[0]["id"])
        except Exception:
            pass

    if not utt_data:
        return None, None, '文字起こしデータが見つかりません'

    utterances = utt_data.get('utterances', [])
    fmt        = utt_data.get('format', 'docx')

    early_result = None
    late_result  = None

    if rerun_phase == 'all':
        early_result = _phase_early(utterances, ca, candidate, fmt, proposed, accepted, client)
        if not early_result.get('overall'):
            return None, None, 'フェーズ1〜3の分析に失敗しました'
        late_result = _phase_late(utterances, ca, candidate, client) or {}
    else:  # 'late'
        early_result = {
            'phase1_opening': d.get('phase1_opening', {}),
            'phase2_basics':  d.get('phase2_basics', {}),
            'phase3_pitch':   d.get('phase3_pitch', {}),
            'grip_drivers':   d.get('grip_drivers', {}),
            'overall':        d.get('overall', {}),
            'notes':          d.get('notes', ''),
        }
        late_result = _phase_late(utterances, ca, candidate, client) or {}
        if not late_result.get('phase4_proposal'):
            return None, None, 'フェーズ4〜6の分析に失敗しました'

    behaviors = {k: v for k, v in d.get('behaviors', {}).items()}
    _save_proposal(ca, grip, candidate, fmt, proposed, accepted,
                   utterances, behaviors, early_result, late_result)
    return early_result, late_result, None


# ── データ読み込み ────────────────────────────────────────
@st.cache_data(ttl=600)
def load_proposal_records():
    records = []

    # ① サマリーJSONを試す（1ファイル取得で完結・高速）
    try:
        summary = download_proposal_summary()
        if summary:
            return pd.DataFrame(summary)
    except Exception:
        pass

    # ② サマリーなし → 全件DLしてサマリーを自動生成
    try:
        drive_files = list_json_files(subfolder="json_proposal")
        drive_files = [f for f in drive_files if not f["name"].startswith("_")]
        for f in sorted(drive_files, key=lambda x: x["name"]):
            try:
                d = gdrive_download_json(f["id"])
                _append(records, d, f["name"], f["id"])
            except Exception:
                continue
        if records:
            try:
                upload_proposal_summary(records)
            except Exception:
                pass
            return pd.DataFrame(records)
    except Exception:
        pass

    # ③ ローカルフォールバック
    if OUTPUT_JSON_PROPOSAL.exists():
        for f in sorted(OUTPUT_JSON_PROPOSAL.glob("*.json")):
            try:
                d = json.loads(f.read_text(encoding='utf-8'))
                _append(records, d, f.name, str(f))
            except Exception:
                continue

    return pd.DataFrame(records)

def rebuild_proposal_summary():
    """全件DLしてサマリーを再構築。「再読み込み」ボタンから呼ぶ。"""
    records = []
    try:
        drive_files = list_json_files(subfolder="json_proposal")
        drive_files = [f for f in drive_files if not f["name"].startswith("_")]
        for f in sorted(drive_files, key=lambda x: x["name"]):
            try:
                d = gdrive_download_json(f["id"])
                _append(records, d, f["name"], f["id"])
            except Exception:
                continue
        if records:
            upload_proposal_summary(records)
    except Exception:
        pass
    return records

def _append(records, d, filename, path_or_id):
    try:
        if d.get('meeting_type') != '求人提案':
            return
        gd  = d.get('grip_drivers', {})
        bh  = d.get('behaviors', {})
        ov  = d.get('overall', {})
        p3  = d.get('phase3_pitch', {})
        p5  = d.get('phase5_agreement', {})
        p1  = d.get('phase1_opening', {})
        p2  = d.get('phase2_basics', {})
        p4  = d.get('phase4_proposal', {})
        p6  = d.get('phase6_closing', {})
        dsb = d.get('dashboard', {})

        axes  = ['求人マッチング説明','動機付け','懸念解消','認識統一','次回接続']
        scores = {ax: gd.get(ax, {}).get('score', 0) for ax in axes}
        total  = sum(scores.values())
        grade  = ov.get('grade', '')
        if not grade:
            if total >= 13: grade = 'S'
            elif total >= 10: grade = 'A'
            elif total >= 7:  grade = 'B'
            elif total >= 4:  grade = 'C'
            else:             grade = 'D'

        proposed = d.get('提案件数') or 0
        accepted = d.get('応諾件数') or 0

        records.append({
            '_file':    filename,
            '_path':    path_or_id,
            '_raw':     d,
            'CA':       d.get('ca', ''),
            'グリップ':  d.get('grip', 'X'),
            '候補者':    d.get('candidate', ''),
            'グレード':  grade,
            '総合スコア': total,
            # 5軸
            '求人マッチング説明': scores['求人マッチング説明'],
            '動機付け':   scores['動機付け'],
            '懸念解消':   scores['懸念解消'],
            '認識統一':   scores['認識統一'],
            '次回接続':   scores['次回接続'],
            # 提案結果
            '提案件数':   proposed,
            '応諾件数':   accepted,
            '応諾率':     round(accepted / proposed * 100) if proposed > 0 else 0,
            # フェーズ3
            '提案スタイル':   p3.get('提案スタイル') or dsb.get('提案スタイル', ''),
            'ミライ型スコア':  p3.get('ミライ型スコア') if p3.get('ミライ型スコア') is not None else dsb.get('ミライ型スコア'),
            'マッチングスコア': p3.get('マッチングスコア'),
            # フェーズ5
            '渋り検出':       p5.get('渋り検出', False),
            '能動性スコア':    p5.get('能動性スコア'),
            '応諾の質スコア':  p5.get('応諾の質スコア'),
            # 行動指標
            '求職者発話比率':  round(bh.get('求職者発話比率', 0) * 100),
            '懸念スルー率':    bh.get('懸念スルー率', 0),
            'バックトラッキング': bh.get('バックトラッキング', 0),
            'フィラー回数':    bh.get('フィラー回数', 0),
            'ポジティブ反応':  bh.get('ポジティブ反応', 0),
        })
    except Exception:
        pass

# ══════════════════════════════════════════════════════
# データ読み込み
# ══════════════════════════════════════════════════════
df_all = load_proposal_records()

if df_all.empty:
    st.warning('求人提案の分析データが見つかりません。まず「求人提案FBツール」で分析してください。')
    st.stop()

# ── サイドバー：フィルター ────────────────────────────────
with st.sidebar:
    st.header('🔍 フィルター')

    ca_opts    = ['全員'] + sorted(df_all['CA'].dropna().unique().tolist())
    grip_opts  = sorted(df_all['グリップ'].dropna().unique().tolist())
    grade_opts = ['S', 'A', 'B', 'C', 'D']

    sel_ca    = st.selectbox('CA', ca_opts)
    sel_grip  = st.multiselect('グリップランク', grip_opts, default=grip_opts)
    sel_grade = st.multiselect('グレード', grade_opts, default=grade_opts)
    cand_q    = st.text_input('候補者名で検索', placeholder='例：岡寺')

    st.divider()
    st.caption(f'総データ数：{len(df_all)}件')
    if st.button('🔄 データを再読み込み'):
        with st.spinner('サマリーを再構築中...'):
            rebuild_proposal_summary()
        load_proposal_records.clear()
        st.rerun()

# ── フィルタリング ────────────────────────────────────────
df = df_all.copy()
if sel_ca   != '全員':      df = df[df['CA'] == sel_ca]
if sel_grip:                df = df[df['グリップ'].isin(sel_grip)]
if sel_grade:               df = df[df['グレード'].isin(sel_grade)]
if cand_q:                  df = df[df['候補者'].str.contains(cand_q, na=False)]

# ══════════════════════════════════════════════════════
# ヘッダー
# ══════════════════════════════════════════════════════
st.title('💼 求人提案 分析ダッシュボード')
st.caption(f'フィルター後：{len(df)}件 / 全{len(df_all)}件')

if df.empty:
    st.info('条件に一致するデータがありません。フィルターを変更してください。')
    st.stop()

# ══════════════════════════════════════════════════════
# KPIサマリー
# ══════════════════════════════════════════════════════
st.markdown('<div class="section-title">📈 サマリー指標</div>', unsafe_allow_html=True)

k1,k2,k3,k4,k5,k6,k7 = st.columns(7)
k1.metric('件数',           f'{len(df)}件')
k2.metric('平均総合スコア',  f'{df["総合スコア"].mean():.1f}/15')
avg_rate = df["応諾率"].mean()
k3.metric('平均応諾率',      f'{avg_rate:.0f}%',
          delta='✅良好' if avg_rate >= 50 else '⚠️目標50%', delta_color='off')
avg_mirai = df["ミライ型スコア"].dropna().mean()
k4.metric('平均ミライ型スコア', f'{avg_mirai:.1f}/3' if not pd.isna(avg_mirai) else '─')
avg_slip = df["懸念スルー率"].mean()
k5.metric('平均懸念スルー率', f'{avg_slip:.0f}%',
          delta='✅良好' if avg_slip <= 50 else '⚠️要改善', delta_color='off')
k6.metric('平均求職者発話比率', f'{df["求職者発話比率"].mean():.0f}%')
style_a_rate = (df["提案スタイル"] == 'A').sum()
k7.metric('提案スタイルA型', f'{style_a_rate}/{len(df)}件',
          delta=f'{style_a_rate/len(df)*100:.0f}%', delta_color='off')

st.divider()

# ══════════════════════════════════════════════════════
# CA別比較
# ══════════════════════════════════════════════════════
if sel_ca == '全員' and df['CA'].nunique() > 1:
    st.markdown('<div class="section-title">👥 CA別比較</div>', unsafe_allow_html=True)

    ca_grp = df.groupby('CA').agg(
        件数=('総合スコア', 'count'),
        平均スコア=('総合スコア', 'mean'),
        平均応諾率=('応諾率', 'mean'),
        求人マッチング=('求人マッチング説明', 'mean'),
        動機付け=('動機付け', 'mean'),
        懸念解消=('懸念解消', 'mean'),
        認識統一=('認識統一', 'mean'),
        次回接続=('次回接続', 'mean'),
        ミライ型=('ミライ型スコア', 'mean'),
        懸念スルー率=('懸念スルー率', 'mean'),
        発話比率=('求職者発話比率', 'mean'),
        フィラー=('フィラー回数', 'mean'),
    ).round(2).reset_index()

    tab1, tab2, tab3 = st.tabs(['📊 5軸スコア', '🎯 求人提案指標', '📈 行動指標'])

    with tab1:
        c1, c2 = st.columns([1, 1])
        with c1:
            st.bar_chart(ca_grp.set_index('CA')[['求人マッチング','動機付け','懸念解消','認識統一','次回接続']])
        with c2:
            st.dataframe(
                ca_grp[['CA','件数','平均スコア','求人マッチング','動機付け','懸念解消','認識統一','次回接続']]
                .sort_values('平均スコア', ascending=False),
                use_container_width=True, hide_index=True)

    with tab2:
        c1, c2 = st.columns([1, 1])
        with c1:
            st.bar_chart(ca_grp.set_index('CA')[['平均応諾率','ミライ型']])
        with c2:
            style_tbl = df.groupby(['CA','提案スタイル']).size().unstack(fill_value=0).reset_index()
            st.dataframe(
                ca_grp[['CA','平均応諾率','ミライ型','懸念スルー率']].sort_values('平均応諾率', ascending=False),
                use_container_width=True, hide_index=True)
            if not style_tbl.empty:
                st.caption('提案スタイル内訳（A=職種軸型 / B=一社ずつ型）')
                st.dataframe(style_tbl, use_container_width=True, hide_index=True)

    with tab3:
        c1, c2 = st.columns([1, 1])
        with c1:
            st.bar_chart(ca_grp.set_index('CA')[['発話比率','フィラー']])
        with c2:
            st.dataframe(
                ca_grp[['CA','発話比率','フィラー','懸念スルー率']].sort_values('発話比率', ascending=False),
                use_container_width=True, hide_index=True)

    st.divider()

# ══════════════════════════════════════════════════════
# 面談一覧
# ══════════════════════════════════════════════════════
st.markdown('<div class="section-title">📋 面談一覧　✅ チェックした候補者の詳細を下に表示</div>',
            unsafe_allow_html=True)

sort_col = st.radio('並び替え',
    ['総合スコア','応諾率','ミライ型スコア','懸念スルー率','CA','グリップ'],
    horizontal=True)
sort_asc = st.toggle('昇順', value=False)
df_sorted = df.sort_values(sort_col, ascending=sort_asc).reset_index(drop=True)

if 'proposal_checked' not in st.session_state:
    st.session_state['proposal_checked'] = set()

display_cols = ['CA','グリップ','候補者','グレード','総合スコア',
                '提案件数','応諾件数','応諾率',
                '提案スタイル','ミライ型スコア','マッチングスコア',
                '懸念スルー率','求職者発話比率','フィラー回数']

df_edit = df_sorted[display_cols].copy()
df_edit.insert(0, '選択', df_sorted['_path'].isin(st.session_state['proposal_checked']))

edited = st.data_editor(
    df_edit,
    use_container_width=True,
    hide_index=True,
    column_config={
        '選択':           st.column_config.CheckboxColumn('選択', default=False),
        '総合スコア':     st.column_config.ProgressColumn('総合スコア', min_value=0, max_value=15, format='%d'),
        '応諾率':         st.column_config.NumberColumn('応諾率', format='%d%%'),
        'ミライ型スコア': st.column_config.NumberColumn('ミライ型', format='%d/3'),
        'マッチングスコア': st.column_config.NumberColumn('マッチング', format='%d/3'),
        '懸念スルー率':   st.column_config.NumberColumn('懸念スルー率', format='%d%%'),
        '求職者発話比率': st.column_config.NumberColumn('発話比率', format='%d%%'),
    }
)

checked_paths = set()
for i, row in edited.iterrows():
    if row['選択']:
        checked_paths.add(df_sorted.at[i, '_path'])
st.session_state['proposal_checked'] = checked_paths

selected_rows = df_sorted[df_sorted['_path'].isin(checked_paths)]

st.caption(f'{len(checked_paths)}件選択中')
if st.button('選択をクリア'):
    st.session_state['proposal_checked'] = set()
    st.rerun()

# ══════════════════════════════════════════════════════
# 選択した面談の詳細表示
# ══════════════════════════════════════════════════════
if not selected_rows.empty:
    st.divider()
    st.markdown('<div class="section-title">🔍 選択した面談の詳細</div>', unsafe_allow_html=True)

    for _, row in selected_rows.iterrows():
        d = row.get('_raw') if isinstance(row.get('_raw'), dict) else None
        if d is None:
            # サマリー方式では _raw がないので _path（file ID or ローカルパス）から取得
            try:
                path_or_id = row['_path']
                if Path(path_or_id).exists():
                    d = json.loads(Path(path_or_id).read_text(encoding='utf-8'))
                else:
                    d = gdrive_download_json(path_or_id)
            except Exception:
                pass
        if not isinstance(d, dict):
            st.warning(f'{row.get("候補者","?")} のデータ取得に失敗しました')
            continue

        ov    = d.get('overall', {})
        gd    = d.get('grip_drivers', {})
        p1    = d.get('phase1_opening', {})
        p2    = d.get('phase2_basics', {})
        p3    = d.get('phase3_pitch', {})
        p4    = d.get('phase4_proposal', {})
        p5    = d.get('phase5_agreement', {})
        p6    = d.get('phase6_closing', {})
        bh    = d.get('behaviors', {})
        ca_an = d.get('concern_analysis', {})
        phrases = d.get('next_phrases', [])

        grade       = ov.get('grade', '─')
        grade_color = {'S':'#1a5276','A':'#1e8449','B':'#2471a3','C':'#d35400','D':'#c0392b'}.get(grade, '#555')
        proposed    = d.get('提案件数') or 0
        accepted    = d.get('応諾件数') or 0

        with st.expander(
            f"**{d.get('ca','')}CA  /  {d.get('candidate','')}  /  グリップ{d.get('grip','')}  "
            f"/  提案{proposed}件・応諾{accepted}件  /  グレード{grade}**",
            expanded=True
        ):
            # ── グレード + 応諾率 ──────────────────────
            col_g, col_r = st.columns([2, 1])
            with col_g:
                st.markdown(
                    f'<div style="background:{grade_color};color:white;padding:12px 18px;'
                    f'border-radius:8px"><span style="font-size:2rem;font-weight:bold">{grade}</span>'
                    f'&nbsp;&nbsp;{ov.get("grade_reason","")}</div>',
                    unsafe_allow_html=True)
            with col_r:
                if proposed > 0:
                    rate = accepted / proposed * 100
                    cls = 'score-high' if rate >= 50 else 'score-low'
                    st.markdown(
                        f'<div class="{cls}" style="text-align:center;font-size:1.4rem;font-weight:bold">'
                        f'応諾率 {rate:.0f}%<br><small>{accepted}/{proposed}件</small></div>',
                        unsafe_allow_html=True)

            st.markdown('---')

            # ── サマリー ───────────────────────────────
            col_s, col_i = st.columns(2)
            with col_s:
                st.markdown('**💪 強み**')
                st.success(ov.get('top_strength', '─'))
            with col_i:
                st.markdown('**🎯 改善TOP3**')
                for idx, issue in enumerate(ov.get('top_issues', [])[:3]):
                    if isinstance(issue, dict):
                        st.markdown(f'**{idx+1}. {issue.get("issue","")}**')
                        if issue.get('detail'): st.caption('📌 ' + issue['detail'])
                        if issue.get('fix'):    st.code(issue['fix'], language=None)

            if ov.get('one_thing'):
                st.warning('🚀 **次に試すこと:** ' + ov['one_thing'])

            st.markdown('---')

            # ── 5軸ルーブリック ────────────────────────
            st.markdown('**🎯 ルーブリック採点（5軸）**')
            AXES = ['求人マッチング説明','動機付け','懸念解消','認識統一','次回接続']
            for axis in AXES:
                ax_d = gd.get(axis, {})
                s    = ax_d.get('score', 0)
                bar  = '█' * s + '░' * (3 - s)
                cls  = 'score-high' if s >= 2 else ('score-mid' if s == 1 else 'score-low')
                na   = ax_d.get('next_action', '')
                ev   = ax_d.get('evidence', [])
                st.markdown(
                    f'<div class="{cls}"><b>{axis}</b> &nbsp; <code>{bar}</code> <b>{s}/3</b>'
                    + (f'<br><small>根拠：{sq(ev[0])}</small>' if ev else '')
                    + (f'<br>🚀 {na}' if na else '')
                    + '</div>', unsafe_allow_html=True)

            st.markdown('---')

            # ── フェーズ別評価 ─────────────────────────
            st.markdown('**📌 フェーズ別評価**')

            def _phase_badge(label, ok):
                return f'{"✅" if ok else "❌"} {label}'

            with st.expander('P1　冒頭・前回振り返り', expanded=False):
                c1, c2, c3 = st.columns(3)
                c1.metric('振り返り', '✅' if p1.get('振り返り実施') else '❌')
                c2.metric('温度感確認', '✅' if p1.get('温度感確認') else '❌')
                c3.metric('メリデメ', f'{p1.get("メリデメ説明スコア",0)}/3')
                if p1.get('振り返り評価'): st.caption('📋 ' + p1['振り返り評価'])
                if p1.get('改善アドバイス'): st.info('💡 ' + p1['改善アドバイス'])

            with st.expander('P2　転職基礎説明', expanded=False):
                if not p2.get('対象'):
                    st.caption('対象外（転職活動経験あり）')
                else:
                    c1, c2, c3, c4 = st.columns(4)
                    c1.metric('実施', '✅' if p2.get('実施') else '❌')
                    c2.metric('フロー説明', '✅' if p2.get('転職フロー説明') else '❌')
                    c3.metric('選考数字', '✅' if p2.get('選考基準数字説明') else '❌')
                    c4.metric('必要応募数', '✅' if p2.get('必要応募数説明') else '❌')
                    if p2.get('コメント'): st.caption(p2['コメント'])

            p3_mirai = p3.get('ミライ型スコア', 0)
            p3_label = f'P3　提案職種説明 ★　ミライ型 {p3_mirai}/3　マッチング {p3.get("マッチングスコア",0)}/3'
            with st.expander(p3_label, expanded=True):
                c1, c2, c3, c4 = st.columns(4)
                style = p3.get('提案スタイル', '─')
                c1.metric('提案スタイル', f'{style}型',
                          delta='✅職種軸' if style == 'A' else '⚠️一社ずつ', delta_color='off')
                c2.metric('ミライ型★', f'{p3_mirai}/3')
                c3.metric('マッチング', f'{p3.get("マッチングスコア",0)}/3'
                          + (' ★' if p3.get('プロフェッショナル') else ''))
                c4.metric('市場トレンド', '✅' if p3.get('市場トレンド説明') else '❌')
                if p3.get('ミライ型根拠'): st.caption('🔮 ' + p3['ミライ型根拠'])
                if p3.get('ミライ型具体例'):
                    st.markdown(f'<div class="concern-hit">✅ ミライ型の例: {sq(p3["ミライ型具体例"])}</div>',
                                unsafe_allow_html=True)
                if p3.get('マッチングスコア根拠'): st.caption('🎯 ' + p3['マッチングスコア根拠'])
                if p3.get('改善アドバイス'): st.warning('💡 ' + p3['改善アドバイス'])

            p4_label = (f'P4　求人提案　価値提供 {p4.get("価値提供スコア",0)}/3　'
                        f'難易度 {p4.get("難易度伝達スコア",0)}/3') if p4 else 'P4　求人提案　（データなし）'
            with st.expander(p4_label, expanded=False):
                if not p4:
                    st.caption('フェーズ4のデータなし')
                else:
                    c1, c2, c3 = st.columns(3)
                    c1.metric('価値提供', f'{p4.get("価値提供スコア",0)}/3')
                    c2.metric('難易度伝達', f'{p4.get("難易度伝達スコア",0)}/3')
                    c3.metric('会話バランス', f'{p4.get("CA一方向比率スコア",0)}/3')
                    if p4.get('価値提供コメント'): st.caption(p4['価値提供コメント'])
                    if p4.get('改善アドバイス'): st.warning('💡 ' + p4['改善アドバイス'])

            p5_shiburi = '⚠️渋りあり' if (p5 or {}).get('渋り検出') else ''
            p5_label = (f'P5　応募意思確認 ★　能動性 {p5.get("能動性スコア",0)}/3　'
                        f'応諾の質 {p5.get("応諾の質スコア",0)}/3　{p5_shiburi}') if p5 else 'P5　応募意思確認 ★　（データなし）'
            with st.expander(p5_label, expanded=bool(p5_shiburi)):
                if not p5:
                    st.caption('フェーズ5のデータなし')
                else:
                    c1, c2, c3, c4 = st.columns(4)
                    c1.metric('能動性', f'{p5.get("能動性スコア",0)}/3')
                    c2.metric('応諾の質', f'{p5.get("応諾の質スコア",0)}/3')
                    c3.metric('応諾後不安', '⚠️あり' if p5.get('応諾後不安') else '✅なし')
                    c4.metric('NGワード', '⚠️検出' if p5.get('NGワード検出') else '✅なし')
                    if p5.get('渋り検出'):
                        st.markdown(
                            f'<div class="concern-miss">⚠️ 渋り検出: {sq(p5.get("渋り場面",""))}'
                            f'<br>アウト返し効果: <b>{p5.get("アウト返し効果","─")}</b></div>',
                            unsafe_allow_html=True)
                    if p5.get('改善アドバイス'): st.warning('💡 ' + p5['改善アドバイス'])

            p6_label = f'P6　次回アポ　日程: {p6.get("日程具体性","─")}' if p6 else 'P6　次回アポ'
            with st.expander(p6_label, expanded=False):
                c1, c2, c3 = st.columns(3)
                c1.metric('面接対策アポ', '✅' if p6.get('面接対策アポ') else '❌')
                c2.metric('日程具体性', p6.get('日程具体性', '─'))
                c3.metric('書類確認', '✅' if p6.get('書類次ステップ確認') else '❌')
                if p6.get('コメント'): st.caption(p6['コメント'])
                if p6.get('改善アドバイス'): st.info('💡 ' + p6['改善アドバイス'])

            st.markdown('---')

            # ── 懸念深掘りマップ ───────────────────────
            st.markdown('**🔍 懸念深掘りマップ**')
            miss = [m for m in (ca_an.get('missed_scenes') or []) if isinstance(m, dict)][:5]
            good = [g for g in (ca_an.get('good_scenes') or []) if isinstance(g, dict)][:5]
            if ca_an.get('summary'):
                st.caption(ca_an['summary'])
            if miss:
                st.markdown(f'<div class="sub-title">❌ スルーした場面（代表{len(miss)}件） → こう返すべきだった</div>',
                            unsafe_allow_html=True)
                for ms in miss:
                    col_l, col_r = st.columns(2)
                    with col_l:
                        st.markdown(
                            f'<div class="concern-miss">'
                            f'求職者: <b>{sq(ms.get("cd_text",""))}</b><br>'
                            f'<small>↓ 実際のCA: {sq(ms.get("ca_actual",""))}</small></div>',
                            unsafe_allow_html=True)
                    with col_r:
                        st.markdown(
                            f'<div class="concern-hit">'
                            f'💡 <b style="color:#375623">{sq(ms.get("ca_suggested",""))}</b><br>'
                            f'<small>→ {ms.get("why","")}</small></div>',
                            unsafe_allow_html=True)
            if good:
                st.markdown(f'<div class="sub-title">✅ うまく対応できた場面（代表{len(good)}件）</div>', unsafe_allow_html=True)
                for gs in good:
                    if not isinstance(gs, dict): continue
                    st.markdown(
                        f'<div class="concern-hit">'
                        f'求職者: {sq(gs.get("cd_text",""))}<br>'
                        f'CA: <b>{sq(gs.get("ca_text",""))}</b><br>'
                        f'<small>💡 {gs.get("why_good","")}</small></div>',
                        unsafe_allow_html=True)

            # ── 行動指標 ───────────────────────────────
            st.markdown('---')
            st.markdown('**📈 行動指標**')
            ok_b = lambda v, thr, rev=False: '✅' if (v <= thr if rev else v >= thr) else '❌'
            bh_rows = [
                ('求職者発話比率',  f'{round(bh.get("求職者発話比率",0)*100)}%', '35%以上',  ok_b(bh.get('求職者発話比率',0)*100, 35)),
                ('懸念スルー率',    f'{bh.get("懸念スルー率",0)}%',              '50%以下',  ok_b(bh.get('懸念スルー率',0), 50, rev=True)),
                ('バックトラッキング', f'{bh.get("バックトラッキング",0)}回',     '5回以上',  ok_b(bh.get('バックトラッキング',0), 5)),
                ('マッチング説明数', f'{bh.get("マッチング説明数",0)}回',         '2回以上',  ok_b(bh.get('マッチング説明数',0), 2)),
                ('名前呼称回数',    f'{bh.get("名前呼称回数",0)}回',             '3回以上',  ok_b(bh.get('名前呼称回数',0), 3)),
                ('フィラー回数',    f'{bh.get("フィラー回数",0)}回',             '30回以下', ok_b(bh.get('フィラー回数',0), 30, rev=True)),
                ('ポジティブ反応',  f'{bh.get("ポジティブ反応",0)}回',           '5回以上',  ok_b(bh.get('ポジティブ反応',0), 5)),
            ]
            cols_bh = st.columns(7)
            for col, (label, val, target, judge) in zip(cols_bh, bh_rows):
                col.metric(f'{judge} {label}', val, delta=target, delta_color='off')

            # ── フレーズ集 ───────────────────────────────
            if phrases:
                st.markdown('---')
                st.markdown('**🗣️ 次の面談で使えるフレーズ集**')
                for ph in phrases:
                    if not isinstance(ph, dict) or not ph.get('phrase'): continue
                    st.markdown(
                        f'<div style="margin:6px 0">'
                        f'<small style="background:#f0f0f0;padding:2px 8px;border-radius:10px">'
                        f'{ph.get("situation","")}</small><br>'
                        f'<div class="phrase-box">{sq(ph.get("phrase",""))}</div>'
                        f'<small style="color:#555">💡 {ph.get("why","")}</small></div>',
                        unsafe_allow_html=True)

            # ── アクション（Word / Slack / 再分析） ──────
            st.markdown('---')
            st.markdown('<div class="section-title">⚡ アクション</div>', unsafe_allow_html=True)
            uid = f'{d.get("ca","")}_{d.get("candidate","")}_{d.get("grip","")}'
            act1, act2, act3 = st.columns(3)

            # Word出力
            with act1:
                st.markdown('**📄 Word文書として出力**')
                if st.button('Word文書を生成', key=f'gen_word_{uid}', use_container_width=True):
                    with st.spinner('生成中...'):
                        try:
                            docx_bytes = generate_word_doc_proposal(d)
                            fname = f'求人提案分析_{d.get("ca","")}_{d.get("candidate","")}.docx'
                            st.download_button(
                                '📥 ダウンロード', data=docx_bytes, file_name=fname,
                                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                use_container_width=True, key=f'dl_{uid}')
                        except Exception as e:
                            st.error(f'生成失敗：{e}')

            # Slack送信
            with act2:
                st.markdown('**💬 Slackに送信**')
                if st.button('Slackに送信', key=f'send_slack_{uid}', use_container_width=True):
                    if not SLACK_WEBHOOK:
                        st.error('Slack Webhook URLが設定されていません')
                    else:
                        with st.spinner('送信中...'):
                            ok_flag, msg = send_to_slack_proposal(SLACK_WEBHOOK, d)
                            if ok_flag: st.success('✅ 送信しました')
                            else:       st.error(f'送信失敗：{msg}')

            # 再分析
            with act3:
                st.markdown('**🔄 再分析する**')
                api_key = os.environ.get('ANTHROPIC_API_KEY', '')
                if not api_key:
                    st.error('APIキー未設定')
                else:
                    has_p4 = bool(d.get('phase4_proposal'))
                    has_p1 = bool(d.get('phase1_opening'))

                    if not has_p1:
                        st.caption('⚠️ フェーズ1〜3のデータなし → 最初から再分析')
                        rerun_label = '🔄 最初から再分析'
                        rerun_phase = 'all'
                    elif not has_p4:
                        st.caption('⚠️ フェーズ4〜6のデータなし → フェーズ4〜6のみ再分析')
                        rerun_label = '🔄 フェーズ4〜6のみ再実行'
                        rerun_phase = 'late'
                    else:
                        rerun_label = '🔄 全フェーズ再分析'
                        rerun_phase = 'all'

                    rerun_mode = st.radio(
                        '再分析モード',
                        ['全フェーズ', 'フェーズ4〜6のみ'],
                        horizontal=True, key=f'rerun_mode_{uid}',
                        index=0 if rerun_phase == 'all' else 1)
                    rerun_phase = 'all' if rerun_mode == '全フェーズ' else 'late'

                    if st.button(rerun_label, key=f'reanalyze_{uid}',
                                 use_container_width=True, type='primary'):
                        client_inst = anthropic.Anthropic(api_key=api_key)
                        phase_label = '全フェーズ' if rerun_phase == 'all' else 'フェーズ4〜6'
                        with st.spinner(f'🤖 {phase_label}を再分析中...'):
                            new_early, new_late, err = reanalyze_proposal(d, client_inst, rerun_phase)
                        if err:
                            st.error(f'再分析失敗：{err}')
                        else:
                            load_proposal_records.clear()
                            st.success('✅ 再分析完了・保存しました。ページを再読み込みしてください。')
                            st.rerun()
