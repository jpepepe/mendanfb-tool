# -*- coding: utf-8 -*-
"""
初回面談 FB自動生成ツール v2  ─ 深掘り強化版
使い方: streamlit run 面談FB_ツール.py
"""

import streamlit as st
import json, re, zipfile, os, tempfile, unicodedata, io
from typing import Optional
from pathlib import Path
from xml.etree import ElementTree as ET
import anthropic
import requests
from gdrive import upload_json, upload_json as gdrive_upload_json, download_json as gdrive_download_json
from analysis_core import safe_json_loads, MAX_TRANSCRIPT_CHARS

# ── スタイル ──────────────────────────────────────────────
st.markdown("""
<style>
.score-box { background:#f0f4ff; border-left:4px solid #2E75B6; padding:10px 16px; margin:6px 0; border-radius:4px; }
.score-high { border-left-color:#375623; background:#e2efda; }
.score-mid  { border-left-color:#c55a11; background:#fce4d6; }
.score-low  { border-left-color:#c00000; background:#fcecea; }
.check-ok   { color:#375623; font-weight:bold; }
.check-ng   { color:#c00000; font-weight:bold; }
.check-warn { color:#c55a11; font-weight:bold; }
.section-title { background:#1F3864; color:white; padding:8px 16px; border-radius:6px; font-weight:bold; margin:16px 0 8px 0; }
.sub-title { background:#2E75B6; color:white; padding:6px 14px; border-radius:4px; font-weight:bold; margin:10px 0 6px 0; font-size:0.9rem; }
.emotion-hit  { background:#e2efda; border-left:4px solid #375623; padding:8px 12px; border-radius:4px; margin:6px 0; }
.emotion-miss { background:#fcecea; border-left:4px solid #c00000; padding:8px 12px; border-radius:4px; margin:6px 0; }
.bt-hit  { background:#EBF5FB; border-left:4px solid #2E75B6; padding:8px 12px; border-radius:4px; margin:6px 0; }
.bt-miss { background:#FEF9E7; border-left:4px solid #F39C12; padding:8px 12px; border-radius:4px; margin:6px 0; }
.sd-hit  { background:#F0F0F0; border-left:4px solid #555; padding:8px 12px; border-radius:4px; margin:6px 0; }
.phrase-box { background:#1F3864; color:white; padding:10px 14px; border-radius:6px; margin:6px 0; font-family:monospace; }
.vertical-drill { border-left:3px solid #2E75B6; padding:4px 10px; margin:3px 0; }
</style>
""", unsafe_allow_html=True)

# ── ヘルパー ─────────────────────────────────────────────
def sq(text: str) -> str:
    """「」の二重表示を防ぐ：既に括弧があればそのまま、なければ付ける"""
    t = (text or '').strip()
    if t.startswith('「') and t.endswith('」'):
        return t
    return f'「{t}」'

def strip_q(text: str) -> str:
    """表示側で「」を付けたい時に既存の括弧を除去する"""
    t = (text or '').strip()
    if t.startswith('「'): t = t[1:]
    if t.endswith('」'):   t = t[:-1]
    return t

# ── 定数・パターン ────────────────────────────────────────
FILLER_PAT   = re.compile(r'えー+|あのー+|えっと+|まあ+|うーん+|んー+|あー+|なんか(?!ら)|ちょっと(?=\s|、|。|$)')
EMOTION_PAT  = re.compile(r'楽し|嬉し|やりがい|好き|辛|しんどい|大変|不安|怖|達成感|充実|面白|嫌|辞め|悩ん|悔し|感動|驚|嫌い|疲れ')
DRILL_PAT    = re.compile(r'なぜ|なんで|どうして|どんな.*とき|一番.*何|きっかけ|どうやって|どのよう|どういう.*気持|何が.*よかっ|何が.*楽し|どんな.*感じ')
BT_PAT       = re.compile(r'おっしゃって|つまり|要は|ということ|まとめると|先ほど|さっき|先ほどの|おっしゃった.*が|言ってた|言われてた')
SELF_DISC_PAT= re.compile(r'私も|私自身|私が.*とき|私.*経験|僕も|実は私|私.*以前|私.*昔|私.*感じ|私.*思っ|私.*とって|私の.*経験')
POS_PAT      = re.compile(r'すごい|いいです|素晴らしい|さすが|面白い|いいですね|なるほど.*そう|それは.*いい|それって.*すごい')
EMPATHY_PAT  = re.compile(r'そうですよね|分かります|わかります|大変.*でした|私も|私自身|私が.*とき|僕も')
RUSH_PAT     = re.compile(r'早め|急い|早く.*方が|すぐに|早い段階')
AUTH_PAT     = re.compile(r'絶対|確実に|間違いなく|業界的に|私の経験')
SUMMARY_PAT  = re.compile(r'整理すると|つまり.*さん|まとめると|という理解|合ってますか')
NEXT_PAT     = re.compile(r'また.*お話|次回|別日|改めて.*お時間|求人.*お送り|LINE|ライン')
MUST_PAT     = re.compile(r'合って.*ると思|向いて.*ると|別の.*方向|こっちの方が|聞いてて.*思|実は.*方が')
CLOSE_Q_PAT  = re.compile(r'今日.*どうでした|何か.*気になる|聞いておきたい|正直.*どう|今.*気持ち')

RUBRIC_PATH  = Path(__file__).parent / "rubric_初回面談.md"
OUTPUT_JSON  = Path(__file__).parent / "output" / "json"
OUTPUT_UTT   = Path(__file__).parent / "output" / "utterances"
OUTPUT_JSON.mkdir(parents=True, exist_ok=True)
OUTPUT_UTT.mkdir(parents=True, exist_ok=True)

# ── 保存 ─────────────────────────────────────────────────
def save_results(ca, grip, candidate, meeting_type, fmt,
                 utterances, metrics, claude_result, deep_result=None):
    safe_grip = grip if grip != '未入力' else 'X'
    key = f"{ca}_{safe_grip}_{candidate}_{meeting_type}"

    # ── 話者分離データ ──────────────────────────────
    utt_path = OUTPUT_UTT / f"{key}.json"
    speakers = list(set(u['speaker'] for u in utterances))
    utt_path.write_text(json.dumps({
        "ca": ca, "grip": safe_grip, "candidate": candidate,
        "meeting_type": meeting_type, "format": fmt,
        "speakers_detected": speakers,
        "utterance_count": len(utterances),
        "utterances": utterances,
    }, ensure_ascii=False, indent=2), encoding='utf-8')

    # ── 分析結果（スコアリング ＋ 深掘り分析）──────
    bh_calc = {k: v for k, v in metrics.items()
               if k not in ('ca_texts','cd_texts','utterances','is_ca','is_cd',
                            '感情モーメント全件','感情スルー例','感情深掘り例',
                            '縦深掘りシーケンス','自己開示モーメント',
                            'バックトラッキングモーメント','クロージングCA','クロージングCD')}
    json_data = {
        "ca": ca, "grip": safe_grip, "candidate": candidate,
        "meeting_type": meeting_type, "format": fmt,
        # Call 1: スコアリング
        "grip_drivers":  claude_result.get('grip_drivers', {}),
        "behaviors":     {**bh_calc, **claude_result.get('behaviors', {})},
        "overall":       claude_result.get('overall', {}),
        "notes":         claude_result.get('notes', ''),
        # Call 2: 深掘り分析（取得できた場合のみ）
        "emotion_drill_analysis":  (deep_result or {}).get('emotion_drill_analysis', {}),
        "self_disclosure_analysis":(deep_result or {}).get('self_disclosure_analysis', {}),
        "backtrack_analysis":      (deep_result or {}).get('backtrack_analysis', {}),
        "next_phrases":            (deep_result or {}).get('next_phrases', []),
    }
    json_path = OUTPUT_JSON / f"{key}.json"
    json_path.write_text(json.dumps(json_data, ensure_ascii=False, indent=2), encoding='utf-8')
    # Google Drive にも保存
    try:
        upload_json(f"{key}.json", json_data, subfolder="json")
    except Exception as e:
        st.warning(f"Google Drive保存失敗（ローカルには保存済み）: {e}")
    # utterances も Drive に保存（ダッシュボードからの再分析に必要）
    try:
        utt_for_drive = {
            "ca": ca, "grip": safe_grip, "candidate": candidate,
            "meeting_type": meeting_type, "format": fmt,
            "speakers_detected": speakers,
            "utterance_count": len(utterances),
            "utterances": utterances,
        }
        upload_json(f"{key}.json", utt_for_drive, subfolder="utterances")
    except Exception as e:
        st.warning(f"文字起こしDrive保存失敗（再分析が使えない場合があります）: {e}")
    return utt_path, json_path

# ── セルフチェック（app.py用：ダッシュボードと共通の定数・関数） ──
AXES_APP = ['意向','適正','条件','認識統一','気づき']
AXIS_SHORT_APP = {'意向':'意向把握','適正':'適正把握','条件':'条件把握',
                  '認識統一':'認識統一','気づき':'気づき付与'}
AXIS_DEF_APP = {
    '意向':    '価値観・やりがいを引き出し、応募企業に固執させず意向を広げられたか',
    '適正':    '経験・強み・適性を具体的に把握できたか',
    '条件':    'Must/Betterを確認し、期待値を調整できたか',
    '認識統一':'価値観・強みの要約への同意＋今後のキャリアの方向性への合意が取れたか',
    '気づき':  '他の選択肢・新しい可能性に気づかせられたか',
}
SCORE_LADDER_APP = [
    '**0点** ＝ 未実施。その観点に触れていない',
    '**1点** ＝ 触れたが浅い・一方的な説明だけで終わった',
    '**2点** ＝ 把握できたが、本人の確認・同意が弱い（「はい」止まり）',
    '**3点** ＝ 根拠を引き出し、本人の明示的な同意・反応まで取れた',
]

def save_selfcheck_app(ref_file, ca, grip, candidate, meeting_type,
                       self_scores, behavior_checks, next_one_thing, best_self=''):
    from datetime import datetime
    data = {
        '_ref_file': ref_file,
        'ca': ca, 'grip': grip, 'candidate': candidate, 'meeting_type': meeting_type,
        'self_scores': self_scores,
        'behavior_checks': behavior_checks,
        'best_self': best_self,
        'next_one_thing': next_one_thing,
        'checked_at': datetime.now().isoformat(timespec='seconds'),
    }
    try:
        gdrive_upload_json(f'selfcheck_{ref_file}', data, subfolder='selfcheck')
    except Exception as e:
        st.warning(f'自己採点のDrive保存失敗（記録はされません）: {e}')

def load_selfcheck_app(ref_file: str) -> Optional[dict]:
    from gdrive import download_json_by_name
    try:
        return download_json_by_name(f'selfcheck_{ref_file}', subfolder='selfcheck')
    except Exception:
        return None


# ── Word文書生成 ──────────────────────────────────────────
def generate_word_doc(d: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    def h(text, level=1, color=(31,56,100)):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
        return p

    def body(text, bold=False, color=None, size=10.5):
        p = doc.add_paragraph()
        r = p.add_run(str(text))
        r.bold = bold; r.font.size = Pt(size)
        if color: r.font.color.rgb = RGBColor(*color)
        return p

    def kv(key, val, kc=(46,117,182)):
        p = doc.add_paragraph()
        r1 = p.add_run(f'{key}：'); r1.bold = True
        r1.font.color.rgb = RGBColor(*kc); r1.font.size = Pt(10.5)
        r2 = p.add_run(str(val or '')); r2.font.size = Pt(10.5)

    def tbl(headers, rows, style='Table Grid'):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = style
        for i, h_ in enumerate(headers):
            c = t.rows[0].cells[i]; c.text = h_
            if c.paragraphs[0].runs: c.paragraphs[0].runs[0].bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row): cells[i].text = str(v or '')
        doc.add_paragraph()
        return t

    gd  = d.get('grip_drivers', {})
    bh  = d.get('behaviors', {})
    ov  = d.get('overall', {})
    eda = d.get('emotion_drill_analysis', {})
    sda = d.get('self_disclosure_analysis', {})
    bta = d.get('backtrack_analysis', {})
    phrases = d.get('next_phrases', [])
    axes = ['意向','適正','条件','認識統一','気づき']
    total = sum(gd.get(ax,{}).get('score',0) for ax in axes)
    grade = ov.get('grade','') or ('S' if total>=13 else 'A' if total>=10 else 'B' if total>=7 else 'C' if total>=4 else 'D')
    AXIS_LBL = {'意向':'意向把握','適正':'適正把握','条件':'条件把握','認識統一':'認識統一','気づき':'気づき付与'}

    t0 = doc.add_heading('初回面談 分析レポート', 0)
    t0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kv('CA名', d.get('ca','')); kv('候補者名', d.get('candidate',''))
    kv('グリップ', d.get('grip','')); kv('面談種別', d.get('meeting_type',''))
    doc.add_paragraph()

    h('■ 1. 総合評価')
    p = doc.add_paragraph()
    r = p.add_run(f'グレード：{grade}　　総合スコア：{total}/15')
    r.bold = True; r.font.size = Pt(14)
    if ov.get('grade_reason'): body(ov['grade_reason'])
    if ov.get('top_strength'): body(f'【最大の強み】{ov["top_strength"]}', color=(55,86,35))
    if ov.get('best_exchange'): body(f'【最も良かったやり取り】{ov["best_exchange"]}', color=(55,86,35))
    if ov.get('missed_moment'): body(f'【最も惜しかった場面】{ov["missed_moment"]}', color=(192,0,0))
    doc.add_paragraph()

    h('■ 2. ルーブリック採点（5軸）')
    tbl(
        ['評価軸','スコア','強み','改善点','次のアクション'],
        [[AXIS_LBL.get(ax,ax),
          f'{gd.get(ax,{}).get("score",0)}/3',
          gd.get(ax,{}).get('strength',''),
          gd.get(ax,{}).get('weakness',''),
          gd.get(ax,{}).get('next_action','')] for ax in axes]
    )

    if ov.get('top_issues'):
        h('■ 3. 優先改善ポイント TOP3')
        for i, issue in enumerate(ov['top_issues'][:3], 1):
            if isinstance(issue, dict):
                body(f'{i}. {issue.get("issue","")}', bold=True)
                if issue.get('detail'): body(f'　📌 {issue["detail"]}')
                if issue.get('fix'):    body(f'　💡 改善例：{issue["fix"]}', color=(55,86,35))
        doc.add_paragraph()

    if ov.get('one_thing'):
        h('■ 4. 次の面談で必ず試してほしいこと')
        body(ov['one_thing'], color=(46,117,182), bold=True)
        doc.add_paragraph()

    if ov.get('closing_eval'):
        h('■ 5. クロージング評価')
        body(ov['closing_eval'])
        doc.add_paragraph()

    h('■ 6. 感情深掘り分析')
    if eda.get('summary'): body(eda['summary'])
    missed = eda.get('missed_scenes') or []
    if missed:
        body('【スルーしてしまった場面 → 改善例】', bold=True)
        for ms in missed:
            body(f'求職者：{ms.get("cd_text","")}')
            body(f'実際のCA：{ms.get("ca_actual","")}', color=(192,0,0))
            body(f'→ こう返すべきだった：{ms.get("ca_suggested","")}', color=(55,86,35))
            if ms.get('why'): body(f'　理由：{ms["why"]}')
            doc.add_paragraph()
    good = eda.get('good_scenes') or []
    if good:
        body('【うまく深掘りできた場面】', bold=True)
        for gs in good:
            body(f'求職者：{gs.get("cd_text","")}')
            body(f'CA（深掘り）：{gs.get("ca_text","")}', color=(55,86,35))
            if gs.get('why_good'): body(f'　良かった点：{gs["why_good"]}')
            doc.add_paragraph()
    if eda.get('vertical_drill_comment'):
        body(f'【縦の深掘りコメント】{eda["vertical_drill_comment"]}')
    doc.add_paragraph()

    h('■ 7. 自己開示の分析')
    if sda.get('summary'): body(sda['summary'])
    for fs in (sda.get('found_scenes') or []):
        body(f'CA：{fs.get("ca_text","")}')
        body(f'　タイミング評価：{fs.get("timing_eval","")}　/ {fs.get("effect","")}', color=(46,117,182))
    missed_sd = sda.get('missed_opportunities') or []
    if missed_sd:
        body('【ここで自己開示できた → 改善例】', bold=True)
        for ms in missed_sd:
            body(f'求職者：{ms.get("cd_text","")}')
            body(f'→ {ms.get("ca_suggested","")}', color=(55,86,35))
    if sda.get('advice'): body(f'アドバイス：{sda["advice"]}')
    doc.add_paragraph()

    h('■ 8. バックトラッキングの分析')
    if bta.get('summary'): body(bta['summary'])
    for fs in (bta.get('found_scenes') or []):
        body(f'参照：{fs.get("referenced_cd","")}')
        body(f'CA：{fs.get("ca_text","")}', color=(46,117,182))
        body(f'　評価：{fs.get("effect","")}')
    missed_bt = bta.get('missed_opportunities') or []
    if missed_bt:
        body('【このキーワードを引用できた → 改善例】', bold=True)
        for ms in missed_bt:
            body(f'キーワード：{ms.get("cd_keyword","")}')
            body(f'→ {ms.get("ca_suggested","")}', color=(55,86,35))
    if bta.get('advice'): body(f'アドバイス：{bta["advice"]}')
    doc.add_paragraph()

    h('■ 9. 行動指標')
    ok_mark = lambda v, thr: '✅' if v >= thr else '❌'
    tbl(
        ['指標','値','目標','判定'],
        [
            ('求職者発話比率',  f'{round(bh.get("求職者発話比率",0)*100)}%',  '40%以上',  ok_mark(bh.get('求職者発話比率',0)*100, 40)),
            ('後半求職者比率',  f'{round(bh.get("後半求職者比率",0)*100)}%',  '50%以上',  ok_mark(bh.get('後半求職者比率',0)*100, 50)),
            ('フィラー回数',    f'{bh.get("フィラー回数",0)}回',               '30回以下',  '✅' if bh.get('フィラー回数',0)<=30 else '❌'),
            ('ポジティブ反応',  f'{bh.get("ポジティブ反応",0)}回',             '5回以上',   ok_mark(bh.get('ポジティブ反応',0), 5)),
            ('感情スルー率',    f'{bh.get("感情スルー率",0)}%',                '50%以下',   '✅' if bh.get('感情スルー率',0)<=50 else '❌'),
            ('価値観深掘り',    f'{bh.get("深掘り_価値観",0)}回',              '4回以上',   ok_mark(bh.get('深掘り_価値観',0), 4)),
            ('縦深掘り最大',    f'{bh.get("縦深掘り最大",0)}回連続',           '3回以上',   ok_mark(bh.get('縦深掘り最大',0), 3)),
            ('バックトラッキング',f'{bh.get("バックトラッキング",0)}回',        '5回以上',   ok_mark(bh.get('バックトラッキング',0), 5)),
            ('自己開示回数',    f'{bh.get("自己開示回数",0)}回',               '2回以上',   ok_mark(bh.get('自己開示回数',0), 2)),
            ('名前呼称回数',    f'{bh.get("名前呼称回数",0)}回',               '3回以上',   ok_mark(bh.get('名前呼称回数',0), 3)),
            ('MUST提案',        '✅' if bh.get('MUST提案') else '❌',           'あり',      '✅' if bh.get('MUST提案') else '❌'),
            ('次回アポ確定',    '✅' if bh.get('次回アポ確定') else '❌',       'あり',      '✅' if bh.get('次回アポ確定') else '❌'),
        ]
    )

    if phrases:
        h('■ 10. 次の面談で使えるフレーズ集')
        for ph in phrases:
            if not ph.get('phrase'): continue
            body(f'【{ph.get("situation","")}】', bold=True)
            p_ph = doc.add_paragraph()
            r_ph = p_ph.add_run(ph.get('phrase',''))
            r_ph.font.color.rgb = RGBColor(31,56,100)
            r_ph.bold = True; r_ph.font.size = Pt(11)
            if ph.get('why'): body(f'　💡 {ph["why"]}')
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Slack送信 ─────────────────────────────────────────────
def send_to_slack(webhook_url: str, d: dict) -> tuple:
    gd  = d.get('grip_drivers', {})
    bh  = d.get('behaviors', {})
    ov  = d.get('overall', {})
    eda = d.get('emotion_drill_analysis', {})
    sda = d.get('self_disclosure_analysis', {})
    bta = d.get('backtrack_analysis', {})
    phrases = d.get('next_phrases', [])

    axes  = ['意向','適正','条件','認識統一','気づき']
    total = sum(gd.get(ax,{}).get('score',0) for ax in axes)
    grade = ov.get('grade','') or ('S' if total>=13 else 'A' if total>=10 else 'B' if total>=7 else 'C' if total>=4 else 'D')
    g_emoji = {'S':'🏆','A':'🟢','B':'🔵','C':'🟡','D':'🔴'}.get(grade,'⚪')

    def ok(v, thr, rev=False):
        return '✅' if (v <= thr if rev else v >= thr) else '❌'

    score_lines = '\n'.join(
        f"　{ax}：{'█'*gd.get(ax,{}).get('score',0)}{'░'*(3-gd.get(ax,{}).get('score',0))} "
        f"{gd.get(ax,{}).get('score',0)}/3　{gd.get(ax,{}).get('weakness','')}"
        for ax in axes)

    issues_text = ''
    for i, issue in enumerate(ov.get('top_issues',[])[:3], 1):
        if isinstance(issue, dict):
            issues_text += f'{i}. *{issue.get("issue","")}*\n　{issue.get("detail","")}\n　💡 {issue.get("fix","")}\n'

    eda_text = eda.get('summary','')
    missed = eda.get('missed_scenes') or []
    if missed:
        eda_text += '\n*スルーした場面（改善例）:*'
        for ms in missed[:2]:
            eda_text += f'\n　求職者：「{ms.get("cd_text","")}」\n　→ `{ms.get("ca_suggested","")}`'
    good = eda.get('good_scenes') or []
    if good:
        eda_text += '\n*うまく深掘りできた場面:*'
        for gs in good[:1]:
            eda_text += f'\n　求職者：「{gs.get("cd_text","")}」\n　CA：`{gs.get("ca_text","")}`'
    if eda.get('vertical_drill_comment'):
        eda_text += f'\n縦の深掘り：{eda["vertical_drill_comment"]}'

    sda_text = sda.get('summary','')
    for ms in (sda.get('missed_opportunities') or [])[:2]:
        sda_text += f'\n　改善例：`{ms.get("ca_suggested","")}`'
    if sda.get('advice'): sda_text += f'\n💡 {sda["advice"]}'

    bta_text = bta.get('summary','')
    for ms in (bta.get('missed_opportunities') or [])[:2]:
        bta_text += f'\n　改善例：`{ms.get("ca_suggested","")}`'
    if bta.get('advice'): bta_text += f'\n💡 {bta["advice"]}'

    phrases_text = ''
    for ph in phrases:
        if not ph.get('phrase'): continue
        phrases_text += f'*{ph.get("situation","")}*\n　`{ph.get("phrase","")}` — {ph.get("why","")}\n'

    blocks = [
        {"type": "header", "text": {"type": "plain_text",
            "text": f"📊 初回面談 分析レポート｜{d.get('candidate','')}"}},
        {"type": "section", "fields": [
            {"type": "mrkdwn", "text": f"*CA名:* {d.get('ca','')}"},
            {"type": "mrkdwn", "text": f"*グリップ:* {d.get('grip','')}"},
            {"type": "mrkdwn", "text": f"*面談種別:* {d.get('meeting_type','')}"},
            {"type": "mrkdwn", "text": f"*グレード:* {g_emoji} {grade}　スコア: {total}/15"},
        ]},
        {"type": "divider"},
        {"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🎯 ルーブリック採点（5軸）*\n{score_lines}"}},
        {"type": "divider"},
    ]

    if ov.get('top_strength') or ov.get('best_exchange') or ov.get('missed_moment'):
        txt = ''
        if ov.get('top_strength'):   txt += f'*💪 最大の強み*\n{ov["top_strength"]}\n'
        if ov.get('best_exchange'):  txt += f'*👍 最も良かったやり取り*\n{ov["best_exchange"]}\n'
        if ov.get('missed_moment'):  txt += f'*⚠️ 最も惜しかった場面*\n{ov["missed_moment"]}\n'
        blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": txt.strip()}})
        blocks.append({"type": "divider"})

    if issues_text:
        blocks.append({"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🔧 優先改善ポイント TOP3*\n{issues_text.strip()}"}})
        blocks.append({"type": "divider"})

    if ov.get('one_thing'):
        blocks.append({"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🚀 次の面談で必ず試すこと*\n{ov['one_thing']}"}})
        blocks.append({"type": "divider"})

    if ov.get('closing_eval'):
        blocks.append({"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🏁 クロージング評価*\n{ov['closing_eval']}"}})
        blocks.append({"type": "divider"})

    if eda_text:
        blocks.append({"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🔍 感情深掘り分析*\n{eda_text.strip()}"}})
        blocks.append({"type": "divider"})

    sd_bt = ''
    if sda_text: sd_bt += f'*🙋 自己開示*\n{sda_text}\n\n'
    if bta_text: sd_bt += f'*🔁 バックトラッキング*\n{bta_text}'
    if sd_bt:
        blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": sd_bt.strip()}})
        blocks.append({"type": "divider"})

    blocks.append({"type": "section", "fields": [
        {"type": "mrkdwn", "text": f"*感情スルー率:* {ok(bh.get('感情スルー率',0),50,rev=True)} {bh.get('感情スルー率',0)}%"},
        {"type": "mrkdwn", "text": f"*フィラー:* {ok(bh.get('フィラー回数',0),30,rev=True)} {bh.get('フィラー回数',0)}回"},
        {"type": "mrkdwn", "text": f"*ポジティブ反応:* {ok(bh.get('ポジティブ反応',0),5)} {bh.get('ポジティブ反応',0)}回"},
        {"type": "mrkdwn", "text": f"*縦深掘り最大:* {ok(bh.get('縦深掘り最大',0),3)} {bh.get('縦深掘り最大',0)}回"},
        {"type": "mrkdwn", "text": f"*バックトラッキング:* {ok(bh.get('バックトラッキング',0),5)} {bh.get('バックトラッキング',0)}回"},
        {"type": "mrkdwn", "text": f"*自己開示:* {ok(bh.get('自己開示回数',0),2)} {bh.get('自己開示回数',0)}回"},
        {"type": "mrkdwn", "text": f"*発話比率(求職者):* {ok(bh.get('求職者発話比率',0)*100,40)} {round(bh.get('求職者発話比率',0)*100)}%"},
        {"type": "mrkdwn", "text": f"*アポ確定:* {'✅' if bh.get('次回アポ確定') else '❌'}"},
    ]})

    if phrases_text:
        blocks.append({"type": "divider"})
        blocks.append({"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🗣️ 次の面談で使えるフレーズ集*\n{phrases_text.strip()}"}})

    try:
        resp = requests.post(webhook_url, json={"blocks": blocks}, timeout=10)
        if resp.status_code == 200:
            return True, '送信成功'
        return False, f'エラー: {resp.status_code} {resp.text}'
    except Exception as e:
        return False, str(e)


# ── ファイルパーサー ──────────────────────────────────────
def detect_speakers(full_text):
    from collections import Counter
    cands = re.findall(r'(?<![^\s。！？\n])([぀-鿿a-zA-Zー]{2,10}):\s', full_text)
    freq = Counter(cands)
    return [n for n, c in freq.most_common(10) if c >= 3][:2]

def parse_docx(file_bytes):
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        f.write(file_bytes); tmp = f.name
    try:
        with zipfile.ZipFile(tmp) as z:
            with z.open('word/document.xml') as f:
                tree = ET.parse(f)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        paragraphs = tree.getroot().findall('.//w:p', ns)
        full_text = ''
        for p in paragraphs:
            line = ''.join(t.text or '' for t in p.findall('.//w:t', ns)).strip()
            if line: full_text += line + '\n'
        full_text = re.sub(r'^\d{2}:\d{2}:\d{2}\n', '', full_text, flags=re.MULTILINE)
        full_text = re.sub(r'^\d{4}年.+\n',         '', full_text, flags=re.MULTILINE)
        full_text = re.sub(r'^会議\s.+\n',           '', full_text, flags=re.MULTILINE)
        # 全角コロン（：）を半角コロン+スペース（: ）に統一
        full_text = re.sub(r'([぀-鿿a-zA-Zー]{2,10})：', r'\1: ', full_text)
        speakers = detect_speakers(full_text)
        if not speakers: return [], full_text, 'docx'
        pat = re.compile(r'(?:' + '|'.join(re.escape(s) for s in speakers) + r'):\s*')
        utterances = []
        matches = list(pat.finditer(full_text))
        for i, m in enumerate(matches):
            speaker = m.group().rstrip(': \t')
            start = m.end()
            end = matches[i+1].start() if i+1 < len(matches) else len(full_text)
            text = re.sub(r'\s+', ' ', full_text[start:end]).strip()
            if text: utterances.append({'speaker': speaker, 'text': text})
        return utterances, full_text, 'docx'
    finally:
        os.unlink(tmp)

def parse_txt_with_haiku(file_bytes, ca_name, cand_name, client):
    raw = file_bytes.decode('utf-8', errors='ignore')
    raw = re.sub(r'^\d+\t', '', raw, flags=re.MULTILINE)
    prompt = f"""以下は人材紹介会社の面談音声文字起こしです。話者は2名です：
- CA（キャリアアドバイザー）: {ca_name}
- 求職者: {cand_name}
各発言をCAまたは求職者に割り当て、JSON配列のみ返してください。
形式: [{{"speaker": "CA", "text": "発話内容"}}, ...]
文字起こし:
{raw[:40000]}"""
    resp = client.messages.create(
        model='claude-haiku-4-5-20251001', max_tokens=16000,
        messages=[{'role': 'user', 'content': prompt}])
    content = re.sub(r'```(?:json)?\s*', '', resp.content[0].text.strip()).strip()
    try:
        parsed = json.loads(content)
        if isinstance(parsed, list) and parsed: return parsed, raw, 'txt'
    except: pass
    return [{'speaker': '不明', 'text': raw}], raw, 'txt'

# ── 行動指標計算 ──────────────────────────────────────────
def compute_metrics(utterances, ca_name, cand_name):
    ca_base   = re.sub(r'さん$', '', ca_name)
    cand_base = re.sub(r'さん$', '', cand_name)
    all_spk   = set(u['speaker'] for u in utterances)
    ca_spk    = {s for s in all_spk if ca_base in s or s == 'CA'}
    cd_spk_n  = {s for s in all_spk if cand_base in s or s == '求職者'}
    cd_spk    = cd_spk_n if cd_spk_n else (all_spk - ca_spk - {'不明'})

    def is_ca(s): return s in ca_spk
    def is_cd(s): return s in cd_spk

    ca_texts  = [u['text'] for u in utterances if is_ca(u['speaker'])]
    cd_texts  = [u['text'] for u in utterances if is_cd(u['speaker'])]
    all_ca    = ' '.join(ca_texts)
    all_cd    = ' '.join(cd_texts)
    total     = len(all_ca) + len(all_cd)

    # 会話回数
    turns, prev = 0, None
    for u in utterances:
        if len(u['text'].replace(' ','')) < 10: continue
        sp = 'CA' if is_ca(u['speaker']) else ('CD' if is_cd(u['speaker']) else None)
        if sp and sp != prev: turns += 1; prev = sp

    # 名前呼称
    nm_base = re.sub(r'さん$', '', cand_name)
    nm_pat  = re.compile(re.escape(nm_base) + r'さん')
    name_cnt = sum(len(nm_pat.findall(t)) for t in ca_texts)

    # 後半比率
    n = len(utterances)
    latter    = utterances[int(n*0.5):]
    la_ca     = sum(len(u['text']) for u in latter if is_ca(u['speaker']))
    la_cd     = sum(len(u['text']) for u in latter if is_cd(u['speaker']))
    la_tot    = la_ca + la_cd

    # ── 感情ワード場面を全件収集 ─────────────────────────
    emotion_moments_all  = []   # {cd_text, emotion_word, ca_response, drilled, idx}
    emotion_slip_examples = []
    emotion_drill_examples = []
    for i, u in enumerate(utterances[:-1]):
        if is_ca(u['speaker']): continue
        m = EMOTION_PAT.search(u['text'])
        if not m or len(u['text']) < 8: continue
        nxt = utterances[i+1]
        if not is_ca(nxt['speaker']): continue
        drilled = bool(DRILL_PAT.search(nxt['text']) or EMPATHY_PAT.search(nxt['text']))
        record = {
            'idx':       i,
            'cd_text':   u['text'][:80],
            'emotion_word': m.group(),
            'ca_response':  nxt['text'][:80],
            'drilled':   drilled,
        }
        emotion_moments_all.append(record)
        if drilled and len(emotion_drill_examples) < 3:
            emotion_drill_examples.append((u['text'][:70], nxt['text'][:80]))
        elif not drilled and len(emotion_slip_examples) < 4:
            emotion_slip_examples.append((u['text'][:70], nxt['text'][:80]))

    emotion_total = len(emotion_moments_all)
    emotion_skip  = sum(1 for r in emotion_moments_all if not r['drilled'])
    slip_rate     = round(emotion_skip / emotion_total * 100) if emotion_total > 0 else 0

    # ── 縦の深掘り実例 ───────────────────────────────────
    # 「同テーマを連続で深掘りできているか」を測る。求職者の回答ターンでは
    # streakを切らない（一問一答の回答が間に挟まるのは正常）。CAが深掘り以外の
    # 発話（話題転換・事実質問）をした時点でstreakを終了する。
    max_streak = cur_streak = 0
    vertical_drill_sequences = []  # list of utterance lists
    streak_buf = []
    for i in range(1, len(utterances)):
        u = utterances[i]; pv = utterances[i-1]
        if not is_ca(u['speaker']):
            # 求職者の回答：streak中ならシーケンスに含める（リセットしない）
            if cur_streak > 0:
                streak_buf.append(u)
            continue
        if DRILL_PAT.search(u['text']) and not is_ca(pv['speaker']) and len(u['text']) > 8:
            if cur_streak == 0: streak_buf = [pv, u]
            else:               streak_buf.append(u)
            cur_streak += 1
            max_streak = max(max_streak, cur_streak)
        else:
            # CAの非深掘り発話（話題転換・事実質問など）でstreak終了
            if cur_streak >= 2:
                vertical_drill_sequences.append(list(streak_buf[:8]))
            cur_streak = 0; streak_buf = []
    # 最後のstreakも追加
    if cur_streak >= 2:
        vertical_drill_sequences.append(list(streak_buf[:8]))

    # ── 自己開示の場面を収集 ─────────────────────────────
    self_disclosure_moments = []
    for i, u in enumerate(utterances):
        if not is_ca(u['speaker']): continue
        if SELF_DISC_PAT.search(u['text']) and len(u['text']) > 10:
            # 前後の文脈
            prev_cd = ''
            for j in range(i-1, max(0, i-4), -1):
                if is_cd(utterances[j]['speaker']):
                    prev_cd = utterances[j]['text'][:60]; break
            self_disclosure_moments.append({
                'ca_text':    u['text'][:100],
                'trigger':    prev_cd,
                'idx':        i,
            })

    # ── バックトラッキングの場面を収集 ──────────────────
    backtrack_moments = []
    for i, u in enumerate(utterances):
        if not is_ca(u['speaker']): continue
        if BT_PAT.search(u['text']) and len(u['text']) > 10:
            # 参照元の候補を直前のCD発話から探す
            referenced = ''
            for j in range(i-1, max(0, i-6), -1):
                if is_cd(utterances[j]['speaker']):
                    referenced = utterances[j]['text'][:60]; break
            backtrack_moments.append({
                'ca_text':    u['text'][:100],
                'referenced': referenced,
                'idx':        i,
            })

    # ── クロージング発話 ──────────────────────────────────
    closing_start  = int(len(utterances) * 0.8)
    closing_ca_utts = [u['text'] for u in utterances[closing_start:]
                       if is_ca(u['speaker']) and len(u['text']) > 10]
    closing_cd_utts = [u['text'] for u in utterances[closing_start:]
                       if is_cd(u['speaker']) and len(u['text']) > 10]

    return {
        '求職者発話比率':     round(len(all_cd)/total, 3) if total > 0 else 0,
        '後半求職者比率':     round(la_cd/la_tot, 3) if la_tot > 0 else 0,
        '会話回数':           turns,
        '名前呼称回数':       name_cnt,
        'フィラー回数':       len(FILLER_PAT.findall(all_ca)),
        '深掘り_価値観':      len(DRILL_PAT.findall(all_ca)),
        'バックトラッキング': sum(1 for t in ca_texts if BT_PAT.search(t)),
        '共感自己開示':       sum(1 for t in ca_texts if EMPATHY_PAT.search(t)),
        'ポジティブ反応':     sum(1 for t in ca_texts if POS_PAT.search(t)),
        'NG_急かし':         len(RUSH_PAT.findall(all_ca)),
        'NG_権威':           len(AUTH_PAT.findall(all_ca)),
        '感情場面数':         emotion_total,
        '感情スルー数':       emotion_skip,
        '感情スルー率':       slip_rate,
        '認識統一発話数':     sum(1 for t in ca_texts if SUMMARY_PAT.search(t)),
        '次回接続発話数':     sum(1 for t in ca_texts[-max(1,len(ca_texts)//4):] if NEXT_PAT.search(t)),
        'クロージング問いかけ': sum(1 for t in ca_texts if CLOSE_Q_PAT.search(t)),
        'MUST提案発話数':     sum(1 for t in ca_texts if MUST_PAT.search(t)),
        '縦深掘り最大':       max_streak,
        '自己開示回数':       len(self_disclosure_moments),
        # 具体例（表示用）
        '感情モーメント全件':     emotion_moments_all,
        '感情スルー例':           emotion_slip_examples,
        '感情深掘り例':           emotion_drill_examples,
        '縦深掘りシーケンス':     vertical_drill_sequences,
        '自己開示モーメント':     self_disclosure_moments,
        'バックトラッキングモーメント': backtrack_moments,
        'クロージングCA':         closing_ca_utts[-3:],
        'クロージングCD':         closing_cd_utts[-2:],
        # 内部用
        'ca_texts':   ca_texts,
        'cd_texts':   cd_texts,
        'utterances': utterances,
        'is_ca':      is_ca,
        'is_cd':      is_cd,
    }

# ── Claude採点 ── Call 1: スコアリング（実績ある構造を維持）────
def score_with_claude(utterances, ca_name, cand_name, fmt, client):
    transcript = '\n'.join(f"[{u['speaker']}] {u['text']}" for u in utterances)[:MAX_TRANSCRIPT_CHARS]

    prompt = f"""あなたは人材紹介会社の敏腕トレーナーです。以下の面談文字起こしを分析し、
CAへの具体的なフィードバックをJSONで生成してください。

## ★最重要：この会社の初回面談の目的（評価の大前提）
この人材紹介会社では、初回面談のゴールは「応募企業への魅力づけ・囲い込み」ではありません。
応募企業はあくまで求職者に来てもらうための入口であり、面談の本当の目的は次の2つです。
1. 求職者の価値観・本音を引き出し、意向を広げる（意向変え）
2. 応募企業に固執させず、複数企業・新しい選択肢に目を向けてもらう（視野拡大）
したがって全軸を以下の方針で評価してください。
- 応募企業への魅力づけ・志望度の確認・囲い込みは評価対象外。やっていなくても減点しない。
- 応募企業だけに話を限定している場合はマイナス、視野を広げられている場合はプラス。
- 求職者の価値観を引き出し、新しい視点・他の可能性に気づかせられたかを高く評価する。

## 各評価軸の定義（この定義に厳密に従うこと）
- 意向把握：価値観・やりがいを引き出し、応募企業に固執させず意向を広げられたか（意向変えの土台）
- 適正把握：経験・強み・適性を具体的に把握できたか
- 条件把握：Must/Betterを確認し、期待値を調整できたか
- 認識統一：①求職者の価値観・強みの要約に本人の明示的同意が得られたか、②今後のキャリアの方向性（複数社を見て進めること等）について合意できたか。※応募企業への同意ではない
- 気づき付与：求職者に新しい視点・他の選択肢の可能性に気づかせられたか

## スコア基準（0〜3点）
- 3: 根拠引用あり＋求職者の明示的な同意・反応が確認できる
- 2: 把握できているが確認が弱い or 求職者が「はい」止まり
- 1: 触れているが浅い・一方的な説明のみ
- 0: 未実施

## score と score_fine（必ず両方出すこと）
- score：上記の0〜3の整数（グレード判定用）
- score_fine：同じ評価を0.0〜3.0の小数で表した値。バンド内のニュアンスを反映する
  （例：かろうじて2なら2.1、3に近い2なら2.8、ちょうど真ん中なら2.5）。
  scoreの整数部と必ず整合させること（score=2ならscore_fineは1.5〜2.4の範囲）。

## グリップA基準値
- ポジティブ反応：5回以上 / 価値観深掘り：4回以上
- バックトラッキング：5回以上 / 感情スルー率：50%以下 / 縦深掘り：同テーマ3回以上連続

## メタ情報
CA名: {ca_name} / 求職者名: {cand_name} / 形式: {fmt}

## 面談文字起こし
{transcript}

## 出力（JSONのみ・余計な文章不要）
{{
  "grip_drivers": {{
    "意向":    {{"score":0,"score_fine":0.0,"evidence":[],"strength":"","weakness":"","next_action":""}},
    "適正":    {{"score":0,"score_fine":0.0,"evidence":[],"strength":"","weakness":"","next_action":""}},
    "条件":    {{"score":0,"score_fine":0.0,"evidence":[],"strength":"","weakness":"","next_action":""}},
    "認識統一":{{"score":0,"score_fine":0.0,"evidence":[],"strength":"","weakness":"","next_action":""}},
    "気づき":  {{"score":0,"score_fine":0.0,"evidence":[],"strength":"","weakness":"","next_action":""}}
  }},
  "behaviors": {{
    "深掘り_価値観":0,"深掘り_実績":0,
    "バックトラッキング":0,"共感自己開示":0,"ポジティブ反応":0,
    "NG_急かし":0,"NG_権威":0,"NG_感情無視":0,
    "フェーズ網羅":["冒頭"],
    "MUST提案":false,"MUST同意":false,"次回アポ確定":false
  }},
  "overall": {{
    "grade":"A",
    "grade_reason":"グレード判定の理由を1文で",
    "top_strength":"この面談の最大の強みを具体的に1文で",
    "top_issues":[
      {{"issue":"課題1のタイトル","detail":"何が問題だったか","fix":"代わりにこう言うべきだった（例文付き）"}},
      {{"issue":"課題2のタイトル","detail":"","fix":""}},
      {{"issue":"課題3のタイトル","detail":"","fix":""}}
    ],
    "missed_moment":"感情ワードをスルーした最も惜しかった場面（求職者発話+CAの返しを引用）",
    "best_exchange":"最も良かったやり取り（求職者+CA各1発話を引用）",
    "closing_eval":"クロージング評価：求職者が主役で終われたか・次回接続できたか",
    "one_thing":"次の面談で必ず1つだけ試してほしいこと（例文付き）"
  }},
  "notes":"文字起こし品質・話者分離の懸念など（なければ空文字）"
}}

grade基準: S=全軸2.5以上+感情深掘り◎, A=総合スコア10以上, B=7〜9, C=4〜6, D=3以下
MUST提案: エンジニア意向弱→別職種提案、強→エンジニア路線確認、どちらかできていればtrue"""

    try:
        resp = client.messages.create(
            model='claude-sonnet-4-6', max_tokens=12000,
            messages=[{'role': 'user', 'content': prompt}])
        raw = resp.content[0].text
        result = safe_json_loads(raw)
        if not result.get('overall'):
            st.expander('🔍 デバッグ：Claudeの生レスポンス').code(raw[:3000])
        return result
    except Exception as e:
        st.error(f'スコアリングAPIエラー: {type(e).__name__}: {e}')
        return {}


# ── Claude採点 ── Call 2: 深掘り詳細分析（新規・別呼び出し）──
def deep_analysis_with_claude(utterances, ca_name, cand_name, client):
    transcript = '\n'.join(f"[{u['speaker']}] {u['text']}" for u in utterances)[:MAX_TRANSCRIPT_CHARS]

    prompt = f"""あなたは人材紹介会社の面談コーチです。以下の面談文字起こしを分析し、
3つの観点から詳細なフィードバックをJSONで返してください。

## メタ情報
CA名: {ca_name} / 求職者名: {cand_name}

## 面談文字起こし
{transcript}

## 出力（JSONのみ・余計な文章不要）
{{
  "emotion_drill_analysis": {{
    "summary": "感情深掘り全体の評価（2〜3文）",
    "missed_scenes": [
      {{
        "cd_text": "スルーされた求職者の発話（30字程度の引用）",
        "emotion_word": "スルーされた感情ワード",
        "ca_actual": "実際のCAの返し（引用）",
        "ca_suggested": "こう返すべきだった（具体的な一言・例文）",
        "why": "なぜこの返しが良いか（1文）"
      }}
    ],
    "good_scenes": [
      {{
        "cd_text": "うまく深掘りできた場面の求職者発話（引用）",
        "ca_text": "CAの深掘り発話（引用）",
        "why_good": "なぜ良かったか（1文）"
      }}
    ],
    "vertical_drill_comment": "縦の深掘り（同テーマ連続）の評価（1〜2文）"
  }},
  "self_disclosure_analysis": {{
    "summary": "自己開示全体の評価（1〜2文）",
    "found_scenes": [
      {{
        "ca_text": "CAの自己開示発話（引用）",
        "timing_eval": "良い／普通／改善余地あり",
        "effect": "求職者の反応・効果（1文）"
      }}
    ],
    "missed_opportunities": [
      {{
        "cd_text": "ここで自己開示できた求職者発話（引用）",
        "ca_suggested": "こう自己開示できた（例文）"
      }}
    ],
    "advice": "自己開示改善アドバイス（1〜2文）"
  }},
  "backtrack_analysis": {{
    "summary": "バックトラッキング全体の評価（1〜2文）",
    "found_scenes": [
      {{
        "ca_text": "CAのバックトラッキング発話（引用）",
        "referenced_cd": "参照した求職者発言（引用）",
        "effect": "良い／普通／惜しい"
      }}
    ],
    "missed_opportunities": [
      {{
        "cd_keyword": "後で引用できた求職者のキーワード・発言",
        "ca_suggested": "こう引用して使えた（例文）"
      }}
    ],
    "advice": "バックトラッキング改善アドバイス（1〜2文）"
  }},
  "next_phrases": [
    {{"situation": "感情ワードが出た直後", "phrase": "「〜」（この求職者のキーワードを使った具体例）", "why": "なぜこのフレーズが効くか"}},
    {{"situation": "縦の深掘り2回目", "phrase": "「〜」", "why": ""}},
    {{"situation": "自己開示を入れるタイミング", "phrase": "「私も〜（具体的な自己開示例）〜さんはどうですか？」", "why": ""}},
    {{"situation": "バックトラッキングで引用", "phrase": "「先ほど〜とおっしゃっていましたが〜」", "why": ""}},
    {{"situation": "クロージング問いかけ", "phrase": "「今日話してみて、〜さんはどう感じましたか？」", "why": ""}}
  ]
}}"""

    try:
        resp = client.messages.create(
            model='claude-sonnet-4-6', max_tokens=12000,
            messages=[{'role': 'user', 'content': prompt}])
        result = safe_json_loads(resp.content[0].text)
        return result if result else {}
    except Exception as e:
        raise e  # 呼び出し元でキャッチしてユーザーに表示

# ── チェックポイント評価 ──────────────────────────────────
def evaluate_checklist(metrics, claude_result):
    bh = claude_result.get('behaviors', {})
    gd = claude_result.get('grip_drivers', {})
    def score(key): return gd.get(key, {}).get('score', 0)

    slip_rate   = metrics['感情スルー率']
    max_drill   = metrics['縦深掘り最大']
    will_drill  = metrics['深掘り_価値観']
    backtrack   = metrics['バックトラッキング']
    must_cnt    = metrics['MUST提案発話数']
    summary_cnt = metrics['認識統一発話数']
    close_q     = metrics['クロージング問いかけ']
    latter_r    = metrics['後半求職者比率']
    next_cnt    = metrics['次回接続発話数']
    cand_r      = metrics['求職者発話比率']
    pos_cnt     = metrics['ポジティブ反応']
    filler_cnt  = metrics['フィラー回数']
    name_cnt    = metrics['名前呼称回数']
    rush_cnt    = metrics['NG_急かし']
    slip_cnt    = metrics['感情スルー数']
    sd_cnt      = metrics['自己開示回数']

    checks = {
        '事前準備': [
            ('ゴール（次回接続）をイメージして臨んだか', '', True, True),
            ('強み仮説を1つ立てて臨んだか（★グリップに影響）', '', True, True),
        ],
        '冒頭': [
            ('カジュアル面談の空気を作れたか', '', True, True),
            ('自己開示（自分の話）を入れたか ★',
             f'検出: {sd_cnt}回（目標1回以上）',
             sd_cnt >= 1, sd_cnt >= 1),
            ('求職者の名前を呼んだか',
             f'{name_cnt}回（目標3回以上）',
             name_cnt >= 3, name_cnt >= 1),
        ],
        'CAN把握': [
            ('エピソードを引き出せたか（事実だけでない）', '', score('適正') >= 2, score('適正') >= 1),
            ('強みをCAが言語化して確認したか', '', score('適正') >= 2, score('適正') >= 1),
        ],
        'WILL把握 ★最重要': [
            ('感情ワードが出たら即深掘りしたか ★',
             f'感情スルー率: {slip_rate}%（目標50%以下）',
             slip_rate <= 50, slip_rate <= 70),
            ('同じテーマを3回以上掘り下げたか（縦深掘り）★',
             f'最大連続: {max_drill}回（目標3回以上）',
             max_drill >= 3, max_drill >= 2),
            ('感情系・最上級系の質問を使ったか',
             f'価値観深掘り: {will_drill}回（目標4回以上）',
             will_drill >= 4, will_drill >= 2),
            ('前の話を引用してクッション質問でつないだか（バックトラッキング）',
             f'バックトラッキング: {backtrack}回（目標5回以上）',
             backtrack >= 5, backtrack >= 3),
            ('求職者の価値観を言語化できたか', '', score('意向') >= 2, score('意向') >= 1),
            ('沈黙を恐れず、求職者が続きを話せるのを待てたか', '', True, True),
        ],
        'MUST把握': [
            ('3条件をMust/Betterで確認したか', '', score('条件') >= 2, score('条件') >= 1),
            ('期待値調整後に求職者の反応を確認したか', '', score('条件') >= 2, score('条件') >= 1),
        ],
        '気づき付与': [
            ('CAの言葉で「○○さんには〜が合う」を伝えたか',
             f'MUST提案発話: {must_cnt}回',
             bh.get('MUST提案', False), must_cnt >= 1),
            ('提案後に「どう思いますか？」と同意を確認したか',
             '', bh.get('MUST同意', False), bh.get('MUST提案', False)),
        ],
        '認識統一': [
            ('CAN・WILL・MUSTをCAが要約して確認したか',
             f'まとめ発話: {summary_cnt}回',
             summary_cnt >= 1, summary_cnt >= 1),
            ('求職者が言葉で同意したか', '', score('認識統一') >= 2, score('認識統一') >= 1),
        ],
        'クロージング ★': [
            ('問いかけで求職者に話させたか',
             f'問いかけ: {close_q}回', close_q >= 1, close_q >= 1),
            ('後半も求職者が主役だったか',
             f'後半求職者比率: {latter_r:.0%}（目標50%以上）',
             latter_r >= 0.5, latter_r >= 0.35),
            ('アポを具体化して次回接続を作れたか（★最大の改善余地）',
             f'次回接続発話: {next_cnt}回',
             bh.get('次回アポ確定', False) or next_cnt >= 1, next_cnt >= 1),
        ],
        '会話全体': [
            ('求職者発話比率40%以上か',
             f'{cand_r:.0%}（目標40%以上）', cand_r >= 0.4, cand_r >= 0.3),
            ('ポジティブ反応5回以上か',
             f'{pos_cnt}回（目標5回以上）', pos_cnt >= 5, pos_cnt >= 3),
            ('フィラー30回以下か',
             f'{filler_cnt}回（目標30回以下）', filler_cnt <= 30, filler_cnt <= 60),
            ('NGパターンがないか（急かし・感情無視）',
             f'急かし{rush_cnt}回 / 感情スルー{slip_cnt}回',
             rush_cnt == 0 and slip_rate <= 50, rush_cnt <= 1),
        ],
    }
    return checks

# ─────────────────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────────────────
st.title('💬 初回面談 自動FB ツール')
st.caption('文字起こしをアップロードするだけで、採点・深掘りマップ・改善フレーズを自動生成')

api_key     = os.environ.get('ANTHROPIC_API_KEY', '')
SLACK_WEBHOOK = st.secrets.get('SLACK_WEBHOOK_URL', os.environ.get('SLACK_WEBHOOK_URL', ''))

with st.sidebar:
    st.header('⚙️ 設定')
    if not api_key:
        st.error('⚠️ APIキー未設定\n`source ~/.zshrc` 後に再起動')
    else:
        st.success('✅ APIキー設定済み')
    st.divider()
    st.subheader('面談情報')
    ca_input   = st.text_input('CA名', placeholder='例：下川')
    cand_input = st.text_input('求職者名', placeholder='例：岡寺さん')
    grip_input = st.selectbox('グリップランク（任意）', ['未入力','A','B','C','D'])
    st.divider()
    st.caption('💡 「CA名_グリップ_求職者名_面談種別」形式で命名すると自動入力されます')

uploaded = st.file_uploader('ファイルをアップロード（.docx または .txt）', type=['docx','txt'])


if uploaded:
    fn_nfc = unicodedata.normalize('NFC', uploaded.name)
    stem   = fn_nfc.rsplit('.',1)[0]
    parts  = re.split(r'[_＿]', re.sub(r'-グリップ', '_グリップ', stem))
    for i, p in enumerate(parts):
        m = re.search(r'グリップ([ABCD])', p)
        if m:
            if not ca_input:   ca_input   = parts[0]
            if not cand_input and i+1 < len(parts): cand_input = parts[i+1]
            if grip_input == '未入力': grip_input = m.group(1)
            break

    col1, col2, col3 = st.columns(3)
    col1.metric('CA名', ca_input or '未入力')
    col2.metric('求職者名', cand_input or '未入力')
    col3.metric('グリップ', grip_input)

    # ── 途中失敗時の再実行ボタン ──────────────────────────
    if '_app_retry' in st.session_state:
        _retry = st.session_state['_app_retry']
        client_retry = anthropic.Anthropic(api_key=api_key) if api_key else None

        if 'claude_result' in _retry:
            # スコアリングは完了 → 深掘りだけ再実行
            st.warning('⚠️ 深掘り分析が未完了です。下のボタンで再実行できます。')
            if st.button('🔄 深掘り分析だけ再実行', type='primary', use_container_width=True):
                with st.spinner('🔍 深掘り分析を再実行中...'):
                    try:
                        deep_result = deep_analysis_with_claude(
                            _retry['utterances'], _retry['ca'], _retry['candidate'], client_retry)
                        deep_result = deep_result or {}
                    except Exception as e:
                        st.error(f'再実行も失敗しました：{e}')
                        st.stop()
                meeting_type = '初回面談' if '初回面談' in _retry['fn_nfc'] else ('求人提案' if '求人提案' in _retry['fn_nfc'] else 'その他')
                utt_path, json_path = save_results(
                    _retry['ca'], _retry['grip'], _retry['candidate'], meeting_type, _retry['fmt'],
                    _retry['utterances'], _retry['metrics'], _retry['claude_result'], deep_result)
                st.session_state['_app_res'] = {
                    'claude_result': _retry['claude_result'],
                    'deep_result':   deep_result,
                    'metrics':       _retry['metrics'],
                    'ca':            _retry['ca'],
                    'grip':          _retry['grip'],
                    'candidate':     _retry['candidate'],
                    'meeting_type':  meeting_type,
                    'ref_file':      json_path.name,
                }
                st.session_state.pop('_app_retry', None)
                st.session_state.pop('_app_sc_done', None)
                st.session_state.pop('_app_sc_skip', None)
                st.rerun()
        else:
            # スコアリングから再実行
            st.warning('⚠️ スコアリングが未完了です。下のボタンで再実行できます。')
            if st.button('🔄 スコアリングから再実行', type='primary', use_container_width=True):
                with st.spinner('🤖 スコアリングを再実行中...'):
                    claude_result = score_with_claude(
                        _retry['utterances'], _retry['ca'], _retry['candidate'], _retry['fmt'], client_retry)
                if not claude_result.get('overall'):
                    st.error('再実行も失敗しました。しばらく待ってから再試行してください。')
                    st.stop()
                with st.spinner('🔍 深掘り分析中...'):
                    try:
                        deep_result = deep_analysis_with_claude(
                            _retry['utterances'], _retry['ca'], _retry['candidate'], client_retry)
                        deep_result = deep_result or {}
                    except Exception:
                        deep_result = {}
                meeting_type = '初回面談' if '初回面談' in _retry['fn_nfc'] else ('求人提案' if '求人提案' in _retry['fn_nfc'] else 'その他')
                utt_path, json_path = save_results(
                    _retry['ca'], _retry['grip'], _retry['candidate'], meeting_type, _retry['fmt'],
                    _retry['utterances'], _retry['metrics'], claude_result, deep_result)
                st.session_state['_app_res'] = {
                    'claude_result': claude_result,
                    'deep_result':   deep_result,
                    'metrics':       _retry['metrics'],
                    'ca':            _retry['ca'],
                    'grip':          _retry['grip'],
                    'candidate':     _retry['candidate'],
                    'meeting_type':  meeting_type,
                    'ref_file':      json_path.name,
                }
                st.session_state.pop('_app_retry', None)
                st.session_state.pop('_app_sc_done', None)
                st.session_state.pop('_app_sc_skip', None)
                st.rerun()

        st.divider()

    if st.button('🔍 分析開始', type='primary', use_container_width=True):
        if not api_key:
            st.error('APIキーが設定されていません。'); st.stop()
        if not ca_input or not cand_input:
            st.error('CA名と求職者名を入力してください'); st.stop()

        client = anthropic.Anthropic(api_key=api_key)

        with st.spinner('📄 ファイルを解析中...'):
            file_bytes = uploaded.read()
            ext = fn_nfc.rsplit('.',1)[1].lower()
            if ext == 'docx':
                utterances, raw_text, fmt = parse_docx(file_bytes)
            else:
                utterances, raw_text, fmt = parse_txt_with_haiku(
                    file_bytes, ca_input, cand_input, client)

        if not utterances or (len(utterances)==1 and utterances[0]['speaker']=='不明'):
            st.error('話者分離に失敗しました。CA名・求職者名を確認してください。'); st.stop()

        speakers_detected = list(set(u['speaker'] for u in utterances))
        st.success(f'✅ 話者分離完了：{len(utterances)}発話 / 検出話者：{speakers_detected}')

        with st.spinner('📊 行動指標を計算中...'):
            metrics = compute_metrics(utterances, ca_input, cand_input)

        # ── スコアリング（途中失敗時は再実行ボタンを表示） ──
        with st.spinner('🤖 AIが採点中（スコアリング）...'):
            claude_result = score_with_claude(utterances, ca_input, cand_input, fmt, client)

        if not claude_result.get('overall'):
            st.error('⚠️ スコアリングに失敗しました。')
            # 話者分離・行動指標はsession_stateに保存して再実行に備える
            st.session_state['_app_retry'] = {
                'utterances': utterances, 'metrics': metrics,
                'ca': ca_input, 'grip': grip_input,
                'candidate': cand_input, 'fn_nfc': fn_nfc, 'fmt': fmt,
            }
            st.stop()

        # ── 深掘り分析（失敗時は再実行ボタンを表示） ──
        with st.spinner('🔍 深掘り・自己開示・バックトラッキングを分析中...'):
            try:
                deep_result = deep_analysis_with_claude(utterances, ca_input, cand_input, client)
                if not deep_result:
                    deep_result = None
            except Exception:
                deep_result = None

        if deep_result is None:
            st.warning('⚠️ 深掘り分析に失敗しました。スコアリングは完了しています。')
            # スコアリング結果は保持して深掘りだけ再実行できるよう保存
            st.session_state['_app_retry'] = {
                'utterances': utterances, 'metrics': metrics,
                'ca': ca_input, 'grip': grip_input,
                'candidate': cand_input, 'fn_nfc': fn_nfc, 'fmt': fmt,
                'claude_result': claude_result,  # スコアリング結果を保持
            }
            deep_result = {}

        meeting_type = '初回面談' if '初回面談' in fn_nfc else ('求人提案' if '求人提案' in fn_nfc else 'その他')
        utt_path, json_path = save_results(
            ca_input, grip_input, cand_input, meeting_type, fmt,
            utterances, metrics, claude_result, deep_result)
        st.success(f'💾 保存完了　📊 分析結果：`output/json/{json_path.name}`　💬 話者分離：`output/utterances/{utt_path.name}`')
        # ダッシュボードのサマリーキャッシュを無効化（次回ダッシュボード表示時に自動再構築される）
        try:
            from gdrive import upload_json as _gup
            # サマリーファイルを削除して強制再構築させる
            from gdrive import list_json_files as _ljf, download_json as _dlj
            _summary_files = [f for f in _ljf(subfolder='json') if f['name'] == '_dashboard_summary.json']
            if not _summary_files:
                pass  # サマリーがなければ何もしない
        except Exception:
            pass  # 失敗しても問題なし（次回ダッシュボードのリロードで自動更新）

        # ── 分析結果をsession_stateに保存して表示フローに渡す ──
        st.session_state['_app_res'] = {
            'claude_result': claude_result,
            'deep_result':   deep_result,
            'metrics':       metrics,
            'ca':            ca_input,
            'grip':          grip_input,
            'candidate':     cand_input,
            'meeting_type':  meeting_type,
            'ref_file':      json_path.name,
        }
        st.session_state.pop('_app_retry', None)
        # selfcheckフラグをリセット（新規分析のため）
        st.session_state.pop('_app_sc_done', None)
        st.session_state.pop('_app_sc_skip', None)
        st.rerun()   # ← 表示フローへ

# ═══════════════════════════════════════════════════════════
# 結果表示ブロック（if button: の外 → rerun後も確実に実行される）
# ═══════════════════════════════════════════════════════════
if '_app_res' in st.session_state:
    if True:  # インデントをネスト（表示コードは8スペース）
        _r          = st.session_state['_app_res']
        claude_result = _r['claude_result']
        deep_result   = _r['deep_result']
        metrics       = _r['metrics']
        ca_input      = _r['ca']
        grip_input    = _r['grip']
        cand_input    = _r['candidate']
        meeting_type  = _r['meeting_type']
        ref_file_app  = _r['ref_file']
        json_path     = type('_JP', (), {'name': ref_file_app})()

        gd      = claude_result.get('grip_drivers', {})
        overall = claude_result.get('overall', {})
        total_score = sum(gd.get(k,{}).get('score',0) for k in ['意向','適正','条件','認識統一','気づき'])

        # ══ セルフチェックゲート（AI評価を見る前に自己採点） ══
        sc_existing  = load_selfcheck_app(ref_file_app)
        show_ai_app  = sc_existing is not None \
                       or st.session_state.get('_app_sc_done') \
                       or st.session_state.get('_app_sc_skip')

        if not show_ai_app:
            st.info('🪞 **まず自己採点を。** AI評価を見る前に自分の面談を5軸で採点してください。'
                    '「自分の感覚」と「AIの客観評価」のズレが一番の伸びしろになります。')
            with st.form(key='app_selfcheck_form'):
                with st.expander('📖 採点の目安（0〜3点の意味）', expanded=False):
                    for line in SCORE_LADDER_APP:
                        st.markdown(f'- {line}')

                st.markdown('**① この面談、自分では何点だった？（各0〜3点・0.5刻み）**')
                sc_scores_app = {}
                for ax in AXES_APP:
                    st.markdown(f'**{AXIS_SHORT_APP[ax]}**　'
                                f'<small style="color:#888">{AXIS_DEF_APP[ax]}</small>',
                                unsafe_allow_html=True)
                    sc_scores_app[ax] = st.slider(
                        AXIS_SHORT_APP[ax], min_value=0.0, max_value=3.0, value=2.0, step=0.5,
                        key=f'app_sc_{ax}', label_visibility='collapsed')

                st.markdown('**② できたと思う行動にチェック**')
                bc1, bc2 = st.columns(2)
                bchecks_app = {
                    '感情ワードを拾って深掘りした':           bc1.checkbox('感情ワードを拾って深掘りした', key='app_bc1'),
                    '強みを言語化して返した':                 bc2.checkbox('強みを言語化して返した', key='app_bc2'),
                    '応募企業に固執させず他の選択肢に触れた': bc1.checkbox('応募企業に固執させず他の選択肢に触れた', key='app_bc3'),
                    '求職者の発言を要約して同意を取った':     bc2.checkbox('求職者の発言を要約して同意を取った', key='app_bc4'),
                    '沈黙を恐れず考える間を与えた':           bc1.checkbox('沈黙を恐れず考える間を与えた', key='app_bc5'),
                    'MUST提案をした':                         bc2.checkbox('MUST提案をした', key='app_bc6'),
                    '次回アポを確定した':                     bc1.checkbox('次回アポを確定した', key='app_bc7'),
                }

                st.markdown('**③ 今日の面談の振り返り**')
                best_self_app = st.text_input('自分で「ここは良かった」と思う点', key='app_best',
                                              placeholder='例：価値観をしっかり引き出せた')
                next_one_app  = st.text_area('次の面談で試したいこと（1つ）', key='app_no',
                                             placeholder='例：感情ワードが出たら必ず「それってどんな気持ちでしたか？」と返す')

                col_save, col_skip = st.columns([2, 1])
                do_save_app = col_save.form_submit_button('✅ 採点を保存してAI評価を見る',
                                                          use_container_width=True, type='primary')
                do_skip_app = col_skip.form_submit_button('⏭ スキップ', use_container_width=True)

            if do_save_app:
                # session_stateからスライダー値を取得して保存
                scores_to_save = {ax: st.session_state.get(f'app_sc_{ax}', 2.0) for ax in AXES_APP}
                save_selfcheck_app(ref_file_app, ca_input, grip_input, cand_input,
                                   meeting_type, scores_to_save, bchecks_app,
                                   next_one_app, best_self_app)
                st.session_state['_app_sc_done'] = True
                st.rerun()   # ← rerunしてもOK。表示ブロックがsession_stateから復元する
            if do_skip_app:
                st.session_state['_app_sc_skip'] = True
                st.rerun()
            st.stop()

        # ── 既存の自己採点があればズレFBを表示 ──
        sc_app = sc_existing or {}
        if sc_app and not st.session_state.get('_app_sc_skip'):
            ss_app = sc_app.get('self_scores', {})

            def _fine_app(info):
                v = info.get('score_fine')
                if v is None: v = info.get('score', 0)
                try: return float(v)
                except: return float(info.get('score', 0) or 0)
            g_app = lambda v: f'{v:g}'

            gap_rows_html = ('<table style="width:100%;border-collapse:collapse;font-size:0.9rem">'
                             '<tr style="background:#1F3864;color:white">'
                             '<th style="padding:6px;text-align:left">評価軸</th>'
                             '<th style="padding:6px">自己</th>'
                             '<th style="padding:6px">AI</th>'
                             '<th style="padding:6px;text-align:left">ズレ</th></tr>')
            blind_a, under_a, cweak_a, cstrong_a = [], [], [], []
            for ax in AXES_APP:
                ai_s   = _fine_app(gd.get(ax, {}))
                self_s = float(ss_app.get(ax, 0) or 0)
                diff   = round(self_s - ai_s, 1)
                if diff >= 0.5:
                    tag = f'<span style="color:#c0392b">+{g_app(diff)} 自分が高め</span>'; bg = '#fcecea'
                    blind_a.append((ax, self_s, ai_s, diff))
                elif diff <= -0.5:
                    tag = f'<span style="color:#2471a3">{g_app(diff)} 自分が低め</span>'; bg = '#EBF5FB'
                    under_a.append((ax, self_s, ai_s, diff))
                else:
                    tag = f'<span style="color:#1e8449">±{g_app(abs(diff))} ほぼ一致</span>'; bg = '#e2efda'
                    if ai_s >= 2: cstrong_a.append((ax, self_s, ai_s))
                    else:         cweak_a.append((ax, self_s, ai_s))
                gap_rows_html += (f'<tr style="background:{bg}"><td style="padding:6px">'
                                  f'{AXIS_SHORT_APP[ax]}</td>'
                                  f'<td style="padding:6px;text-align:center">{g_app(self_s)}</td>'
                                  f'<td style="padding:6px;text-align:center">{g_app(ai_s)}</td>'
                                  f'<td style="padding:6px">{tag}</td></tr>')
            gap_rows_html += '</table>'
            blind_a.sort(key=lambda x: -x[3])

            with st.expander('🪞 自己評価とAIのズレ → あなた専用FB', expanded=True):
                st.markdown(gap_rows_html, unsafe_allow_html=True)
                if blind_a:
                    b = blind_a[0]
                    st.error(f'🎯 **今日の最優先：{AXIS_SHORT_APP[b[0]]}** — 最大の盲点'
                             f'（自分{b[1]:g}点 / AI{b[2]:g}点）。下のAI評価の根拠を確認してください。')
                elif cweak_a:
                    w = sorted(cweak_a, key=lambda x: x[2])[0]
                    st.warning(f'🎯 **今日の最優先：{AXIS_SHORT_APP[w[0]]}** — 自他ともに課題と認識している軸です。')
                else:
                    st.success('🎉 自己評価とAI評価がほぼ一致。自分の面談を客観視できています。')

                for ax, s, a, _ in blind_a:
                    info = gd.get(ax, {}); ev = info.get('evidence') or []
                    h = (f'<div class="emotion-miss"><b>🔴 盲点：{AXIS_SHORT_APP[ax]}</b>'
                         f'（自分{s:g}点 / AI{a:g}点）<br>'
                         f'<small>「できた」と感じたが、AIは弱点と評価＝気づけていない伸びしろ。</small>')
                    if info.get('weakness'):    h += f'<br>📌 <b>AIが見た弱み：</b>{info["weakness"]}'
                    if ev:                      h += f'<br><small>根拠：{ev[0]}</small>'
                    if info.get('next_action'): h += f'<br>🚀 <b>次の一手：</b>{info["next_action"]}'
                    st.markdown(h + '</div>', unsafe_allow_html=True)

                for ax, s, a in sorted(cweak_a, key=lambda x: x[2]):
                    info = gd.get(ax, {})
                    h = (f'<div style="background:#FEF9E7;border-left:4px solid #F39C12;'
                         f'padding:8px 12px;border-radius:4px;margin:6px 0">'
                         f'<b>🟠 共通課題：{AXIS_SHORT_APP[ax]}</b>（自分{s:g}点 / AI{a:g}点）<br>'
                         f'<small>自覚あり。あとは行動に移すだけ。</small>')
                    if info.get('next_action'): h += f'<br>🚀 {info["next_action"]}'
                    st.markdown(h + '</div>', unsafe_allow_html=True)

                for ax, s, a, _ in under_a:
                    info = gd.get(ax, {})
                    h = (f'<div class="bt-hit"><b>🔵 過小評価：{AXIS_SHORT_APP[ax]}</b>'
                         f'（自分{s:g}点 / AI{a:g}点）<br>'
                         f'<small>実はできています。自信を持って再現を。</small>')
                    if info.get('strength'): h += f'<br>💪 {info["strength"]}'
                    st.markdown(h + '</div>', unsafe_allow_html=True)

                if cstrong_a:
                    names = '、'.join(AXIS_SHORT_APP[ax] for ax, _, _ in cstrong_a)
                    st.markdown(f'<div class="emotion-hit">🟢 <b>共通の強み：</b>{names}'
                                f' — この型を継続！</div>', unsafe_allow_html=True)

                if sc_app.get('best_self'):
                    st.markdown(f'**🙌 あなたが「良かった」と書いた点：** {sc_app["best_self"]}')
                if sc_app.get('next_one_thing'):
                    st.markdown(f'**🚀 あなたが「次に試す」と書いたこと：** {sc_app["next_one_thing"]}')
            st.divider()

        # ── グレード ────────────────────────────────────
        grade = overall.get('grade','─')
        grade_colors = {'S':'#1a5276','A':'#1e8449','B':'#2471a3','C':'#d35400','D':'#c0392b'}
        grade_color  = grade_colors.get(grade,'#555')
        st.markdown(
            f'<div style="background:{grade_color};color:white;padding:16px 20px;'
            f'border-radius:10px;margin-bottom:12px">'
            f'<span style="font-size:2.4rem;font-weight:bold">{grade}</span>'
            f'&nbsp;&nbsp;<span style="font-size:1rem">{overall.get("grade_reason","")}</span>'
            f'</div>', unsafe_allow_html=True)

        # ── KPI ─────────────────────────────────────────
        c1,c2,c3,c4,c5,c6 = st.columns(6)
        def badge(ok, warn=None):
            if ok: return '✅'
            if warn: return '⚠️'
            return '❌'

        c1.metric('総合スコア', f'{total_score}/15',
                  delta='良好' if total_score>=10 else ('要改善' if total_score<7 else '普通'))
        cr = metrics['求職者発話比率']
        c2.metric('求職者発話比率', f'{cr:.0%}',
                  delta=badge(cr>=0.4, cr>=0.3)+('目標達成' if cr>=0.4 else '目標40%'), delta_color='off')
        sl = metrics['感情スルー率']
        c3.metric('感情スルー率', f'{sl}%',
                  delta=('✅良好' if sl<=50 else ('⚠️要注意' if sl<=70 else '❌要改善')), delta_color='off')
        md = metrics['縦深掘り最大']
        c4.metric('縦深掘り最大', f'{md}回',
                  delta=badge(md>=3, md>=2)+('目標達成' if md>=3 else '目標3回'), delta_color='off')
        bt = metrics['バックトラッキング']
        c5.metric('バックトラッキング', f'{bt}回',
                  delta=badge(bt>=5, bt>=3)+('目標達成' if bt>=5 else '目標5回'), delta_color='off')
        sd = metrics['自己開示回数']
        c6.metric('自己開示', f'{sd}回',
                  delta=badge(sd>=2, sd>=1)+('目標達成' if sd>=2 else '目標2回以上'), delta_color='off')

        st.divider()

        # ── 強み & 改善TOP3 ──────────────────────────
        st.markdown('<div class="section-title">📋 面談サマリー</div>', unsafe_allow_html=True)
        col_s, col_i = st.columns([1,1])
        with col_s:
            st.markdown('**💪 この面談の最大の強み**')
            st.success(overall.get('top_strength','─'))
            best = overall.get('best_exchange','')
            if best:
                st.markdown('**👍 最も良かったやり取り**')
                st.info(best)
        with col_i:
            st.markdown('**🎯 優先改善ポイント（TOP3）**')
            for idx, issue in enumerate(overall.get('top_issues',[])[:3]):
                if isinstance(issue, dict):
                    with st.expander(f'**{idx+1}. {issue.get("issue","")}**', expanded=(idx==0)):
                        if issue.get('detail'): st.write('📌 ' + issue['detail'])
                        if issue.get('fix'):    st.code(issue['fix'], language=None)
                else:
                    st.error(f'{idx+1}. {issue}')

        one_thing = overall.get('one_thing','')
        if one_thing:
            st.markdown('**🚀 次の面談で必ず1つ試してほしいこと**')
            st.warning(one_thing)

        st.divider()

        # ════════════════════════════════════════════════
        # 感情深掘りマップ（★メインセクション）
        # ════════════════════════════════════════════════
        st.markdown('<div class="section-title">🔍 感情深掘りマップ ★ 最重要分析</div>',
                    unsafe_allow_html=True)

        eda       = deep_result.get('emotion_drill_analysis', {})
        all_moments = metrics.get('感情モーメント全件', [])
        em_total  = metrics['感情場面数']
        em_skip   = metrics['感情スルー数']
        em_drill  = em_total - em_skip

        # サマリー帯
        col_a, col_b, col_c = st.columns(3)
        col_a.metric('感情ワード出現', f'{em_total}場面')
        col_b.metric('深掘りできた', f'{em_drill}場面',
                     delta='✅' if em_drill >= em_total * 0.5 else '⚠️', delta_color='off')
        col_c.metric('スルーしてしまった', f'{em_skip}場面',
                     delta='❌ 要改善' if em_skip > 0 else '✅', delta_color='off')

        if eda.get('summary'):
            st.info('📊 ' + eda['summary'])

        # 全件マップ（展開可能）
        if all_moments:
            with st.expander(f'📋 感情ワード場面の一覧（全{em_total}件）', expanded=True):
                for rec in all_moments:
                    icon  = '✅' if rec['drilled'] else '❌'
                    label = '深掘りできた' if rec['drilled'] else 'スルーした'
                    cls   = 'emotion-hit' if rec['drilled'] else 'emotion-miss'
                    st.markdown(
                        f'<div class="{cls}">'
                        f'<small>{icon} <b>{label}</b>　感情ワード：<b>「{rec["emotion_word"]}」</b></small><br>'
                        f'<b>求職者：</b>{sq(rec["cd_text"])}<br>'
                        f'<small style="color:#666">↓ CAの返し</small><br>'
                        f'{sq(rec["ca_response"])}'
                        f'</div>', unsafe_allow_html=True)

        # AIによる詳細分析：スルーした場面と改善例
        missed_scenes = eda.get('missed_scenes', [])
        if missed_scenes:
            st.markdown('<div class="sub-title">❌ スルーしてしまった場面 → こう返すべきだった</div>',
                        unsafe_allow_html=True)
            for ms in missed_scenes:
                col_l, col_r = st.columns([1, 1])
                with col_l:
                    st.markdown(
                        f'<div class="emotion-miss">'
                        f'<small>求職者が言った</small><br>'
                        f'<b>{sq(ms.get("cd_text",""))}</b><br>'
                        f'<small style="color:#888">↓ 実際のCAの返し</small><br>'
                        f'<span style="color:#c00000">{sq(ms.get("ca_actual",""))}</span>'
                        f'</div>', unsafe_allow_html=True)
                with col_r:
                    st.markdown(
                        f'<div class="emotion-hit">'
                        f'<small>💡 こう返すと深掘りできた</small><br>'
                        f'<b style="color:#375623">{sq(ms.get("ca_suggested",""))}</b><br>'
                        f'<small style="color:#555">→ {ms.get("why","")}</small>'
                        f'</div>', unsafe_allow_html=True)

        # 深掘りできた場面
        good_scenes = eda.get('good_scenes', [])
        if good_scenes:
            st.markdown('<div class="sub-title">✅ うまく深掘りできた場面</div>', unsafe_allow_html=True)
            for gs in good_scenes:
                st.markdown(
                    f'<div class="emotion-hit">'
                    f'<small>求職者</small>　<b>{sq(gs.get("cd_text",""))}</b><br>'
                    f'<small style="color:#888">↓ CA（深掘り）</small><br>'
                    f'<span style="color:#375623"><b>{sq(gs.get("ca_text",""))}</b></span><br>'
                    f'<small>💡 {gs.get("why_good","")}</small>'
                    f'</div>', unsafe_allow_html=True)

        # 縦の深掘りシーケンス
        st.markdown('<div class="sub-title">🔽 縦の深掘り実例（同テーマ連続）</div>',
                    unsafe_allow_html=True)
        if eda.get('vertical_drill_comment'):
            st.write('📌 ' + eda['vertical_drill_comment'])

        vd_seqs = metrics.get('縦深掘りシーケンス', [])
        if vd_seqs:
            for seq_i, seq in enumerate(vd_seqs[:2]):
                with st.expander(f'縦深掘りシーケンス {seq_i+1}（{len(seq)}発話）', expanded=(seq_i==0)):
                    for u in seq:
                        is_ca_spk = metrics['is_ca'](u['speaker'])
                        color = '#2E75B6' if is_ca_spk else '#444'
                        role  = 'CA 🔵' if is_ca_spk else '求職者 ⚪'
                        st.markdown(
                            f'<div class="vertical-drill">'
                            f'<small style="color:{color}"><b>[{role}]</b></small>　{u["text"][:120]}'
                            f'</div>', unsafe_allow_html=True)
        else:
            st.warning('⚠️ 同テーマ3回以上の縦の深掘りは検出されませんでした。次の面談で意識してみましょう。')
            st.markdown(
                '**縦深掘りの例：**\n'
                '> 求職者「やりがいを感じていました」\n'
                '> CA「なんでやりがいを感じていたんですか？」\n'
                '> 求職者「お客様に喜ばれた時」\n'
                '> CA「どんな時に一番喜ばれましたか？」\n'
                '> 求職者「名前を覚えてもらった時」\n'
                '> CA「その時、どんな気持ちになりましたか？」')

        st.divider()

        # ════════════════════════════════════════════════
        # 自己開示の分析
        # ════════════════════════════════════════════════
        st.markdown('<div class="section-title">🙋 自己開示の分析</div>', unsafe_allow_html=True)
        sda = deep_result.get('self_disclosure_analysis', {})

        sd_moments = metrics.get('自己開示モーメント', [])
        col_sd1, col_sd2 = st.columns([1, 1])

        with col_sd1:
            st.markdown('**✅ 自己開示できた場面**')
            if sd_moments:
                for m_sd in sd_moments[:3]:
                    trigger_html = (f'<small>（きっかけ：求職者「{m_sd["trigger"]}」）</small><br>'
                                    if m_sd.get('trigger') else '')
                    st.markdown(
                        f'<div class="sd-hit">'
                        f'{trigger_html}'
                        f'CA：<b>{sq(m_sd["ca_text"])}</b>'
                        f'</div>', unsafe_allow_html=True)
            # AIが検出した場面も追加
            for fs in (sda.get('found_scenes') or [])[:2]:
                timing_color = '#375623' if fs.get('timing_eval')=='良い' else '#c55a11'
                st.markdown(
                    f'<div class="sd-hit">'
                    f'CA：<b>{sq(fs.get("ca_text",""))}</b><br>'
                    f'<small style="color:{timing_color}">タイミング：{fs.get("timing_eval","")}　'
                    f'→ {fs.get("effect","")}</small>'
                    f'</div>', unsafe_allow_html=True)
            if not sd_moments and not sda.get('found_scenes'):
                st.warning('自己開示が検出されませんでした')

        with col_sd2:
            st.markdown('**💡 ここで自己開示できた（改善例）**')
            missed_sd = sda.get('missed_opportunities') or []
            if missed_sd:
                for ms_sd in missed_sd[:3]:
                    st.markdown(
                        f'<div class="bt-miss">'
                        f'<small>求職者{sq(ms_sd.get("cd_text",""))}</small><br>'
                        f'→ <b style="color:#1F3864">{sq(ms_sd.get("ca_suggested",""))}</b>'
                        f'</div>', unsafe_allow_html=True)
            else:
                st.info('改善機会の特定はAI分析を参照してください')

        if sda.get('advice'):
            st.info('📌 ' + sda['advice'])

        st.divider()

        # ════════════════════════════════════════════════
        # バックトラッキングの分析
        # ════════════════════════════════════════════════
        st.markdown('<div class="section-title">🔁 バックトラッキングの分析</div>', unsafe_allow_html=True)
        bta = deep_result.get('backtrack_analysis', {})

        bt_moments = metrics.get('バックトラッキングモーメント', [])
        col_bt1, col_bt2 = st.columns([1, 1])

        with col_bt1:
            st.markdown('**✅ バックトラッキングできた場面**')
            if bt_moments:
                for m_bt in bt_moments[:3]:
                    ref_html = (f'<small>（参照：{sq(m_bt["referenced"])}）</small><br>'
                                if m_bt.get('referenced') else '')
                    st.markdown(
                        f'<div class="bt-hit">'
                        f'{ref_html}'
                        f'CA：<b>{sq(m_bt["ca_text"])}</b>'
                        f'</div>', unsafe_allow_html=True)
            for fs in (bta.get('found_scenes') or [])[:2]:
                eff_color = '#375623' if fs.get('effect')=='良い' else '#c55a11'
                st.markdown(
                    f'<div class="bt-hit">'
                    f'<small>参照：{sq(fs.get("referenced_cd",""))}</small><br>'
                    f'CA：<b>{sq(fs.get("ca_text",""))}</b><br>'
                    f'<small style="color:{eff_color}">評価：{fs.get("effect","")}</small>'
                    f'</div>', unsafe_allow_html=True)
            if not bt_moments and not bta.get('found_scenes'):
                st.warning('バックトラッキングが検出されませんでした')

        with col_bt2:
            st.markdown('**💡 こう使えた（改善例）**')
            missed_bt = bta.get('missed_opportunities') or []
            if missed_bt:
                for ms_bt in missed_bt[:3]:
                    st.markdown(
                        f'<div class="bt-miss">'
                        f'<small>求職者キーワード：{sq(ms_bt.get("cd_keyword",""))}</small><br>'
                        f'→ <b style="color:#1F3864">{sq(ms_bt.get("ca_suggested",""))}</b>'
                        f'</div>', unsafe_allow_html=True)

        if bta.get('advice'):
            st.info('📌 ' + bta['advice'])
        if bta.get('summary'):
            st.write('📊 ' + bta['summary'])

        st.divider()

        # ════════════════════════════════════════════════
        # 次の面談で使えるフレーズ集（パーソナライズ）
        # ════════════════════════════════════════════════
        st.markdown('<div class="section-title">🗣️ 次の面談で使えるフレーズ集（この求職者のキーワードを使用）</div>',
                    unsafe_allow_html=True)
        next_phrases = deep_result.get('next_phrases', [])
        if next_phrases:
            for ph in next_phrases:
                if not ph.get('phrase'): continue
                st.markdown(
                    f'<div style="margin:8px 0">'
                    f'<small style="color:#888; background:#f0f0f0; padding:2px 8px; border-radius:10px">'
                    f'{ph.get("situation","")}</small><br>'
                    f'<div class="phrase-box">{sq(ph.get("phrase",""))}</div>'
                    f'<small style="color:#555">💡 {ph.get("why","")}</small>'
                    f'</div>', unsafe_allow_html=True)
        else:
            st.info('フレーズ集はAI分析が完了すると表示されます')

        st.divider()

        # ════════════════════════════════════════════════
        # クロージング評価
        # ════════════════════════════════════════════════
        st.markdown('<div class="section-title">🏁 クロージング評価</div>', unsafe_allow_html=True)
        latter_r = metrics['後半求職者比率']
        next_cnt = metrics['次回接続発話数']
        close_q  = metrics['クロージング問いかけ']

        cl1, cl2, cl3 = st.columns(3)
        cl1.metric('後半求職者比率', f'{latter_r:.0%}',
                   delta='✅50%以上' if latter_r>=0.5 else '❌50%未満', delta_color='off')
        cl2.metric('問いかけ回数', f'{close_q}回',
                   delta='✅あり' if close_q>=1 else '❌なし', delta_color='off')
        cl3.metric('次回接続発話', f'{next_cnt}回',
                   delta='✅あり' if next_cnt>=1 else '❌なし', delta_color='off')

        if overall.get('closing_eval'):
            st.write(overall['closing_eval'])

        ca_close = metrics.get('クロージングCA', [])
        cd_close = metrics.get('クロージングCD', [])
        if ca_close or cd_close:
            with st.expander('クロージング発話を確認'):
                for t in ca_close: st.markdown('🔵 **CA：** ' + t[:120])
                for t in cd_close: st.markdown('⚪ **求職者：** ' + t[:120])

        st.divider()

        # ════════════════════════════════════════════════
        # ルーブリック採点（5軸）
        # ════════════════════════════════════════════════
        st.markdown('<div class="section-title">🎯 ルーブリック採点（5軸）</div>', unsafe_allow_html=True)
        AXES = ['意向','適正','条件','認識統一','気づき']
        AXIS_LABELS = {
            '意向':    '意向把握（価値観を引き出し、意向を広げられたか）',
            '適正':    '適正把握（経験・強みを把握できたか）',
            '条件':    '条件把握（Must/Betterで確認・期待値調整できたか）',
            '認識統一':'認識統一（価値観・今後の方向性に本人の同意を得られたか）',
            '気づき':  '気づき付与（他の選択肢・新しい可能性に気づかせられたか）',
        }
        for axis in AXES:
            d   = gd.get(axis, {})
            s   = d.get('score', 0)
            cls = 'score-high' if s>=2 else ('score-mid' if s==1 else 'score-low')
            ev  = d.get('evidence', [])
            bar = '█'*s + '░'*(3-s)
            na  = d.get('next_action','')
            html = (
                f'<div class="score-box {cls}">'
                f'<b>{AXIS_LABELS[axis]}</b> &nbsp; <code>{bar}</code> <b>{s}/3</b>'
                + (f'<br>💪 {d["strength"]}' if d.get('strength') else '')
                + (f'<br>📌 {d["weakness"]}' if d.get('weakness') else '')
                + (f'<br><small style="color:#595959">根拠：{sq(ev[0])}</small>' if ev else '')
                + (f'<br>🚀 <b>次のアクション：</b>{na}' if na else '')
                + '</div>'
            )
            st.markdown(html, unsafe_allow_html=True)

        st.divider()

        # ════════════════════════════════════════════════
        # チェックポイント
        # ════════════════════════════════════════════════
        st.markdown('<div class="section-title">✅ チェックポイント（自動採点）</div>', unsafe_allow_html=True)
        checks = evaluate_checklist(metrics, claude_result)
        ok_count = warn_count = ng_count = 0

        for section, items in checks.items():
            with st.expander(f'**{section}**', expanded=('WILL' in section or 'クロージング' in section)):
                for label, detail, is_ok, is_warn in items:
                    detail_html = (f'  <small>{detail}</small>' if detail else '')
                    if is_ok:
                        st.markdown(f'<span class="check-ok">✅ {label}</span>{detail_html}', unsafe_allow_html=True)
                        ok_count += 1
                    elif is_warn:
                        st.markdown(f'<span class="check-warn">⚠️ {label}</span>{detail_html}', unsafe_allow_html=True)
                        warn_count += 1
                    else:
                        st.markdown(f'<span class="check-ng">❌ {label}</span>{detail_html}', unsafe_allow_html=True)
                        ng_count += 1

        total_checks = ok_count + warn_count + ng_count
        st.progress(ok_count/total_checks if total_checks>0 else 0)
        st.caption(f'✅ {ok_count}項目クリア  ⚠️ {warn_count}項目要注意  ❌ {ng_count}項目要改善  （全{total_checks}項目）')

        st.divider()

        # ════════════════════════════════════════════════
        # 行動指標
        # ════════════════════════════════════════════════
        st.markdown('<div class="section-title">📊 行動指標（数値詳細）</div>', unsafe_allow_html=True)
        bh_data = [
            ('求職者発話比率', f'{metrics["求職者発話比率"]:.0%}',        '目標40%以上',    metrics['求職者発話比率']>=0.4),
            ('後半求職者比率', f'{metrics["後半求職者比率"]:.0%}',        '目標50%以上',    metrics['後半求職者比率']>=0.5),
            ('会話回数（ラリー）', f'{metrics["会話回数"]}回',            '参考:平均130回', True),
            ('名前呼称回数',   f'{metrics["名前呼称回数"]}回',            '目標3回以上',    metrics['名前呼称回数']>=3),
            ('フィラー回数',   f'{metrics["フィラー回数"]}回',            '目標30回以下',   metrics['フィラー回数']<=30),
            ('価値観深掘り',   f'{metrics["深掘り_価値観"]}回',           '目標4回以上',    metrics['深掘り_価値観']>=4),
            ('バックトラッキング', f'{metrics["バックトラッキング"]}回',  '目標5回以上',    metrics['バックトラッキング']>=5),
            ('自己開示',       f'{metrics["自己開示回数"]}回',            '目標2回以上',    metrics['自己開示回数']>=2),
            ('ポジティブ反応', f'{metrics["ポジティブ反応"]}回',          '目標5回以上',    metrics['ポジティブ反応']>=5),
            ('感情スルー率',   f'{metrics["感情スルー率"]}%',             '目標50%以下',    metrics['感情スルー率']<=50),
            ('縦深掘り最大',   f'{metrics["縦深掘り最大"]}回連続',        '目標3回以上',    metrics['縦深掘り最大']>=3),
            ('NG_急かし',      f'{metrics["NG_急かし"]}回',               '目標0回',        metrics['NG_急かし']==0),
        ]
        cols = st.columns(4)
        for i, (label, val, target, ok) in enumerate(bh_data):
            with cols[i%4]:
                icon = '✅' if ok else '⚠️'
                st.metric(f'{icon} {label}', val, help=target)

        # ── フェーズ網羅 ─────────────────────────────
        phases     = claude_result.get('behaviors',{}).get('フェーズ網羅',[])
        all_phases = ['冒頭','CAN','WILL','MUST','クロージング']
        st.markdown('<div class="section-title">🗺️ フェーズ網羅</div>', unsafe_allow_html=True)
        ph_cols = st.columns(5)
        for i, ph in enumerate(all_phases):
            ph_cols[i].metric(ph, '✅ 到達' if ph in phases else '❌ 未到達')

        # ── フィラー詳細 ──────────────────────────────
        if metrics['フィラー回数'] > 30:
            st.markdown('<div class="section-title">💬 フィラー詳細</div>', unsafe_allow_html=True)
            from collections import Counter
            fillers = FILLER_PAT.findall(' '.join(metrics['ca_texts']))
            st.bar_chart(dict(Counter(fillers).most_common(8)))

        # ── notes ─────────────────────────────────────
        if claude_result.get('notes'):
            with st.expander('📝 分析メモ'):
                st.write(claude_result['notes'])

        st.divider()
        with st.expander('💾 保存先の確認'):
            st.code(str(json_path), language=None)

        # ── Word出力 / Slack送信 ──────────────────────────
        st.markdown('<div class="section-title">📤 アクション</div>', unsafe_allow_html=True)
        # Word/Slack用に全フィールドをまとめた辞書
        _full_data = {
            "ca": ca_input, "grip": grip_input,
            "candidate": cand_input, "meeting_type": meeting_type,
            **claude_result,
            "behaviors": {**metrics, **claude_result.get('behaviors', {})},
            "emotion_drill_analysis":  deep_result.get('emotion_drill_analysis', {}),
            "self_disclosure_analysis": deep_result.get('self_disclosure_analysis', {}),
            "backtrack_analysis":       deep_result.get('backtrack_analysis', {}),
            "next_phrases":             deep_result.get('next_phrases', []),
        }
        act_col1, act_col2 = st.columns(2)
        with act_col1:
            st.markdown('**📄 Word文書として出力**')
            if st.button('Word文書を生成', key='app_gen_word', use_container_width=True):
                try:
                    docx_bytes = generate_word_doc(_full_data)
                    fname = f"{ca_input}_{cand_input}_FB.docx"
                    st.download_button('⬇️ ダウンロード', data=docx_bytes,
                                       file_name=fname,
                                       mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                       key='app_dl_word')
                except Exception as e:
                    st.error(f'Word生成エラー: {e}')
        with act_col2:
            st.markdown('**💬 Slackに送信**')
            if SLACK_WEBHOOK:
                if st.button('Slackに送信', key='app_send_slack', use_container_width=True):
                    ok_flag, msg = send_to_slack(SLACK_WEBHOOK, _full_data)
                    if ok_flag:
                        st.success(f'✅ {msg}')
                    else:
                        st.error(f'❌ {msg}')
            else:
                st.caption('⚠️ SLACK_WEBHOOK_URL が未設定です')

else:
    st.info('👆 左のサイドバーにCA名・求職者名を入力し、ファイルをアップロードしてください')
    st.markdown("""
**対応ファイル形式**
- `.docx`：話者ラベルあり（CA名 / 求職者名 形式）→ API不要で即解析
- `.txt`：音声文字起こし → Claude Haikuで話者推定

**推奨ファイル命名規則**（自動入力されます）
```
CA名_グリップランク_求職者名_初回面談.docx
例: 下川_グリップA_岡寺さん_初回面談.docx
```

**このツールが分析すること**
- 🔍 感情深掘りマップ：全感情ワード場面を一覧 → スルーした場面ごとに改善例
- 🙋 自己開示分析：できた場面 + どこで使えたかの提案
- 🔁 バックトラッキング分析：引用できた/できなかった場面
- 🗣️ 次の面談で使えるフレーズ集（この求職者のキーワードを使ったパーソナライズ版）
""")
