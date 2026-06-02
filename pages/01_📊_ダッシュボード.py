# -*- coding: utf-8 -*-
"""
面談分析ダッシュボード
CA・グリップ・候補者名・面談種別でフィルタリングして過去の分析結果を閲覧
"""

import streamlit as st
import json, re, os, sys, io
from pathlib import Path
import pandas as pd
import requests

# analysis_core を親ディレクトリからimport
sys.path.insert(0, str(Path(__file__).parent.parent))
import analysis_core as core
import anthropic
from gdrive import (list_json_files, download_json as gdrive_download_json,
                    upload_summary, download_summary, upload_json)

st.set_page_config(page_title="面談分析ダッシュボード", page_icon="📊", layout="wide")

# ── スタイル ──────────────────────────────────────────────
st.markdown("""
<style>
.section-title { background:#1F3864; color:white; padding:8px 16px; border-radius:6px;
                 font-weight:bold; margin:16px 0 8px 0; }
.kpi-card { background:#f0f4ff; border-left:4px solid #2E75B6; padding:10px 14px;
            margin:4px 0; border-radius:4px; }
.score-high { border-left-color:#375623; background:#e2efda; }
.score-mid  { border-left-color:#c55a11; background:#fce4d6; }
.score-low  { border-left-color:#c00000; background:#fcecea; }
.grip-A { background:#1e8449; color:white; padding:2px 8px; border-radius:10px; font-weight:bold; }
.grip-B { background:#2471a3; color:white; padding:2px 8px; border-radius:10px; font-weight:bold; }
.grip-C { background:#d35400; color:white; padding:2px 8px; border-radius:10px; font-weight:bold; }
.grip-D { background:#c0392b; color:white; padding:2px 8px; border-radius:10px; font-weight:bold; }
.grip-X { background:#888;    color:white; padding:2px 8px; border-radius:10px; font-weight:bold; }
.phrase-box { background:#1F3864; color:white; padding:8px 12px; border-radius:6px;
              margin:4px 0; font-family:monospace; font-size:0.9rem; }
.emotion-miss { background:#fcecea; border-left:4px solid #c00000; padding:8px 12px; border-radius:4px; margin:6px 0; }
.emotion-hit  { background:#e2efda; border-left:4px solid #375623; padding:8px 12px; border-radius:4px; margin:6px 0; }
.bt-hit  { background:#EBF5FB; border-left:4px solid #2E75B6; padding:8px 12px; border-radius:4px; margin:6px 0; }
.sd-hit  { background:#F0F0F0; border-left:4px solid #555; padding:8px 12px; border-radius:4px; margin:6px 0; }
</style>
""", unsafe_allow_html=True)

OUTPUT_JSON   = Path(__file__).parent.parent / "output" / "json"
SLACK_WEBHOOK = st.secrets.get("SLACK_WEBHOOK_URL", "")

def sq(text: str) -> str:
    """「」の二重表示を防ぐ：既に括弧があればそのまま、なければ付ける"""
    t = (text or '').strip()
    if t.startswith('「') and t.endswith('」'):
        return t
    return f'「{t}」'

# ── Word文書生成 ──────────────────────────────────────────
def generate_word_doc(d: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

    def h(text, level=1, color=(31,56,100)):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
        return p

    def body(text, bold=False, color=None, size=10.5):
        p = doc.add_paragraph()
        r = p.add_run(str(text))
        r.bold = bold; r.font.size = Pt(size)
        if color: r.font.color.rgb = RGBColor(*color)
        return p

    def kv(key, val, kc=(46,117,182)):
        p = doc.add_paragraph()
        r1 = p.add_run(f'{key}：'); r1.bold = True
        r1.font.color.rgb = RGBColor(*kc); r1.font.size = Pt(10.5)
        r2 = p.add_run(str(val or '')); r2.font.size = Pt(10.5)

    def table(headers, rows, style='Table Grid'):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = style
        for i, h_ in enumerate(headers):
            c = t.rows[0].cells[i]; c.text = h_
            if c.paragraphs[0].runs: c.paragraphs[0].runs[0].bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row): cells[i].text = str(v or '')
        doc.add_paragraph()
        return t

    gd  = d.get('grip_drivers', {})
    bh  = d.get('behaviors', {})
    ov  = d.get('overall', {})
    eda = d.get('emotion_drill_analysis', {})
    sda = d.get('self_disclosure_analysis', {})
    bta = d.get('backtrack_analysis', {})
    phrases = d.get('next_phrases', [])
    axes = ['意向','適正','条件','認識統一','気づき']
    total = sum(gd.get(ax,{}).get('score',0) for ax in axes)
    grade = ov.get('grade','') or ('S' if total>=13 else 'A' if total>=10 else 'B' if total>=7 else 'C' if total>=4 else 'D')
    AXIS_LBL = {'意向':'意向把握','適正':'適正把握','条件':'条件把握','認識統一':'認識統一','気づき':'気づき付与'}

    # ══ タイトル ══
    t0 = doc.add_heading('初回面談 分析レポート', 0)
    t0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kv('CA名', d.get('ca','')); kv('候補者名', d.get('candidate',''))
    kv('グリップ', d.get('grip','')); kv('面談種別', d.get('meeting_type',''))
    doc.add_paragraph()

    # ══ 1. 総合評価 ══
    h('■ 1. 総合評価')
    p = doc.add_paragraph()
    r = p.add_run(f'グレード：{grade}　　総合スコア：{total}/15')
    r.bold = True; r.font.size = Pt(14)
    if ov.get('grade_reason'): body(ov['grade_reason'])
    if ov.get('top_strength'): body(f'【最大の強み】{ov["top_strength"]}', color=(55,86,35))
    if ov.get('best_exchange'): body(f'【最も良かったやり取り】{ov["best_exchange"]}', color=(55,86,35))
    if ov.get('missed_moment'): body(f'【最も惜しかった場面】{ov["missed_moment"]}', color=(192,0,0))
    doc.add_paragraph()

    # ══ 2. ルーブリック採点（5軸） ══
    h('■ 2. ルーブリック採点（5軸）')
    table(
        ['評価軸','スコア','強み','改善点','次のアクション'],
        [[AXIS_LBL.get(ax,ax),
          f'{gd.get(ax,{}).get("score",0)}/3',
          gd.get(ax,{}).get('strength',''),
          gd.get(ax,{}).get('weakness',''),
          gd.get(ax,{}).get('next_action','')] for ax in axes]
    )

    # ══ 3. 優先改善ポイント ══
    if ov.get('top_issues'):
        h('■ 3. 優先改善ポイント TOP3')
        for i, issue in enumerate(ov['top_issues'][:3], 1):
            if isinstance(issue, dict):
                body(f'{i}. {issue.get("issue","")}', bold=True)
                if issue.get('detail'): body(f'　📌 {issue["detail"]}')
                if issue.get('fix'):    body(f'　💡 改善例：{issue["fix"]}', color=(55,86,35))
        doc.add_paragraph()

    if ov.get('one_thing'):
        h('■ 4. 次の面談で必ず試してほしいこと')
        body(ov['one_thing'], color=(46,117,182), bold=True)
        doc.add_paragraph()

    # ══ 5. クロージング評価 ══
    if ov.get('closing_eval'):
        h('■ 5. クロージング評価')
        body(ov['closing_eval'])
        doc.add_paragraph()

    # ══ 6. 感情深掘り分析 ══
    h('■ 6. 感情深掘り分析')
    if eda.get('summary'): body(eda['summary'])
    missed = eda.get('missed_scenes') or []
    if missed:
        body('【スルーしてしまった場面 → 改善例】', bold=True)
        for ms in missed:
            body(f'求職者：{ms.get("cd_text","")}')
            body(f'実際のCA：{ms.get("ca_actual","")}', color=(192,0,0))
            body(f'→ こう返すべきだった：{ms.get("ca_suggested","")}', color=(55,86,35))
            if ms.get('why'): body(f'　理由：{ms["why"]}')
            doc.add_paragraph()
    good = eda.get('good_scenes') or []
    if good:
        body('【うまく深掘りできた場面】', bold=True)
        for gs in good:
            body(f'求職者：{gs.get("cd_text","")}')
            body(f'CA（深掘り）：{gs.get("ca_text","")}', color=(55,86,35))
            if gs.get('why_good'): body(f'　良かった点：{gs["why_good"]}')
            doc.add_paragraph()
    if eda.get('vertical_drill_comment'):
        body(f'【縦の深掘りコメント】{eda["vertical_drill_comment"]}')
    doc.add_paragraph()

    # ══ 7. 自己開示の分析 ══
    h('■ 7. 自己開示の分析')
    if sda.get('summary'): body(sda['summary'])
    for fs in (sda.get('found_scenes') or []):
        body(f'CA：{fs.get("ca_text","")}')
        body(f'　タイミング評価：{fs.get("timing_eval","")}　/ {fs.get("effect","")}', color=(46,117,182))
    missed_sd = sda.get('missed_opportunities') or []
    if missed_sd:
        body('【ここで自己開示できた → 改善例】', bold=True)
        for ms in missed_sd:
            body(f'求職者：{ms.get("cd_text","")}')
            body(f'→ {ms.get("ca_suggested","")}', color=(55,86,35))
    if sda.get('advice'): body(f'アドバイス：{sda["advice"]}')
    doc.add_paragraph()

    # ══ 8. バックトラッキングの分析 ══
    h('■ 8. バックトラッキングの分析')
    if bta.get('summary'): body(bta['summary'])
    for fs in (bta.get('found_scenes') or []):
        body(f'参照：{fs.get("referenced_cd","")}')
        body(f'CA：{fs.get("ca_text","")}', color=(46,117,182))
        body(f'　評価：{fs.get("effect","")}')
    missed_bt = bta.get('missed_opportunities') or []
    if missed_bt:
        body('【このキーワードを引用できた → 改善例】', bold=True)
        for ms in missed_bt:
            body(f'キーワード：{ms.get("cd_keyword","")}')
            body(f'→ {ms.get("ca_suggested","")}', color=(55,86,35))
    if bta.get('advice'): body(f'アドバイス：{bta["advice"]}')
    doc.add_paragraph()

    # ══ 9. 行動指標 ══
    h('■ 9. 行動指標')
    ok = lambda v, thr: '✅' if v >= thr else '❌'
    table(
        ['指標','値','目標','判定'],
        [
            ('求職者発話比率',  f'{round(bh.get("求職者発話比率",0)*100)}%',  '40%以上',  ok(bh.get('求職者発話比率',0)*100, 40)),
            ('後半求職者比率',  f'{round(bh.get("後半求職者比率",0)*100)}%',  '50%以上',  ok(bh.get('後半求職者比率',0)*100, 50)),
            ('フィラー回数',    f'{bh.get("フィラー回数",0)}回',               '30回以下',  '✅' if bh.get('フィラー回数',0)<=30 else '❌'),
            ('ポジティブ反応',  f'{bh.get("ポジティブ反応",0)}回',             '5回以上',   ok(bh.get('ポジティブ反応',0), 5)),
            ('感情スルー率',    f'{bh.get("感情スルー率",0)}%',                '50%以下',   '✅' if bh.get('感情スルー率',0)<=50 else '❌'),
            ('価値観深掘り',    f'{bh.get("深掘り_価値観",0)}回',              '4回以上',   ok(bh.get('深掘り_価値観',0), 4)),
            ('縦深掘り最大',    f'{bh.get("縦深掘り最大",0)}回連続',           '3回以上',   ok(bh.get('縦深掘り最大',0), 3)),
            ('バックトラッキング',f'{bh.get("バックトラッキング",0)}回',        '5回以上',   ok(bh.get('バックトラッキング',0), 5)),
            ('自己開示回数',    f'{bh.get("自己開示回数",0)}回',               '2回以上',   ok(bh.get('自己開示回数',0), 2)),
            ('名前呼称回数',    f'{bh.get("名前呼称回数",0)}回',               '3回以上',   ok(bh.get('名前呼称回数',0), 3)),
            ('MUST提案',        '✅' if bh.get('MUST提案') else '❌',           'あり',      '✅' if bh.get('MUST提案') else '❌'),
            ('次回アポ確定',    '✅' if bh.get('次回アポ確定') else '❌',       'あり',      '✅' if bh.get('次回アポ確定') else '❌'),
        ]
    )

    # ══ 10. フレーズ集 ══
    if phrases:
        h('■ 10. 次の面談で使えるフレーズ集')
        for ph in phrases:
            if not ph.get('phrase'): continue
            body(f'【{ph.get("situation","")}】', bold=True)
            p_ph = doc.add_paragraph()
            r_ph = p_ph.add_run(ph.get('phrase',''))
            r_ph.font.color.rgb = RGBColor(31,56,100)
            r_ph.bold = True; r_ph.font.size = Pt(11)
            if ph.get('why'): body(f'　💡 {ph["why"]}')
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Slack送信（全情報版） ──────────────────────────────────
def send_to_slack(webhook_url: str, d: dict) -> tuple[bool, str]:
    gd  = d.get('grip_drivers', {})
    bh  = d.get('behaviors', {})
    ov  = d.get('overall', {})
    eda = d.get('emotion_drill_analysis', {})
    sda = d.get('self_disclosure_analysis', {})
    bta = d.get('backtrack_analysis', {})
    phrases = d.get('next_phrases', [])

    axes  = ['意向','適正','条件','認識統一','気づき']
    total = sum(gd.get(ax,{}).get('score',0) for ax in axes)
    grade = ov.get('grade','') or ('S' if total>=13 else 'A' if total>=10 else 'B' if total>=7 else 'C' if total>=4 else 'D')
    g_emoji = {'S':'🏆','A':'🟢','B':'🔵','C':'🟡','D':'🔴'}.get(grade,'⚪')

    def ok(v, thr, rev=False):
        return '✅' if (v <= thr if rev else v >= thr) else '❌'

    # ルーブリック
    score_lines = '\n'.join(
        f"　{ax}：{'█'*gd.get(ax,{}).get('score',0)}{'░'*(3-gd.get(ax,{}).get('score',0))} "
        f"{gd.get(ax,{}).get('score',0)}/3　{gd.get(ax,{}).get('weakness','')}"
        for ax in axes)

    # 改善ポイント
    issues_text = ''
    for i, issue in enumerate(ov.get('top_issues',[])[:3], 1):
        if isinstance(issue, dict):
            issues_text += f'{i}. *{issue.get("issue","")}*\n　{issue.get("detail","")}\n　💡 {issue.get("fix","")}\n'

    # 感情深掘り
    eda_text = eda.get('summary','')
    missed = eda.get('missed_scenes') or []
    if missed:
        eda_text += '\n*スルーした場面（改善例）:*'
        for ms in missed[:2]:
            eda_text += f'\n　求職者：「{ms.get("cd_text","")}」\n　→ `{ms.get("ca_suggested","")}`'
    good = eda.get('good_scenes') or []
    if good:
        eda_text += '\n*うまく深掘りできた場面:*'
        for gs in good[:1]:
            eda_text += f'\n　求職者：「{gs.get("cd_text","")}」\n　CA：`{gs.get("ca_text","")}`'
    if eda.get('vertical_drill_comment'):
        eda_text += f'\n縦の深掘り：{eda["vertical_drill_comment"]}'

    # 自己開示
    sda_text = sda.get('summary','')
    for ms in (sda.get('missed_opportunities') or [])[:2]:
        sda_text += f'\n　改善例：`{ms.get("ca_suggested","")}`'
    if sda.get('advice'): sda_text += f'\n💡 {sda["advice"]}'

    # バックトラッキング
    bta_text = bta.get('summary','')
    for ms in (bta.get('missed_opportunities') or [])[:2]:
        bta_text += f'\n　改善例：`{ms.get("ca_suggested","")}`'
    if bta.get('advice'): bta_text += f'\n💡 {bta["advice"]}'

    # フレーズ集
    phrases_text = ''
    for ph in phrases:
        if not ph.get('phrase'): continue
        phrases_text += f'*{ph.get("situation","")}*\n　`{ph.get("phrase","")}` — {ph.get("why","")}\n'

    blocks = [
        {"type": "header", "text": {"type": "plain_text",
            "text": f"📊 初回面談 分析レポート｜{d.get('candidate','')}"}},
        {"type": "section", "fields": [
            {"type": "mrkdwn", "text": f"*CA名:* {d.get('ca','')}"},
            {"type": "mrkdwn", "text": f"*グリップ:* {d.get('grip','')}"},
            {"type": "mrkdwn", "text": f"*面談種別:* {d.get('meeting_type','')}"},
            {"type": "mrkdwn", "text": f"*グレード:* {g_emoji} {grade}　スコア: {total}/15"},
        ]},
        {"type": "divider"},
        # ルーブリック
        {"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🎯 ルーブリック採点（5軸）*\n{score_lines}"}},
        {"type": "divider"},
    ]

    # 強み・惜しかった場面・最も良かった場面
    if ov.get('top_strength') or ov.get('best_exchange') or ov.get('missed_moment'):
        txt = ''
        if ov.get('top_strength'):   txt += f'*💪 最大の強み*\n{ov["top_strength"]}\n'
        if ov.get('best_exchange'):  txt += f'*👍 最も良かったやり取り*\n{ov["best_exchange"]}\n'
        if ov.get('missed_moment'):  txt += f'*⚠️ 最も惜しかった場面*\n{ov["missed_moment"]}\n'
        blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": txt.strip()}})
        blocks.append({"type": "divider"})

    # 改善ポイント
    if issues_text:
        blocks.append({"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🔧 優先改善ポイント TOP3*\n{issues_text.strip()}"}})
        blocks.append({"type": "divider"})

    # 次に試すこと
    if ov.get('one_thing'):
        blocks.append({"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🚀 次の面談で必ず試すこと*\n{ov['one_thing']}"}})
        blocks.append({"type": "divider"})

    # クロージング評価
    if ov.get('closing_eval'):
        blocks.append({"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🏁 クロージング評価*\n{ov['closing_eval']}"}})
        blocks.append({"type": "divider"})

    # 感情深掘り
    if eda_text:
        blocks.append({"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🔍 感情深掘り分析*\n{eda_text.strip()}"}})
        blocks.append({"type": "divider"})

    # 自己開示 / バックトラッキング
    sd_bt = ''
    if sda_text: sd_bt += f'*🙋 自己開示*\n{sda_text}\n\n'
    if bta_text: sd_bt += f'*🔁 バックトラッキング*\n{bta_text}'
    if sd_bt:
        blocks.append({"type": "section", "text": {"type": "mrkdwn", "text": sd_bt.strip()}})
        blocks.append({"type": "divider"})

    # 行動指標
    blocks.append({"type": "section", "fields": [
        {"type": "mrkdwn", "text": f"*感情スルー率:* {ok(bh.get('感情スルー率',0),50,rev=True)} {bh.get('感情スルー率',0)}%"},
        {"type": "mrkdwn", "text": f"*フィラー:* {ok(bh.get('フィラー回数',0),30,rev=True)} {bh.get('フィラー回数',0)}回"},
        {"type": "mrkdwn", "text": f"*ポジティブ反応:* {ok(bh.get('ポジティブ反応',0),5)} {bh.get('ポジティブ反応',0)}回"},
        {"type": "mrkdwn", "text": f"*縦深掘り最大:* {ok(bh.get('縦深掘り最大',0),3)} {bh.get('縦深掘り最大',0)}回"},
        {"type": "mrkdwn", "text": f"*バックトラッキング:* {ok(bh.get('バックトラッキング',0),5)} {bh.get('バックトラッキング',0)}回"},
        {"type": "mrkdwn", "text": f"*自己開示:* {ok(bh.get('自己開示回数',0),2)} {bh.get('自己開示回数',0)}回"},
        {"type": "mrkdwn", "text": f"*発話比率(求職者):* {ok(bh.get('求職者発話比率',0)*100,40)} {round(bh.get('求職者発話比率',0)*100)}%"},
        {"type": "mrkdwn", "text": f"*アポ確定:* {'✅' if bh.get('次回アポ確定') else '❌'}"},
    ]})

    # フレーズ集
    if phrases_text:
        blocks.append({"type": "divider"})
        blocks.append({"type": "section", "text": {"type": "mrkdwn",
            "text": f"*🗣️ 次の面談で使えるフレーズ集*\n{phrases_text.strip()}"}})

    try:
        resp = requests.post(webhook_url, json={"blocks": blocks}, timeout=10)
        if resp.status_code == 200:
            return True, '送信成功'
        return False, f'エラー: {resp.status_code} {resp.text}'
    except Exception as e:
        return False, str(e)


# ── データ読み込み ─────────────────────────────────────────
@st.cache_data(ttl=600)
def load_all_records():
    """サマリーJSONが存在すればそれを使う（高速）。なければ全件ダウンロードして自動保存。"""
    # ① サマリーJSONを試す（1ファイル取得で完結）
    try:
        summary = download_summary()
        if summary:
            return pd.DataFrame(summary)
    except Exception:
        pass

    # ② フォールバック: Google Drive から全件ダウンロード
    records = []
    try:
        drive_files = list_json_files(subfolder="json")
        # サマリーファイル自体は除外
        drive_files = [f for f in drive_files if not f["name"].startswith("_")]
        for f in sorted(drive_files, key=lambda x: x["name"]):
            try:
                d = gdrive_download_json(f["id"])
                _append_record(records, d, f["name"], f["id"])
            except Exception:
                continue
        # 次回以降はサマリーで高速読み込みできるよう自動保存
        if records:
            try:
                upload_summary(records)
            except Exception:
                pass
        return pd.DataFrame(records)
    except Exception:
        pass

    # ③ 最終フォールバック: ローカルから読み込み
    for f in sorted(OUTPUT_JSON.glob("*.json")):
        try:
            d = json.loads(f.read_text(encoding='utf-8'))
            _append_record(records, d, f.name, str(f))
        except Exception:
            continue
    return pd.DataFrame(records)


def rebuild_summary():
    """全件ダウンロードしてサマリーJSONを再構築。「再読み込み」ボタンから呼ぶ。"""
    records = []
    try:
        drive_files = list_json_files(subfolder="json")
        drive_files = [f for f in drive_files if not f["name"].startswith("_")]
        for f in sorted(drive_files, key=lambda x: x["name"]):
            try:
                d = gdrive_download_json(f["id"])
                _append_record(records, d, f["name"], f["id"])
            except Exception:
                continue
        if records:
            upload_summary(records)
    except Exception:
        pass
    return records


# ── セルフチェック（自己評価 vs AI のズレを蓄積） ─────────────
@st.cache_data(ttl=600)
def load_selfchecks():
    """Driveのselfcheckフォルダから全自己採点を取得。{参照jsonファイル名: data} で返す"""
    out = {}
    try:
        files = list_json_files(subfolder="selfcheck")
        for f in files:
            if f["name"].startswith("_"):
                continue
            try:
                d = gdrive_download_json(f["id"])
                ref = d.get("_ref_file")
                if ref:
                    out[ref] = d
            except Exception:
                continue
    except Exception:
        pass
    return out

def save_selfcheck(ref_file, ca, grip, candidate, meeting_type,
                   self_scores, behavior_checks, next_one_thing, best_self=""):
    """自己採点をDriveのselfcheckフォルダに保存"""
    from datetime import datetime
    data = {
        "_ref_file": ref_file,
        "ca": ca, "grip": grip, "candidate": candidate, "meeting_type": meeting_type,
        "self_scores": self_scores,
        "behavior_checks": behavior_checks,
        "best_self": best_self,
        "next_one_thing": next_one_thing,
        "checked_at": datetime.now().isoformat(timespec="seconds"),
    }
    upload_json(f"selfcheck_{ref_file}", data, subfolder="selfcheck")

def _append_record(records, d, filename, path_or_id):
    try:
        gd = d.get('grip_drivers', {})
        bh = d.get('behaviors', {})
        ov = d.get('overall', {})

        axes      = ['意向','適正','条件','認識統一','気づき']
        scores    = {ax: gd.get(ax, {}).get('score', 0) for ax in axes}
        total     = sum(scores.values())

        grade = ov.get('grade', '')
        if not grade:
            if total >= 13: grade = 'S'
            elif total >= 10: grade = 'A'
            elif total >= 7:  grade = 'B'
            elif total >= 4:  grade = 'C'
            else:             grade = 'D'

        records.append({
            '_file':       filename,
            '_path':       path_or_id,
            'CA':          d.get('ca', ''),
            'グリップ':     d.get('grip', 'X'),
            '候補者':       d.get('candidate', ''),
            '面談種別':     d.get('meeting_type', ''),
            'グレード':     grade,
            '総合スコア':   total,
            '意向':         scores['意向'],
            '適正':         scores['適正'],
            '条件':         scores['条件'],
            '認識統一':     scores['認識統一'],
            '気づき':       scores['気づき'],
            '求職者発話比率':  round(bh.get('求職者発話比率', 0) * 100),
            'フィラー回数':    bh.get('フィラー回数', 0),
            'ポジティブ反応':  bh.get('ポジティブ反応', 0),
            '感情スルー率':    bh.get('感情スルー率', 0),
            '深掘り_価値観':   bh.get('深掘り_価値観', 0),
            'バックトラッキング': bh.get('バックトラッキング', 0),
            '縦深掘り最大':    bh.get('縦深掘り最大', 0),
            '自己開示回数':    bh.get('自己開示回数', 0),
            'MUST提案':        bh.get('MUST提案', False),
            '次回アポ確定':    bh.get('次回アポ確定', False),
            '新形式':          'overall' in d,
        })
    except Exception:
        pass

df_all = load_all_records()

if df_all.empty:
    st.warning('分析済みデータが見つかりません。まず「FB生成ツール」でファイルを分析してください。')
    st.stop()

# ── サイドバー：フィルター ────────────────────────────────
with st.sidebar:
    st.header('🔍 フィルター')

    ca_opts    = ['全員'] + sorted(df_all['CA'].unique().tolist())
    grip_opts  = ['全て'] + sorted(df_all['グリップ'].unique().tolist())
    type_opts  = ['全て'] + sorted(df_all['面談種別'].unique().tolist())
    grade_opts = ['全て', 'S', 'A', 'B', 'C', 'D']

    sel_ca    = st.selectbox('CA', ca_opts)
    sel_grip  = st.multiselect('グリップランク', grip_opts[1:], default=grip_opts[1:])
    sel_type  = st.selectbox('面談種別', type_opts)
    sel_grade = st.multiselect('グレード', grade_opts[1:], default=grade_opts[1:])
    cand_q    = st.text_input('候補者名で検索', placeholder='例：岡寺')

    st.divider()
    st.caption('💬 Slack Webhook：設定済み ✅')

    st.divider()
    st.caption(f'総データ数：{len(df_all)}件')
    if st.button('🔄 データを再読み込み'):
        with st.spinner('Drive から全件取得してサマリーを更新中...'):
            rebuild_summary()
        load_all_records.clear()
        st.rerun()

# ── フィルタリング ────────────────────────────────────────
df = df_all.copy()
if sel_ca    != '全員':               df = df[df['CA'] == sel_ca]
if sel_grip:                           df = df[df['グリップ'].isin(sel_grip)]
if sel_type  != '全て':               df = df[df['面談種別'] == sel_type]
if sel_grade:                          df = df[df['グレード'].isin(sel_grade)]
if cand_q:                             df = df[df['候補者'].str.contains(cand_q, na=False)]

# ════════════════════════════════════════════════════════
# ヘッダー
# ════════════════════════════════════════════════════════
st.title('📊 面談分析ダッシュボード')
st.caption(f'フィルター後：{len(df)}件 / 全{len(df_all)}件')

if df.empty:
    st.info('条件に一致するデータがありません。フィルターを変更してください。')
    st.stop()

# ════════════════════════════════════════════════════════
# KPIサマリー行
# ════════════════════════════════════════════════════════
st.markdown('<div class="section-title">📈 サマリー指標（フィルター後）</div>', unsafe_allow_html=True)

k1,k2,k3,k4,k5,k6,k7 = st.columns(7)
k1.metric('件数',          f'{len(df)}件')
k2.metric('平均総合スコア', f'{df["総合スコア"].mean():.1f}/15')
k3.metric('平均感情スルー率',f'{df["感情スルー率"].mean():.0f}%',
          delta='低いほど良い', delta_color='off')
k4.metric('平均フィラー回数', f'{df["フィラー回数"].mean():.0f}回',
          delta='少ないほど良い', delta_color='off')
k5.metric('平均ポジティブ反応', f'{df["ポジティブ反応"].mean():.1f}回')
k6.metric('平均発話比率(求職者)', f'{df["求職者発話比率"].mean():.0f}%')
k7.metric('アポ確定率',
          f'{df["次回アポ確定"].sum()}/{len(df)}件',
          delta=f'{df["次回アポ確定"].mean()*100:.0f}%', delta_color='off')

st.divider()

# ════════════════════════════════════════════════════════
# CA別比較チャート
# ════════════════════════════════════════════════════════
if sel_ca == '全員' and df['CA'].nunique() > 1:
    st.markdown('<div class="section-title">👥 CA別比較</div>', unsafe_allow_html=True)

    tab_score, tab_drill, tab_kpi = st.tabs(['総合スコア', '深掘り指標', '行動指標'])

    ca_grp = df.groupby('CA').agg(
        件数=('総合スコア','count'),
        平均スコア=('総合スコア','mean'),
        意向=('意向','mean'), 適正=('適正','mean'),
        条件=('条件','mean'), 認識統一=('認識統一','mean'), 気づき=('気づき','mean'),
        感情スルー率=('感情スルー率','mean'),
        縦深掘り=('縦深掘り最大','mean'),
        深掘り_価値観=('深掘り_価値観','mean'),
        バックトラッキング=('バックトラッキング','mean'),
        自己開示=('自己開示回数','mean'),
        ポジティブ反応=('ポジティブ反応','mean'),
        フィラー=('フィラー回数','mean'),
        発話比率=('求職者発話比率','mean'),
    ).round(2).reset_index()

    with tab_score:
        c1, c2 = st.columns([1,1])
        with c1:
            st.bar_chart(ca_grp.set_index('CA')[['意向','適正','条件','認識統一','気づき']])
        with c2:
            st.dataframe(
                ca_grp[['CA','件数','平均スコア','意向','適正','条件','認識統一','気づき']]
                .sort_values('平均スコア', ascending=False),
                use_container_width=True, hide_index=True)

    with tab_drill:
        c1, c2 = st.columns([1,1])
        with c1:
            st.bar_chart(ca_grp.set_index('CA')[['縦深掘り','深掘り_価値観','バックトラッキング','自己開示']])
        with c2:
            st.dataframe(
                ca_grp[['CA','感情スルー率','縦深掘り','深掘り_価値観','バックトラッキング','自己開示']]
                .sort_values('感情スルー率'),
                use_container_width=True, hide_index=True)

    with tab_kpi:
        c1, c2 = st.columns([1,1])
        with c1:
            st.bar_chart(ca_grp.set_index('CA')[['ポジティブ反応','発話比率']])
        with c2:
            st.dataframe(
                ca_grp[['CA','ポジティブ反応','フィラー','発話比率']]
                .sort_values('ポジティブ反応', ascending=False),
                use_container_width=True, hide_index=True)

    st.divider()

GRADE_COLOR = {'S':'#1a5276','A':'#1e8449','B':'#2471a3','C':'#d35400','D':'#c0392b'}

AXES = ['意向','適正','条件','認識統一','気づき']
AXIS_SHORT = {'意向':'意向把握','適正':'適正把握','条件':'条件把握',
              '認識統一':'認識統一','気づき':'気づき付与'}
AXIS_DEF = {
    '意向':    '価値観・やりがいを引き出し、応募企業に固執させず意向を広げられたか',
    '適正':    '経験・強み・適性を具体的に把握できたか',
    '条件':    'Must/Betterを確認し、期待値を調整できたか',
    '認識統一':'価値観・強みの要約への同意＋今後のキャリアの方向性への合意が取れたか',
    '気づき':  '他の選択肢・新しい可能性に気づかせられたか',
}
SCORE_LADDER = [
    '**0点** ＝ 未実施。その観点に触れていない',
    '**1点** ＝ 触れたが浅い・一方的な説明だけで終わった',
    '**2点** ＝ 把握できたが、本人の確認・同意が弱い（「はい」止まり）',
    '**3点** ＝ 根拠を引き出し、本人の明示的な同意・反応まで取れた',
]
SCORE_OPT_LABEL = {0:'0 未実施', 1:'1 浅い', 2:'2 把握', 3:'3 確認済み'}

# セルフチェックを一度だけ読み込み（CA別集計・詳細表示で共用）
selfchecks = load_selfchecks()

# ════════════════════════════════════════════════════════
# 自己評価 vs AI のズレ（CA別）
# ════════════════════════════════════════════════════════
if selfchecks:
    gap_rows = []
    for ref, sc in selfchecks.items():
        match = df_all[df_all['_file'] == ref]
        if match.empty:
            continue
        r = match.iloc[0]
        ss = sc.get('self_scores', {})
        for ax in AXES:
            gap_rows.append({
                'CA':   r['CA'],
                '_ref': ref,
                'ズレ':  ss.get(ax, 0) - int(r[ax]),
            })
    if gap_rows:
        gdf = pd.DataFrame(gap_rows)
        agg = gdf.groupby('CA').agg(
            自己採点件数=('_ref', 'nunique'),
            平均ズレ=('ズレ', 'mean'),
            平均絶対ズレ=('ズレ', lambda s: s.abs().mean()),
        ).round(2).reset_index().sort_values('平均ズレ', ascending=False)

        def _tendency(v):
            if v >= 0.5:  return '🔴 自己評価が甘い（過大）'
            if v <= -0.5: return '🔵 自己評価が辛い（過小）'
            return '🟢 AIとほぼ一致'
        agg['傾向'] = agg['平均ズレ'].apply(_tendency)

        st.markdown('<div class="section-title">🪞 自己評価とAIのズレ（CA別）'
                    '　※自己採点を入力した面談のみ集計</div>', unsafe_allow_html=True)
        st.caption('「平均ズレ」がプラス＝自分を高く評価しがち / マイナス＝低く評価しがち。'
                   'ズレが大きいほど、自分の面談を客観視できていない可能性があります。')
        st.dataframe(
            agg[['CA','自己採点件数','平均ズレ','平均絶対ズレ','傾向']],
            use_container_width=True, hide_index=True)
        st.divider()

# ════════════════════════════════════════════════════════
# 面談一覧（チェックボックス選択）
# ════════════════════════════════════════════════════════
st.markdown('<div class="section-title">📋 面談一覧　✅ チェックした候補者の詳細を下に表示</div>',
            unsafe_allow_html=True)

# ソート
sort_col = st.radio('並び替え', ['総合スコア','CA','グリップ','感情スルー率','フィラー回数','ポジティブ反応'],
                    horizontal=True)
sort_asc = st.toggle('昇順', value=False)
df_sorted = df.sort_values(sort_col, ascending=sort_asc).reset_index(drop=True)

# session_stateでチェック状態を保持
if 'checked_paths' not in st.session_state:
    st.session_state['checked_paths'] = set()

display_cols = ['CA','グリップ','候補者','面談種別','グレード','総合スコア',
                '意向','適正','条件','認識統一','気づき',
                '感情スルー率','縦深掘り最大','ポジティブ反応','フィラー回数','求職者発話比率']

# チェックボックス列を追加
df_edit = df_sorted[display_cols].copy()
df_edit.insert(0, '選択', df_sorted['_path'].isin(st.session_state['checked_paths']))

edited = st.data_editor(
    df_edit,
    use_container_width=True,
    hide_index=True,
    column_config={
        '選択':         st.column_config.CheckboxColumn('選択', default=False),
        '総合スコア':   st.column_config.ProgressColumn('総合スコア', min_value=0, max_value=15, format='%d'),
        '意向':         st.column_config.NumberColumn('意向', format='%d/3'),
        '適正':         st.column_config.NumberColumn('適正', format='%d/3'),
        '条件':         st.column_config.NumberColumn('条件', format='%d/3'),
        '認識統一':     st.column_config.NumberColumn('認識統一', format='%d/3'),
        '気づき':       st.column_config.NumberColumn('気づき', format='%d/3'),
        '感情スルー率': st.column_config.NumberColumn('感情スルー率', format='%d%%'),
        '求職者発話比率':st.column_config.NumberColumn('発話比率', format='%d%%'),
    }
)

# チェック状態をsession_stateに反映
checked_paths = set()
for i, row in edited.iterrows():
    if row['選択']:
        checked_paths.add(df_sorted.at[i, '_path'])
st.session_state['checked_paths'] = checked_paths

selected_rows = df_sorted[df_sorted['_path'].isin(checked_paths)]

col_info1, col_info2 = st.columns([3,1])
col_info1.caption(f'チェック済み：{len(selected_rows)}件 / 表示中：{len(df_sorted)}件')
if st.button('チェックをすべて外す', key='clear_check'):
    st.session_state['checked_paths'] = set()
    st.rerun()

if selected_rows.empty:
    st.info('👆 一覧の「選択」列にチェックを入れると詳細が表示されます')
    st.stop()

st.divider()

# ════════════════════════════════════════════════════════
# 表示する分析セクションの選択
# ════════════════════════════════════════════════════════
st.markdown('<div class="section-title">🔎 選択した候補者の詳細分析</div>', unsafe_allow_html=True)

ANALYSIS_OPTIONS = ['ルーブリック採点','改善ポイント','感情深掘り分析','自己開示分析','バックトラッキング分析','フレーズ集']
show_sections = st.multiselect(
    '表示する分析セクションを選択',
    ANALYSIS_OPTIONS,
    default=ANALYSIS_OPTIONS,
    help='必要な分析だけ選んで表示できます'
)

# ════════════════════════════════════════════════════════
# 選択した候補者ごとに詳細を表示
# ════════════════════════════════════════════════════════
AXES = ['意向','適正','条件','認識統一','気づき']
AXIS_LABELS = {
    '意向':    '意向把握（価値観を引き出し、意向を広げられたか）',
    '適正':    '適正把握（経験・強みを把握できたか）',
    '条件':    '条件把握（Must/Betterで確認・期待値調整できたか）',
    '認識統一':'認識統一（価値観・今後の方向性に本人の同意を得られたか）',
    '気づき':  '気づき付与（他の選択肢・新しい可能性に気づかせられたか）',
}

for loop_i, (_, sel_row) in enumerate(selected_rows.iterrows()):
    selected_path = sel_row['_path']
    # Google DriveのファイルIDかローカルパスかを判定
    if Path(selected_path).exists():
        d = json.loads(Path(selected_path).read_text(encoding='utf-8'))
    else:
        d = gdrive_download_json(selected_path)
    gd = d.get('grip_drivers', {})
    bh = d.get('behaviors', {})
    ov = d.get('overall', {})

    grip  = d.get('grip','X')
    total_score = sum(gd.get(ax,{}).get('score',0) for ax in AXES)
    grade = ov.get('grade','') or ('S' if total_score>=13 else 'A' if total_score>=10 else 'B' if total_score>=7 else 'C' if total_score>=4 else 'D')
    gc    = GRADE_COLOR.get(grade,'#555')
    uid = f'{loop_i}_{sel_row.get("_file","").replace(".","_")}'
    ref_file = sel_row.get('_file', '')

    # ════ セルフチェックのゲート（AI評価を見る前に自己採点） ════
    sc      = selfchecks.get(ref_file)
    skipped = st.session_state.get(f'sc_skip_{ref_file}', False)
    show_ai = (sc is not None) or skipped

    if not show_ai:
        # AI評価を隠し、記憶が新しいうちに自己採点を促す
        st.markdown(
            f'<div style="background:#555;color:white;padding:14px 20px;border-radius:10px;margin:16px 0 8px 0">'
            f'<span style="font-size:1rem">'
            f'CA: {d.get("ca","")}　/　グリップ: {grip}　/　{d.get("candidate","")}　/　{d.get("meeting_type","")}'
            f'</span></div>', unsafe_allow_html=True)
        st.info('🪞 **まず自己採点を。** AI評価を見る前に、自分の面談を5軸で採点してください。'
                '「自分の感覚」と「AIの客観評価」のズレが、一番の伸びしろになります。')
        with st.form(key=f'selfcheck_form_{uid}'):
            with st.expander('📖 採点の目安（0〜3点の意味）を見る', expanded=False):
                for line in SCORE_LADDER:
                    st.markdown(f'- {line}')
                st.caption('※ 各軸とも共通の基準です。下の各軸の説明と照らして採点してください。')

            st.markdown('**① この面談、自分では何点だった？（各0〜3点・0.5刻み）**')
            sc_scores = {}
            for ax in AXES:
                st.markdown(f'**{AXIS_SHORT[ax]}**　<small style="color:#888">{AXIS_DEF[ax]}</small>',
                            unsafe_allow_html=True)
                sc_scores[ax] = st.slider(
                    AXIS_SHORT[ax], min_value=0.0, max_value=3.0, value=2.0, step=0.5,
                    key=f'scs_{uid}_{ax}', label_visibility='collapsed')

            st.markdown('**② できたと思う行動にチェック**')
            bcols = st.columns(2)
            bchecks = {
                '感情ワードを拾って深掘りした':           bcols[0].checkbox('感情ワードを拾って深掘りした', key=f'bc1_{uid}'),
                '強みを言語化して返した':                 bcols[1].checkbox('強みを言語化して返した', key=f'bc2_{uid}'),
                '応募企業に固執させず他の選択肢に触れた': bcols[0].checkbox('応募企業に固執させず他の選択肢に触れた', key=f'bc5_{uid}'),
                '求職者の発言を要約して同意を取った':     bcols[1].checkbox('求職者の発言を要約して同意を取った', key=f'bc6_{uid}'),
                '沈黙を恐れず考える間を与えた':           bcols[0].checkbox('沈黙を恐れず考える間を与えた', key=f'bc7_{uid}'),
                'MUST提案をした':                         bcols[1].checkbox('MUST提案をした', key=f'bc3_{uid}'),
                '次回アポを確定した':                     bcols[0].checkbox('次回アポを確定した', key=f'bc4_{uid}'),
            }

            st.markdown('**③ 今日の面談の振り返り**')
            best_self = st.text_input('自分で「ここは良かった」と思う点', key=f'best_{uid}',
                                      placeholder='例：価値観をしっかり引き出せた')
            next_one  = st.text_area(
                '次の面談で試したいこと（1つ）', key=f'no_{uid}',
                placeholder='例：感情ワードが出たら必ず「それってどんな気持ちでしたか？」と返す')
            c_submit, c_skip = st.columns([2,1])
            do_save = c_submit.form_submit_button('✅ 採点を保存してAI評価を見る', use_container_width=True, type='primary')
            do_skip = c_skip.form_submit_button('⏭ スキップ', use_container_width=True)
        if do_save:
            try:
                save_selfcheck(ref_file, d.get('ca',''), grip, d.get('candidate',''),
                               d.get('meeting_type',''), sc_scores, bchecks, next_one, best_self)
                load_selfchecks.clear()
            except Exception as e:
                st.warning(f'保存に失敗しました（AI評価は表示します）: {e}')
                st.session_state[f'sc_skip_{ref_file}'] = True
            st.session_state['checked_paths'] = checked_paths
            st.rerun()
        if do_skip:
            st.session_state[f'sc_skip_{ref_file}'] = True
            st.session_state['checked_paths'] = checked_paths
            st.rerun()
        continue  # AI評価は自己採点 or スキップ後のrerunで表示

    # ── 候補者ヘッダー（AI評価） ──
    st.markdown(
        f'<div style="background:{gc};color:white;padding:14px 20px;border-radius:10px;margin:16px 0 8px 0">'
        f'<span style="font-size:2rem;font-weight:bold">{grade}</span>'
        f'&nbsp;&nbsp;<span style="font-size:1rem">'
        f'CA: {d.get("ca","")}　/　グリップ: {grip}　/　{d.get("candidate","")}　/　{d.get("meeting_type","")}'
        f'</span>'
        + (f'<br><small>{ov.get("grade_reason","")}</small>' if ov.get('grade_reason') else '')
        + '</div>', unsafe_allow_html=True)

    # ── 自己評価 vs AI のズレ ──
    if sc is not None:
        ss = sc.get('self_scores', {})
        gap_html = ('<table style="width:100%;border-collapse:collapse;font-size:0.9rem">'
                    '<tr style="background:#1F3864;color:white">'
                    '<th style="padding:6px;text-align:left">評価軸</th>'
                    '<th style="padding:6px">自己</th><th style="padding:6px">AI</th>'
                    '<th style="padding:6px;text-align:left">ズレ</th></tr>')
        def _fine(info):
            v = info.get('score_fine')
            if v is None:
                v = info.get('score', 0)
            try:
                return float(v)
            except (TypeError, ValueError):
                return float(info.get('score', 0) or 0)
        g = lambda v: f'{v:g}'   # 2.0→"2"、2.5→"2.5"
        for ax in AXES:
            ai_s   = _fine(gd.get(ax, {}))
            self_s = float(ss.get(ax, 0) or 0)
            diff   = round(self_s - ai_s, 1)
            if diff >= 0.5:
                tag = f'<span style="color:#c0392b">+{g(diff)} 自分が高め</span>'; bg = '#fcecea'
            elif diff <= -0.5:
                tag = f'<span style="color:#2471a3">{g(diff)} 自分が低め</span>'; bg = '#EBF5FB'
            else:
                tag = f'<span style="color:#1e8449">±{g(abs(diff))} ほぼ一致</span>'; bg = '#e2efda'
            gap_html += (f'<tr style="background:{bg}"><td style="padding:6px">{AXIS_SHORT[ax]}</td>'
                         f'<td style="padding:6px;text-align:center">{g(self_s)}</td>'
                         f'<td style="padding:6px;text-align:center">{g(ai_s)}</td>'
                         f'<td style="padding:6px">{tag}</td></tr>')
        gap_html += '</table>'
        # ── 4象限に分類（ズレ駆動FBの中核） ──
        blind, under, cweak, cstrong = [], [], [], []
        for ax in AXES:
            ai_s   = _fine(gd.get(ax, {}))
            self_s = float(ss.get(ax, 0) or 0)
            diff   = round(self_s - ai_s, 1)
            if   diff >= 0.5:   blind.append((ax, self_s, ai_s, diff))     # 過信＝盲点
            elif diff <= -0.5:  under.append((ax, self_s, ai_s, diff))     # 過小評価
            elif ai_s >= 2:     cstrong.append((ax, self_s, ai_s))         # 共通強み
            else:               cweak.append((ax, self_s, ai_s))           # 共通課題
        blind.sort(key=lambda x: -x[3])

        with st.expander('🪞 自己評価とAIのズレ → あなた専用FB', expanded=True):
            st.markdown(gap_html, unsafe_allow_html=True)

            # 今日の最優先
            if blind:
                bax = blind[0]
                st.error(f'🎯 **今日の最優先：{AXIS_SHORT[bax[0]]}** — 最大の盲点'
                         f'（自分{bax[1]:g}点 / AI{bax[2]:g}点）。下で根拠を確認してください。')
            elif cweak:
                wax = sorted(cweak, key=lambda x: x[2])[0]
                st.warning(f'🎯 **今日の最優先：{AXIS_SHORT[wax[0]]}** — 自他ともに課題と認識している軸です。')
            else:
                st.success('🎉 自己評価とAI評価がほぼ一致。自分の面談を客観視できています。強みを伸ばしましょう。')

            # 🔴 盲点（自分は高評価、AIは低評価）＝最重要
            for ax, s, a, _ in blind:
                info = gd.get(ax, {})
                ev   = info.get('evidence') or []
                html = (f'<div class="emotion-miss"><b>🔴 盲点：{AXIS_SHORT[ax]}</b>'
                        f'（自分{s:g}点 / AI{a:g}点）<br>'
                        f'<small>あなたは「できた」と感じたが、AIはここを弱点と評価。'
                        f'＝気づけていない伸びしろ。</small>')
                if info.get('weakness'):    html += f'<br>📌 <b>AIが見た弱み：</b>{info["weakness"]}'
                if ev:                      html += f'<br><small>根拠：{sq(ev[0])}</small>'
                if info.get('next_action'): html += f'<br>🚀 <b>次の一手：</b>{info["next_action"]}'
                html += '</div>'
                st.markdown(html, unsafe_allow_html=True)

            # 🟠 共通課題（自他ともに低い）
            for ax, s, a in sorted(cweak, key=lambda x: x[2]):
                info = gd.get(ax, {})
                html = (f'<div style="background:#FEF9E7;border-left:4px solid #F39C12;'
                        f'padding:8px 12px;border-radius:4px;margin:6px 0">'
                        f'<b>🟠 共通課題：{AXIS_SHORT[ax]}</b>（自分{s:g}点 / AI{a:g}点）<br>'
                        f'<small>自覚あり。あとは行動に移すだけ。</small>')
                if info.get('next_action'): html += f'<br>🚀 {info["next_action"]}'
                html += '</div>'
                st.markdown(html, unsafe_allow_html=True)

            # 🔵 過小評価（自分は低評価、AIは高評価）＝自信に変える
            for ax, s, a, _ in under:
                info = gd.get(ax, {})
                html = (f'<div class="bt-hit"><b>🔵 過小評価：{AXIS_SHORT[ax]}</b>'
                        f'（自分{s:g}点 / AI{a:g}点）<br>'
                        f'<small>実はできています。自信を持って再現を。</small>')
                if info.get('strength'): html += f'<br>💪 {info["strength"]}'
                html += '</div>'
                st.markdown(html, unsafe_allow_html=True)

            # 🟢 共通強み
            if cstrong:
                names = '、'.join(AXIS_SHORT[ax] for ax, _, _ in cstrong)
                st.markdown(f'<div class="emotion-hit">🟢 <b>共通の強み：</b>{names}'
                            f' — 自己評価通り。この型を継続！</div>', unsafe_allow_html=True)

            # ── 行動チェックの事実ズレ（自己申告 vs AI記録） ──
            bchk = sc.get('behavior_checks') or {}
            fact_map = {'MUST提案をした': bh.get('MUST提案', False),
                        '次回アポを確定した': bh.get('次回アポ確定', False)}
            mismatches = [(lbl, bchk.get(lbl), aiv) for lbl, aiv in fact_map.items()
                          if bchk.get(lbl) is not None and bchk.get(lbl) != aiv]
            if mismatches:
                msg = '⚠️ **行動の認識ズレ**（自己申告とAIの記録が不一致）<br>'
                for lbl, selfv, aiv in mismatches:
                    msg += (f'・「{lbl}」→ あなた：{"✅した" if selfv else "していない"} / '
                            f'AI記録：{"あり" if aiv else "なし"}<br>')
                st.markdown(f'<div class="emotion-miss">{msg}</div>', unsafe_allow_html=True)

            # 本人の振り返りメモ
            if sc.get('best_self'):
                st.markdown(f'**🙌 あなたが「良かった」と書いた点：** {sc["best_self"]}')
            if sc.get('next_one_thing'):
                st.markdown(f'**🚀 あなたが「次に試す」と書いたこと：** {sc["next_one_thing"]}')

    # KPI行
    k1,k2,k3,k4,k5,k6 = st.columns(6)
    k1.metric('総合スコア',    f'{total_score}/15')
    k2.metric('求職者発話比率', f'{round(bh.get("求職者発話比率",0)*100)}%')
    k3.metric('感情スルー率',   f'{bh.get("感情スルー率",0)}%')
    k4.metric('縦深掘り最大',   f'{bh.get("縦深掘り最大",0)}回')
    k5.metric('ポジティブ反応', f'{bh.get("ポジティブ反応",0)}回')
    k6.metric('フィラー回数',   f'{bh.get("フィラー回数",0)}回')

    # ── ルーブリック採点 ──────────────────────────────────
    if 'ルーブリック採点' in show_sections:
        with st.expander('🎯 ルーブリック採点（5軸）', expanded=True):
            for ax in AXES:
                info = gd.get(ax, {})
                s    = info.get('score', 0)
                cls  = 'score-high' if s>=2 else ('score-mid' if s==1 else 'score-low')
                bar  = '█'*s + '░'*(3-s)
                ev   = info.get('evidence', [])
                na   = info.get('next_action','')
                st.markdown(
                    f'<div class="kpi-card {cls}">'
                    f'<b>{AXIS_LABELS[ax]}</b> &nbsp; <code>{bar}</code> <b>{s}/3</b>'
                    + (f'<br>💪 {info["strength"]}' if info.get('strength') else '')
                    + (f'<br>📌 {info["weakness"]}' if info.get('weakness') else '')
                    + (f'<br><small style="color:#595959">根拠：{sq(ev[0])}</small>' if ev else '')
                    + (f'<br>🚀 <b>次のアクション：</b>{na}' if na else '')
                    + '</div>', unsafe_allow_html=True)
            st.markdown('**📊 行動指標**')
            bh_disp = [
                ('求職者発話比率', f'{round(bh.get("求職者発話比率",0)*100)}%', '目標40%以上', bh.get('求職者発話比率',0)>=0.4),
                ('フィラー回数',   f'{bh.get("フィラー回数",0)}回',             '目標30回以下', bh.get('フィラー回数',0)<=30),
                ('ポジティブ反応', f'{bh.get("ポジティブ反応",0)}回',           '目標5回以上',  bh.get('ポジティブ反応',0)>=5),
                ('深掘り_価値観',  f'{bh.get("深掘り_価値観",0)}回',            '目標4回以上',  bh.get('深掘り_価値観',0)>=4),
                ('バックトラッキング',f'{bh.get("バックトラッキング",0)}回',     '目標5回以上',  bh.get('バックトラッキング',0)>=5),
                ('縦深掘り最大',   f'{bh.get("縦深掘り最大",0)}回連続',         '目標3回以上',  bh.get('縦深掘り最大',0)>=3),
                ('自己開示回数',   f'{bh.get("自己開示回数",0)}回',             '目標2回以上',  bh.get('自己開示回数',0)>=2),
                ('感情スルー率',   f'{bh.get("感情スルー率",0)}%',              '目標50%以下',  bh.get('感情スルー率',0)<=50),
                ('MUST提案',       '✅' if bh.get('MUST提案') else '❌',         '',             bh.get('MUST提案',False)),
                ('次回アポ確定',   '✅' if bh.get('次回アポ確定') else '❌',     '',             bh.get('次回アポ確定',False)),
            ]
            cols = st.columns(5)
            for i, (label, val, target, ok_flag) in enumerate(bh_disp):
                cols[i%5].metric(f'{"✅" if ok_flag else "⚠️"} {label}', val, help=target)

    # ── 改善ポイント ──────────────────────────────────────
    if '改善ポイント' in show_sections:
        with st.expander('💬 改善ポイント', expanded=True):
            if ov.get('top_strength'):
                st.markdown('**💪 最大の強み**'); st.success(ov['top_strength'])
            if ov.get('best_exchange'):
                st.markdown('**👍 最も良かったやり取り**'); st.info(ov['best_exchange'])
            if ov.get('top_issues'):
                st.markdown('**🎯 優先改善ポイント（TOP3）**')
                for idx, issue in enumerate(ov['top_issues'][:3]):
                    if isinstance(issue, dict):
                        with st.expander(f'**{idx+1}. {issue.get("issue","")}**', expanded=(idx==0)):
                            if issue.get('detail'): st.write('📌 ' + issue['detail'])
                            if issue.get('fix'):    st.code(issue['fix'], language=None)
            if ov.get('missed_moment'):
                st.markdown('**📌 最も惜しかった場面**'); st.error(ov['missed_moment'])
            if ov.get('one_thing'):
                st.markdown('**🚀 次の面談で必ず試してほしいこと**'); st.warning(ov['one_thing'])
            if ov.get('closing_eval'):
                st.markdown('**🏁 クロージング評価**'); st.write(ov['closing_eval'])
            if not any([ov.get('top_strength'), ov.get('top_issues'), ov.get('one_thing')]):
                st.info('旧形式のデータです。再分析すると詳細な改善ポイントが追加されます。')

    # ── 感情深掘り分析 ────────────────────────────────────
    if '感情深掘り分析' in show_sections:
        eda = d.get('emotion_drill_analysis', {})
        with st.expander('🔍 感情深掘り分析', expanded=True):
            if not eda:
                st.info('データなし（再分析すると表示されます）')
            else:
                if eda.get('summary'): st.info('📊 ' + eda['summary'])
                missed = eda.get('missed_scenes') or []
                good   = eda.get('good_scenes') or []
                if missed:
                    st.markdown('**❌ スルーした場面 → こう返すべきだった**')
                    for ms in missed:
                        col_l, col_r = st.columns(2)
                        with col_l:
                            st.markdown(
                                f'<div class="emotion-miss">'
                                f'<small>求職者</small><br><b>{sq(ms.get("cd_text",""))}</b><br>'
                                f'<small>↓ 実際のCAの返し</small><br>'
                                f'<span style="color:#c00000">{sq(ms.get("ca_actual",""))}</span>'
                                f'</div>', unsafe_allow_html=True)
                        with col_r:
                            st.markdown(
                                f'<div class="emotion-hit">'
                                f'<small>💡 こう返すと深掘りできた</small><br>'
                                f'<b style="color:#375623">{sq(ms.get("ca_suggested",""))}</b><br>'
                                f'<small>→ {ms.get("why","")}</small>'
                                f'</div>', unsafe_allow_html=True)
                if good:
                    st.markdown('**✅ うまく深掘りできた場面**')
                    for gs in good:
                        st.markdown(
                            f'<div class="emotion-hit">'
                            f'<small>求職者</small>　{sq(gs.get("cd_text",""))}<br>'
                            f'<small>↓ CA（深掘り）</small><br>'
                            f'<span style="color:#375623"><b>{sq(gs.get("ca_text",""))}</b></span>'
                            f'<br><small>💡 {gs.get("why_good","")}</small>'
                            f'</div>', unsafe_allow_html=True)
                if eda.get('vertical_drill_comment'):
                    st.markdown('**🔽 縦の深掘りコメント**')
                    st.write(eda['vertical_drill_comment'])

    # ── 自己開示分析 ──────────────────────────────────────
    if '自己開示分析' in show_sections:
        sda = d.get('self_disclosure_analysis', {})
        with st.expander('🙋 自己開示の分析', expanded=True):
            if not sda:
                st.info('データなし（再分析すると表示されます）')
            else:
                if sda.get('summary'): st.write('📊 ' + sda['summary'])
                for fs in (sda.get('found_scenes') or []):
                    st.markdown(
                        f'<div class="sd-hit">'
                        f'CA：<b>{sq(fs.get("ca_text",""))}</b><br>'
                        f'<small>タイミング：{fs.get("timing_eval","")} / {fs.get("effect","")}</small>'
                        f'</div>', unsafe_allow_html=True)
                for ms in (sda.get('missed_opportunities') or []):
                    st.markdown(
                        f'<div style="background:#FEF9E7;border-left:3px solid #F39C12;padding:8px;border-radius:4px;margin:4px 0">'
                        f'<small>💡 ここで自己開示できた　求職者：{sq(ms.get("cd_text",""))}</small><br>'
                        f'→ <b>{sq(ms.get("ca_suggested",""))}</b>'
                        f'</div>', unsafe_allow_html=True)
                if sda.get('advice'): st.caption('📌 ' + sda['advice'])

    # ── バックトラッキング分析 ────────────────────────────
    if 'バックトラッキング分析' in show_sections:
        bta = d.get('backtrack_analysis', {})
        with st.expander('🔁 バックトラッキングの分析', expanded=True):
            if not bta:
                st.info('データなし（再分析すると表示されます）')
            else:
                if bta.get('summary'): st.write('📊 ' + bta['summary'])
                for fs in (bta.get('found_scenes') or []):
                    st.markdown(
                        f'<div class="bt-hit">'
                        f'<small>参照：{sq(fs.get("referenced_cd",""))}</small><br>'
                        f'CA：<b>{sq(fs.get("ca_text",""))}</b><br>'
                        f'<small>評価：{fs.get("effect","")}</small>'
                        f'</div>', unsafe_allow_html=True)
                for ms in (bta.get('missed_opportunities') or []):
                    st.markdown(
                        f'<div style="background:#FEF9E7;border-left:3px solid #F39C12;padding:8px;border-radius:4px;margin:4px 0">'
                        f'<small>💡 このキーワードを引用できた：{sq(ms.get("cd_keyword",""))}</small><br>'
                        f'→ <b>{sq(ms.get("ca_suggested",""))}</b>'
                        f'</div>', unsafe_allow_html=True)
                if bta.get('advice'): st.caption('📌 ' + bta['advice'])

    # ── フレーズ集 ────────────────────────────────────────
    if 'フレーズ集' in show_sections:
        phrases = d.get('next_phrases', [])
        with st.expander('🗣️ 次の面談で使えるフレーズ集', expanded=True):
            if not phrases:
                st.info('データなし（再分析すると表示されます）')
            else:
                for ph in phrases:
                    if not ph.get('phrase'): continue
                    st.markdown(
                        f'<div style="margin:8px 0">'
                        f'<small style="color:#888;background:#f0f0f0;padding:2px 8px;border-radius:10px">'
                        f'{ph.get("situation","")}</small><br>'
                        f'<div class="phrase-box">{sq(ph.get("phrase",""))}</div>'
                        f'<small style="color:#555">💡 {ph.get("why","")}</small>'
                        f'</div>', unsafe_allow_html=True)

    # ── アクション（Word / Slack / 再分析） ──────────────
    st.markdown('<div class="section-title">⚡ アクション</div>', unsafe_allow_html=True)
    act1, act2, act3 = st.columns(3)

    with act1:
        st.markdown('**📄 Word文書として出力**')
        if st.button('Word文書を生成', key=f'gen_word_{uid}', use_container_width=True):
            with st.spinner('生成中...'):
                try:
                    docx_bytes = generate_word_doc(d)
                    fname = f'面談分析_{d.get("ca","")}_{d.get("candidate","")}_{d.get("meeting_type","")}.docx'
                    st.download_button('📥 ダウンロード', data=docx_bytes, file_name=fname,
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        use_container_width=True, key=f'dl_{uid}')
                except Exception as e:
                    st.error(f'生成失敗：{e}')

    with act2:
        st.markdown('**💬 Slackに送信**')
        if st.button('Slackに送信', key=f'send_slack_{uid}', use_container_width=True):
            with st.spinner('送信中...'):
                ok_flag, msg = send_to_slack(SLACK_WEBHOOK, d)
                if ok_flag: st.success('✅ 送信しました')
                else:        st.error(f'送信失敗：{msg}')

    with act3:
        st.markdown('**🔄 再分析する**')
        is_old   = not d.get('overall')
        utt_file = core.find_utterances_file(
            d.get('ca',''), d.get('grip','X'),
            d.get('candidate',''), d.get('meeting_type',''))
        # ローカルになければGoogle Driveから検索
        if not utt_file:
            safe_grip = d.get('grip','X') if d.get('grip','X') != '未入力' else 'X'
            utt_filename = f"{d.get('ca','')}_{safe_grip}_{d.get('candidate','')}_{d.get('meeting_type','')}.json"
            drive_utts = list_json_files(subfolder="utterances")
            matched = [f for f in drive_utts if f["name"] == utt_filename]
            utt_file = matched[0]["id"] if matched else None
        if is_old:
            st.caption('⚠️ 旧形式。再分析で詳細分析が追加されます')
        if not utt_file:
            st.warning('文字起こしデータが見つかりません')
        else:
            api_key = os.environ.get('ANTHROPIC_API_KEY', '')
            if not api_key:
                st.error('APIキー未設定')
            else:
                btn_label = '🔄 旧データを再分析' if is_old else '🔄 再分析して上書き'
                if st.button(btn_label, key=f'reanalyze_{uid}', use_container_width=True, type='primary'):
                    client   = anthropic.Anthropic(api_key=api_key)
                    if Path(str(utt_file)).exists():
                        utt_data = json.loads(Path(str(utt_file)).read_text(encoding='utf-8'))
                    else:
                        utt_data = gdrive_download_json(str(utt_file))
                    utterances_data = utt_data.get('utterances', [])
                    ca   = d.get('ca','')
                    cand = d.get('candidate','')
                    fmt  = utt_data.get('format','docx')

                    with st.spinner('🤖 スコアリング中...'):
                        new_score = core.score_with_claude(utterances_data, ca, cand, fmt, client)

                    if not new_score.get('overall'):
                        st.error('スコアリングに失敗しました。もう一度お試しください。')
                    else:
                        with st.spinner('🔍 深掘り分析中...'):
                            new_deep = core.deep_analysis_with_claude(utterances_data, ca, cand, client)

                        bh_existing = {k: v for k, v in d.get('behaviors',{}).items()
                                       if not isinstance(v, (list, bool))}
                        json_path = core.save_analysis(ca, d.get('grip','X'), cand,
                                           d.get('meeting_type',''), fmt,
                                           bh_existing, new_score, new_deep)
                        # Google Drive にも上書き保存
                        try:
                            from gdrive import upload_json
                            import json as _json
                            new_data = _json.loads(json_path.read_text(encoding='utf-8'))
                            upload_json(json_path.name, new_data, subfolder="json")
                        except Exception as e:
                            st.warning(f'Drive保存に失敗しました（ローカルには保存済み）: {e}')
                        # サマリーも再構築してキャッシュクリア
                        with st.spinner('サマリーを更新中...'):
                            rebuild_summary()
                        st.session_state['checked_paths'] = checked_paths
                        load_all_records.clear()
                        st.success('✅ 再分析完了！')
                        st.rerun()

    st.divider()
